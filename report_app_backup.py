"""
E听说 成效报告对话系统
上传Excel → 自动分析 → 大模型对话调整 → 公文格式Word导出
"""

import streamlit as st
import pandas as pd
import openpyxl
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import json
import os
import re
import sys
from datetime import datetime
from io import BytesIO

sys.path.insert(0, os.path.dirname(__file__))

# ─────────────────────────────────────────────
# 1. 数据处理
# ─────────────────────────────────────────────

def parse_class_overview(file_obj) -> pd.DataFrame:
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active
    data = []
    headers = [cell.value for cell in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in row):
            data.append(dict(zip(headers, row)))
    return pd.DataFrame(data)

def parse_hw_details(file_obj) -> pd.DataFrame:
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active
    data = []
    headers = [cell.value for cell in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in row):
            data.append(dict(zip(headers, row)))
    return pd.DataFrame(data)

def analyze_data(class_df: pd.DataFrame, hw_df: pd.DataFrame) -> dict:
    results = {}

    # 三、激活概况
    results['schools'] = int(class_df['学校名称'].nunique())
    results['classes'] = int(class_df['班级id'].nunique())
    results['total_students'] = int(class_df['总学生数'].sum())
    results['school_name'] = str(class_df['学校名称'].iloc[0]) if len(class_df) > 0 else '未知学校'

    # 四、应用情况
    hw_df['大类'] = hw_df['作业路径'].apply(
        lambda x: str(x).split('-')[0] if pd.notna(x) and '-' in str(x) else '其他'
    )
    cat_counts = hw_df['大类'].value_counts().to_dict()
    results['category_counts'] = {k: int(v) for k, v in cat_counts.items()}
    results['total_hw'] = int(len(hw_df))

    def get_sub(path):
        if pd.isna(path) or '-' not in str(path):
            return '其他'
        parts = str(path).split('-')
        return parts[1] if len(parts) > 1 else parts[0]

    hw_df['小类'] = hw_df['作业路径'].apply(get_sub)
    sub_raw = hw_df.groupby(['大类', '小类']).size()
    results['sub_counts'] = {str(k): int(v) for k, v in sub_raw.to_dict().items()}

    results['assign_count'] = int(class_df['布置作业次数'].sum())
    results['assign_total'] = int(class_df['布置作业份数'].sum())
    results['completion_rate'] = round(float(class_df['作业完成率'].mean()) * 100, 2)
    results['self_practice'] = int(class_df['自主练习次数'].sum())
    results['vocab_practice'] = int(class_df['词汇自主练习次数'].sum())

    hw_df['月份'] = pd.to_datetime(hw_df['作业开始日期'], errors='coerce').dt.to_period('M').astype(str)
    monthly = hw_df.groupby('月份').size().to_dict()
    results['monthly_hw'] = {k: int(v) for k, v in sorted(monthly.items())}

    cat_monthly_raw = hw_df.groupby(['月份', '大类']).size().unstack(fill_value=0)
    results['cat_monthly'] = {
        m: {str(k): int(v) for k, v in cat_monthly_raw.loc[m].to_dict().items()}
        for m in results['monthly_hw'].keys()
    }

    results['category_pct'] = {
        k: round(v / results['total_hw'] * 100, 1)
        for k, v in results['category_counts'].items()
    }

    # 五、效果分析
    mock_df = hw_df[hw_df['作业路径'].fillna('').str.contains('模拟')].copy()
    mock_df['月份'] = pd.to_datetime(mock_df['作业开始日期'], errors='coerce').dt.to_period('M').astype(str)

    monthly_score_raw = mock_df.groupby('月份')['作业得分率'].mean()
    results['mock_monthly_score'] = {
        str(k): round(float(v) * 100, 2) for k, v in sorted(monthly_score_raw.to_dict().items())
    }

    grade_score_raw = mock_df.groupby(['年级', '月份'])['作业得分率'].mean()
    results['grade_monthly_score'] = {}
    for (grade, month), score in grade_score_raw.to_dict().items():
        g, m = str(grade), str(month)
        if g not in results['grade_monthly_score']:
            results['grade_monthly_score'][g] = {}
        results['grade_monthly_score'][g][m] = round(float(score) * 100, 2)

    class_month = hw_df.groupby(['班级id', '月份']).agg(
        hw_count=('作业ID', 'count'),
        avg_score=('作业得分率', 'mean')
    ).reset_index()
    if len(class_month) > 2:
        mean_x = class_month['hw_count'].mean()
        mean_y = class_month['avg_score'].mean()
        cov = ((class_month['hw_count'] - mean_x) * (class_month['avg_score'] - mean_y)).mean()
        std_x = class_month['hw_count'].std(ddof=0)
        std_y = class_month['avg_score'].std(ddof=0)
        corr = cov / (std_x * std_y) if std_x > 0 and std_y > 0 else 0
        results['correlation'] = round(float(corr), 4)
    else:
        results['correlation'] = 0.0

    # 六、案例
    class_stats_raw = hw_df.groupby(['班级id', '班级名称', '年级']).agg(
        hw_count=('作业ID', 'count'),
        avg_score=('作业得分率', 'mean')
    ).reset_index().sort_values('hw_count', ascending=False)

    results['top_classes'] = [
        {
            'class_id': str(row['班级id']),
            'class_name': str(row['班级名称']),
            'grade': str(row['年级']),
            'hw_count': int(row['hw_count']),
            'avg_score': round(float(row['avg_score']) * 100, 2)
        }
        for _, row in class_stats_raw.head(5).iterrows()
    ]

    if len(class_stats_raw) > 0:
        top_id = str(class_stats_raw.iloc[0]['班级id'])
        top_name = str(class_stats_raw.iloc[0]['班级名称'])
        top_grade = str(class_stats_raw.iloc[0]['年级'])
        top_df = hw_df[hw_df['班级id'] == top_id].copy()
        top_df['月份'] = pd.to_datetime(top_df['作业开始日期'], errors='coerce').dt.to_period('M').astype(str)
        top_monthly_raw = top_df.groupby('月份')['作业得分率'].agg(['mean', 'count'])
        results['top_class_monthly'] = {
            str(m): {'score': round(float(v['mean']) * 100, 2), 'count': int(v['count'])}
            for m, v in top_monthly_raw.to_dict('index').items()
        }
        results['top_class_name'] = top_name
        results['top_class_grade'] = top_grade

    return results

# ─────────────────────────────────────────────
# 2. 报告文本生成
# ─────────────────────────────────────────────

def generate_report_text(data: dict) -> str:
    school = data['school_name']
    months = sorted(data.get('monthly_hw', {}).keys())
    month_range = f"{min(months)} 至 {max(months)}" if months else "N/A"

    lines = [
        f"# {school} 英语AI练习产品 成效报告\n",
        f"**生成时间：{datetime.now().strftime('%Y年%m月%d日')}**\n",
        "---\n",
        "## 三、激活/应用概况\n",
        "| 指标 | 数值 |\n|------|------|\n",
        f"| 参与学校数 | {data['schools']}所 |\n",
        f"| 班级数（去重） | {data['classes']}个 |\n",
        f"| 激活学生总数 | {data['total_students']}人 |\n",
        "\n> 数据来源：班级数据总览\n\n",
        "## 四、应用情况分析\n",
        "### 4.1 训练内容/栏目介绍\n",
        "产品功能模块覆盖四大类，以「同步」类训练为主：\n\n",
        "| 大类 | 次数 | 占比 | 性质 |\n|------|------|------|------|\n",
    ]

    cat_names = {'同步': '日常训练类', '专项': '能力提升类', '模拟': '考试类',
                 '课外拓展': '拓展延伸类', '其他': '其他'}
    for cat, cnt in sorted(data.get('category_counts', {}).items(), key=lambda x: -x[1]):
        pct = data['category_pct'].get(cat, 0)
        lines.append(f"| **{cat}** | {cnt}次 | {pct}% | {cat_names.get(cat,'')} |\n")

    lines.append("\n**子类分布：**\n")
    sub_counts = data.get('sub_counts', {})
    all_cats = sorted(set(k[0] if isinstance(k, tuple) else k for k in sub_counts.keys()))
    for cat in all_cats:
        subs = [(k2, v) for k, k2 in sub_counts.items()
                if isinstance(k, tuple) and k[0] == cat and k[1] != cat]
        subs.sort(key=lambda x: -x[1])
        sub_text = '、'.join([f"{k}({v}次)" for k, v in subs[:6]])
        if sub_text:
            lines.append(f"- **{cat}类**：{sub_text}\n")
    lines.append("\n")

    lines.extend([
        "### 4.2 整体应用数据\n",
        "| 指标 | 数值 |\n|------|------|\n",
        f"| 布置作业次数 | {data['assign_count']}次 |\n",
        f"| 布置作业份数 | {data['assign_total']}份 |\n",
        f"| 作业完成率（均值） | {data['completion_rate']}% |\n",
        f"| 学生自主练习次数 | {data['self_practice']}次 |\n",
        f"| 词汇自主练习次数 | {data['vocab_practice']}次 |\n\n",
    ])

    lines.extend([
        "### 4.3 应用频次分析（按月）\n",
        "**月度作业总量：**\n\n",
        "| 月份 | 作业数 | 趋势 |\n|------|--------|------|\n",
    ])
    for i, m in enumerate(months):
        cnt = data['monthly_hw'][m]
        trend = "—" if i == 0 else ("↑" if cnt > data['monthly_hw'][months[i-1]] else "↓")
        lines.append(f"| {m} | {cnt} | {trend} |\n")

    lines.append("\n**各路径大类按月分布：**\n\n")
    cats = ['同步', '专项', '模拟', '课外拓展']
    lines.append("| 月份 | " + " | ".join(cats) + " |\n")
    lines.append("|" + "|".join(["------"] * (len(cats)+1)) + "\n")
    for m in months:
        cm = data.get('cat_monthly', {}).get(m, {})
        vals = [str(cm.get(c, 0)) for c in cats]
        lines.append(f"| {m} | " + " | ".join(vals) + " |\n")
    lines.append("\n")

    lines.extend([
        "### 4.4 应用方式分析\n",
        "| 大类 | 占比 | 核心子类 |\n|------|------|----------|\n",
    ])
    for cat, pct in sorted(data['category_pct'].items(), key=lambda x: -x[1]):
        subs = data.get('sub_counts', {})
        top_sub = sorted([(k, v) for k, v in subs.items()
                          if isinstance(k, tuple) and k[0] == cat], key=lambda x: -x[1])
        sub_text = '、'.join([k for k, v in top_sub[:3]])
        lines.append(f"| {cat} | {pct}% | {sub_text} |\n")
    lines.append("\n")

    lines.extend([
        "## 五、应用效果分析\n",
        "### 5.1 成绩数据对比（模拟类路径）\n",
    ])
    mock_scores = data.get('mock_monthly_score', {})
    if mock_scores:
        lines.extend([
            "**月均得分率：**\n\n",
            "| 月份 | 平均得分率 |\n|------|------------|\n",
        ])
        for m, score in sorted(mock_scores.items()):
            lines.append(f"| {m} | {score}% |\n")
        lines.append("\n")

    grade_scores = data.get('grade_monthly_score', {})
    if grade_scores:
        lines.append("**分年级得分率趋势：**\n\n")
        for grade, monthly in sorted(grade_scores.items()):
            vals = [f"{m}:{s}%" for m, s in sorted(monthly.items())]
            lines.append(f"- **{grade}**：{' → '.join(vals)}\n")
        lines.append("\n")

    corr = data.get('correlation', 0)
    corr_desc = "弱正相关" if corr > 0.3 else ("负相关" if corr < -0.3 else "相关性不强")
    lines.extend([
        "### 5.2 相关性分析\n",
        f"| 分析维度 | 相关系数 | 结论 |\n",
        f"|---------|---------|------|\n",
        f"| 班级月度作业频次 vs 平均得分率 | {corr} | {corr_desc} |\n\n",
    ])

    top = data.get('top_classes', [])
    lines.append("## 六、典型案例/学校分析\n")
    if top:
        lines.extend([
            f"**标杆班级：{data.get('top_class_name', top[0]['class_name'])}（{data.get('top_class_grade', '')}）**\n\n",
            f"- 总作业次数：{top[0]['hw_count']}次\n",
            f"- 平均得分率：{top[0]['avg_score']}%\n\n",
        ])
        top_monthly = data.get('top_class_monthly', {})
        if top_monthly:
            lines.extend([
                "**月度得分率走势：**\n\n",
                "| 月份 | 得分率 | 作业量 |\n|------|--------|--------|\n",
            ])
            for m, v in sorted(top_monthly.items()):
                lines.append(f"| {m} | {v['score']}% | {v['count']}次 |\n")
    lines.append("\n")

    lines.extend([
        "## 七、总结与建议\n",
        "### 7.1 主要发现\n",
        "**✅ 亮点：**\n",
        f"- 激活率高：全校{data['total_students']}名学生，{data['classes']}个班级全覆盖\n",
        "- 同步训练扎实：课文跟读/朗读为日常主力，日常开口习惯培养效果显著\n",
        f"- 词汇自主练习活跃：{data['vocab_practice']}次，自主学习意愿强\n",
    ])
    if top:
        lines.append(f"- 标杆班级表现优异：{data.get('top_class_name','')}平均得分率{top[0]['avg_score']}%\n")

    low_complete = data.get('completion_rate', 0)
    lines.append("\n**⚠️ 风险点：**\n")
    if low_complete < 40:
        lines.append(f"- 作业完成率偏低（均值{low_complete}%），大量布置未转化为实际练习\n")
    if grade_scores:
        latest_scores = {g: list(vals.values())[-1] if vals else 0
                        for g, vals in grade_scores.items()}
        min_grade = min(latest_scores.items(), key=lambda x: x[1])
        lines.append(f"- {min_grade[0]}年级得分率最低（{min_grade[1]}%），需重点关注\n")

    lines.extend([
        "\n### 7.2 建议\n",
        "| 优先级 | 建议 |\n|--------|------|\n",
        "| 🔴 高 | 提升作业完成率：通过分层任务设计，将完成率提升至50%以上 |\n",
        "| 🔴 高 | 重点关注低得分年级：进行专项教研，诊断成绩下滑根因 |\n",
        "| 🟡 中 | 增加模拟类训练频次：建议每月至少2次听说模拟 |\n",
        "| 🟡 中 | 优化假期作业设计：避免假期作业完成质量下降 |\n",
        "\n---\n",
        f"*数据周期：{month_range}*\n",
        "*数据来源：班级数据总览、作业明细*\n",
    ])
    return ''.join(lines)

# ─────────────────────────────────────────────
# 3. 图表
# ─────────────────────────────────────────────

def make_charts(data: dict) -> dict:
    charts = {}
    cat_colors = {'同步': '#4C78A8', '专项': '#F58518', '模拟': '#E45756', '课外拓展': '#72D7B8', '其他': '#999'}
    cat_monthly = data.get('cat_monthly', {})
    months = sorted(cat_monthly.keys())
    cats = ['同步', '专项', '模拟', '课外拓展']

    # 堆叠柱状图
    fig = go.Figure()
    for cat in cats:
        y = [cat_monthly.get(m, {}).get(cat, 0) for m in months]
        fig.add_trace(go.Bar(name=cat, x=months, y=y, marker_color=cat_colors.get(cat, '#999')))
    fig.update_layout(barmode='stack', title='月度作业总量趋势（按大类）',
                      xaxis_title='月份', yaxis_title='作业次数',
                      height=420, template='plotly_white',
                      legend=dict(orientation='h', yanchor='bottom', y=1.02))
    charts['monthly_trend'] = fig

    # 饼图
    cat_pct = data.get('category_pct', {})
    fig2 = go.Figure()
    fig2.add_trace(go.Pie(
        labels=list(cat_pct.keys()),
        values=list(cat_pct.values()),
        marker_colors=[cat_colors.get(c, '#999') for c in cat_pct.keys()],
        textinfo='label+percent', hole=0.35
    ))
    fig2.update_layout(title='作业类型占比分布', height=400, template='plotly_white')
    charts['category_pie'] = fig2

    # 模拟得分率趋势
    mock_scores = data.get('mock_monthly_score', {})
    if mock_scores:
        ms = sorted(mock_scores.keys())
        sc = [mock_scores[m] for m in ms]
        fig3 = go.Figure()
        fig3.add_trace(go.Scatter(x=ms, y=sc, mode='lines+markers',
                                   line=dict(color='#E45756', width=3), marker=dict(size=8)))
        fig3.update_layout(title='模拟类作业月均得分率趋势',
                           xaxis_title='月份', yaxis_title='得分率（%）',
                           height=350, template='plotly_white', yaxis=dict(range=[0, 100]))
        charts['mock_score_trend'] = fig3

    # 年级趋势
    grade_scores = data.get('grade_monthly_score', {})
    if grade_scores:
        fig4 = go.Figure()
        gc = {'六年级': '#4C78A8', '七年级': '#F58518', '八年级': '#E45756'}
        for grade, monthly in sorted(grade_scores.items()):
            ms = sorted(monthly.keys())
            sc = [monthly[m] for m in ms]
            fig4.add_trace(go.Scatter(name=grade, x=ms, y=sc,
                                       mode='lines+markers', line=dict(width=2),
                                       marker=dict(size=7, color=gc.get(grade, '#999'))))
        fig4.update_layout(title='各年级模拟类得分率月度对比',
                            xaxis_title='月份', yaxis_title='得分率（%）',
                            height=350, template='plotly_white', yaxis=dict(range=[0, 100]))
        charts['grade_score_trend'] = fig4

    # 标杆班级
    top_monthly = data.get('top_class_monthly', {})
    if top_monthly:
        tm = sorted(top_monthly.keys())
        sc = [top_monthly[m]['score'] for m in tm]
        ct = [top_monthly[m]['count'] for m in tm]
        fig5 = make_subplots(specs=[[{"secondary_y": True}]])
        fig5.add_trace(go.Scatter(x=tm, y=sc, name='得分率%', mode='lines+markers',
                                   line=dict(color='#4C78A8', width=2), marker=dict(size=8),
                                   yaxis='y'))
        fig5.add_trace(go.Bar(x=tm, y=ct, name='作业量', opacity=0.4,
                               marker_color='#ccc', yaxis='y2'))
        fig5.update_layout(title=f"标杆班级（{data.get('top_class_name','')}）月度得分率与作业量",
                           template='plotly_white', height=350,
                           legend=dict(orientation='h', yanchor='bottom', y=1.02),
                           yaxis2=dict(title_text='作业次数', overlaying='y', side='right', anchor='x', position=1.0))
        charts['top_class_trend'] = fig5

    return charts

# ─────────────────────────────────────────────
# 4. Word导出（公文格式）
# ─────────────────────────────────────────────

def export_to_docx(report_md: str) -> tuple:
    """导出为公文格式Word"""
    try:
        from docx import Document
        from docx.shared import Pt, Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        return None, "python-docx未安装"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    def set_run_fmt(run, fname, fsize, bold=False):
        run.font.name = fname
        run.font.size = Pt(fsize)
        run.font.bold = bold
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
        except Exception:
            pass

    def add_para(text, fname='仿宋_GB2312', fsize=16, bold=False,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=False):
        para = doc.add_paragraph()
        para.alignment = align
        run = para.add_run(text)
        set_run_fmt(run, fname, fsize, bold)
        if first_indent:
            para.paragraph_format.first_line_indent = Cm(0.74)
        return para

    def add_table_hdr(headers, rows_data):
        tbl = doc.add_table(rows=1+len(rows_data), cols=len(headers))
        tbl.style = 'Table Grid'
        for i, h in enumerate(headers):
            cell = tbl.rows[0].cells[i]
            cell.text = h
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.bold = True
                p.runs[0].font.size = Pt(12)
        for ri, row in enumerate(rows_data):
            for ci, val in enumerate(row):
                cell = tbl.rows[ri+1].cells[ci]
                cell.text = str(val)
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if p.runs:
                    p.runs[0].font.size = Pt(12)
        return tbl

    lines = report_md.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].strip()
        if not line or line.startswith('>') or line.startswith('*数据') or line.startswith('*报告'):
            i += 1
            continue

        if line.startswith('# ') and '成效报告' in line:
            add_para(line.replace('# ', '').strip(), '方正小标宋简体', 22, True, WD_ALIGN_PARAGRAPH.CENTER)

        elif line.startswith('## 三、') or line.startswith('## 四、') or \
             line.startswith('## 五、') or line.startswith('## 六、') or line.startswith('## 七、'):
            add_para(line.replace('## ', '').strip(), '黑体', 16, True, WD_ALIGN_PARAGRAPH.LEFT)

        elif '### 4.' in line or '### 5.' in line or '### 7.' in line:
            add_para(line.replace('### ', '').strip(), '楷体_GB2312', 16, True, WD_ALIGN_PARAGRAPH.LEFT)

        elif line.startswith('|') and '---' not in line:
            rows_data = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                if '---' not in lines[j]:
                    cells = [c.strip() for c in lines[j].strip().split('|')[1:-1]]
                    rows_data.append(cells)
                j += 1
            if rows_data:
                add_table_hdr(rows_data[0], rows_data[1:])
            i = j
            continue

        elif line.startswith('- '):
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line).lstrip('- ')
            add_para('• ' + clean, '仿宋_GB2312', 16, False, WD_ALIGN_PARAGRAPH.JUSTIFY, True)

        else:
            clean = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            if clean.strip():
                add_para(clean.strip(), '仿宋_GB2312', 16, False, WD_ALIGN_PARAGRAPH.JUSTIFY, False)
        i += 1

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, None

# ─────────────────────────────────────────────
# 5. 主应用
# ─────────────────────────────────────────────

st.set_page_config(page_title="E听说成效报告系统", page_icon="📊", layout="wide")

with st.sidebar:
    st.title("📋 使用说明")
    st.markdown("""
    **步骤：**
    1. 上传 `班级数据总览.xlsx`
    2. 上传 `作业明细.xlsx`
    3. 系统自动分析并生成报告
    4. 在「对话调整」标签页与大模型对话
    5. 导出为 公文格式Word

    **文件要求：**
    - 班级数据总览：含班级id、总学生数、布置作业次数、作业完成率等
    - 作业明细：含作业路径、得分率、开始日期等
    """)
    st.divider()
    st.caption("支持：修改结论 · 调整数据口径 · 补充分析 · 换正式语气")

    with st.expander("🔑 大模型API配置"):
        st.markdown("**选择或配置模型服务商：**")

        PROVIDERS = {
            " Minimax（海螺AI）": {
                "base_url": "https://api.minimax.chat/v",
                "model": "MiniMax-Text-01",
                "key_hint": "Bearer Token（maa-...）",
                "key_example": "maa-xxxxxxxxxxxxxxxxxxxxxxxx"
            },
            " DeepSeek": {
                "base_url": "https://api.deepseek.com",
                "model": "deepseek-chat",
                "key_hint": "API Key（sk-...）",
                "key_example": "sk-xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx"
            },
            " 硅基流动（SiliconFlow）": {
                "base_url": "https://api.siliconflow.cn/v1",
                "model": "deepseek-ai/DeepSeek-V3",
                "key_hint": "API Key（sk-...）",
                "key_example": "sk-xxxxxxxxxxxxxxxx"
            },
            " Groq": {
                "base_url": "https://api.groq.com/openai/v1",
                "model": "mixtral-8x7b-32768",
                "key_hint": "API Key（gsk_...）",
                "key_example": "gsk_xxxxxxxxxxxxxxxx"
            },
            " 自定义（OpenAI兼容）": {
                "base_url": "",
                "model": "",
                "key_hint": "API Key",
                "key_example": "sk-..."
            },
        }

        selected = st.selectbox("服务商", list(PROVIDERS.keys()), label_visibility="collapsed")
        provider = PROVIDERS[selected]

        # API Key
        api_key_input = st.text_input(
            f"API Key（{provider['key_hint']}）",
            type="password",
            placeholder=provider['key_example'],
            label_visibility="collapsed"
        )

        # Base URL（自动填入，但对自定义开放编辑）
        if selected == " 自定义（OpenAI兼容）":
            base_url_input = st.text_input(
                "Base URL",
                placeholder="https://api.openai.com/v1",
                value=""
            )
        else:
            base_url_input = provider["base_url"]
            st.text_input(
                "Base URL（自动填充，不可编辑）",
                value=base_url_input,
                disabled=True,
                label_visibility="collapsed"
            )

        # 模型名称
        if selected == " 自定义（OpenAI兼容）":
            model_input = st.text_input("模型名称", placeholder="例如：gpt-4o-mini、deepseek-chat")
        else:
            model_input = provider["model"]
            st.text_input(
                "模型（自动填充，不可编辑）",
                value=model_input,
                disabled=True,
                label_visibility="collapsed"
            )

        if api_key_input and model_input:
            os.environ["LLM_API_KEY"] = api_key_input
            os.environ["LLM_BASE_URL"] = base_url_input
            os.environ["LLM_MODEL"] = model_input
            st.success(f"✅ 已配置：{selected.strip()} · 模型：{model_input}")

        st.caption("💡 配置仅保存在本地会话，刷新页面后需重新输入")

col1, col2 = st.columns(2)
with col1:
    class_file = st.file_uploader("📁 上传「班级数据总览.xlsx」", type=['xlsx'], key="class_file")
with col2:
    hw_file = st.file_uploader("📁 上传「作业明细.xlsx」", type=['xlsx'], key="hw_file")

if class_file and hw_file:
    with st.spinner("正在分析数据，请稍候..."):
        try:
            class_df = parse_class_overview(class_file)
            hw_df = parse_hw_details(hw_file)
            data = analyze_data(class_df, hw_df)
            report_text = generate_report_text(data)
            charts = make_charts(data)
            st.session_state['data'] = data
            st.session_state['report'] = report_text
            st.session_state['charts'] = charts
            st.session_state['messages'] = []
            st.success(f"✅ 分析完成！学校：{data['school_name']}，班级：{data['classes']}个，学生：{data['total_students']}人")
        except Exception as e:
            st.error(f"❌ 数据解析出错：{e}")
            import traceback
            st.code(traceback.format_exc())
            st.stop()

if 'report' in st.session_state:
    data = st.session_state['data']
    charts = st.session_state['charts']
    report_text = st.session_state['report']

    tab1, tab2, tab3, tab4 = st.tabs(["📄 成效报告", "📈 图表分析", "💬 对话调整", "📥 导出Word"])

    with tab1:
        st.divider()
        st.markdown("### 📄 成效报告（初稿）")
        st.divider()
        st.markdown(report_text)

    with tab2:
        st.divider()
        st.markdown("### 📈 数据可视化")
        st.divider()
        chart_map = [
            ('monthly_trend', '📊 月度作业趋势（堆叠柱状图）'),
            ('category_pie', '🥧 作业类型占比'),
            ('mock_score_trend', '📉 模拟类月均得分率趋势'),
            ('grade_score_trend', '📈 各年级得分率对比'),
            ('top_class_trend', '🎯 标杆班级月度分析'),
        ]
        for key, title in chart_map:
            if key in charts:
                st.plotly_chart(charts[key], use_container_width=True)
                st.divider()

    with tab3:
        st.divider()
        st.markdown("### 💬 对话调整报告")
        st.markdown("""
        **支持的操作：**
        - 修改/补充结论（如"将七年级建议改为..."）
        - 调整数据口径或重新计算
        - 补充某班级详细分析
        - 改变建议优先级
        - 用更正式的公文语气重写某章节
        """)
        st.divider()

        if 'messages' not in st.session_state:
            st.session_state['messages'] = []

        for msg in st.session_state['messages']:
            avatar = "👤" if msg['role'] == 'user' else "🤖"
            with st.chat_message(msg['role'], avatar=avatar):
                st.markdown(msg['content'])

        user_input = st.chat_input("输入你的调整要求...")

        if user_input:
            with st.chat_message("user", avatar="👤"):
                st.markdown(user_input)
            st.session_state['messages'].append({'role': 'user', 'content': user_input})

            prompt = f"""你是一个专业的教育数据分析报告编辑助手。请根据用户指示修改报告。

**原始报告：**
---
{report_text}

**用户指示：**
{user_input}

请直接输出修改后的完整报告（Markdown格式），保持原有结构，只修改指定内容。"""

            with st.chat_message("assistant", avatar="🤖"):
                with st.spinner("大模型正在修改报告..."):
                    try:
                        import openai
                        client = openai.OpenAI(
                            api_key=os.environ.get("LLM_API_KEY", ""),
                            base_url=os.environ.get("LLM_BASE_URL") or None
                        )
                        response = client.chat.completions.create(
                            model=os.environ.get("LLM_MODEL", "gpt-4o-mini"),
                            messages=[{"role": "user", "content": prompt}],
                            max_tokens=4000,
                            temperature=0.3
                        )
                        revised = response.choices[0].message.content
                        # 去掉可能的markdown代码块包裹
                        if revised.startswith('```'):
                            lines_r = revised.split('\n')
                            revised = '\n'.join(lines_r[1:-1] if lines_r[-1] == '```' else lines_r[1:])
                        st.markdown(revised)
                        st.session_state['messages'].append({'role': 'assistant', 'content': revised})
                        st.session_state['report'] = revised
                    except Exception as e:
                        st.error(f"❌ 大模型调用失败：{e}")
                        st.info("💡 请在左侧「🔑 大模型API配置」中完成配置后重试")

    with tab4:
        st.divider()
        st.markdown("### 📥 导出为公文格式Word")
        st.markdown("""
        **导出格式（参考公文规范）：**
        - 标题：方正小标宋简体，二号，居中
        - 一级标题：黑体，三号
        - 二级标题：楷体_GB2312，三号
        - 正文：仿宋_GB2312，三号，首行缩进2格
        - 行间距：固定值31磅
        - 页边距：上3.7cm、下3.5cm、左2.8cm、右2.6cm
        """)
        st.divider()

        if st.button("📄 生成Word文档", type="primary"):
            with st.spinner("正在生成Word文档..."):
                buf, err = export_to_docx(report_text)
                if err:
                    st.error(err)
                    st.info("💡 安装 python-docx：pip install python-docx")
                else:
                    school = data.get('school_name', '学校')
                    fname = f"{school}成效报告_{datetime.now().strftime('%Y%m%d')}.docx"
                    st.download_button(
                        label=f"⬇️ 下载 {fname}",
                        data=buf.getvalue(),
                        file_name=fname,
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    st.success("✅ Word文档已生成！")

else:
    st.info("👆 请同时上传两个Excel文件开始生成报告")
    st.markdown("""
    ---
    **预期输出内容：**

    | 章节 | 内容 |
    |------|------|
    | 三、激活/应用概况 | 学校数、班级数、激活学生数 |
    | 四、应用情况分析 | 栏目介绍、应用数据、频次分析、方式分析 |
    | 五、应用效果分析 | 模拟类成绩对比、相关性分析 |
    | 六、典型案例 | 标杆班级月度趋势 |
    | 七、总结与建议 | 亮点、风险、建议 |
    """)
