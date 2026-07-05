"""
E听说 成效报告对话系统 v2.5
"""
import pandas as pd
import openpyxl
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import os, re, sys, math
from datetime import datetime
from io import BytesIO

sys.path.insert(0, os.path.dirname(__file__))

REPORT_VERSION = "v2.7"

# 联考六大题型（与联考 Word 报告一致），用于将日常练习题型对齐到联考题型
EXAM_QTYPES = ['听句子选答案', '听对话选答案', '听短文选答案', '短文朗读', '情景对话', '情景表达']

# 题型 → 通用课堂训练动作（教研增强，v2.7）；教材单元级对应需备课组对照表后引用
QT_ACTIONS = {
    '情景表达': '看图说话 / 功能句复述；课前 3 分钟情景问答，聚焦提问与建议功能句',
    '情景对话': '同桌角色扮演对练；积累场景词块（购物、问路、就医等高频场景）',
    '短文朗读': '跟读模仿＋录音自评；重点抓连读、弱读与意群停顿',
    '听短文选答案': '听前预测＋关键词捕捉训练；先审题圈定信息点再听',
    '听对话选答案': '场景推断专项；训练抓转折词（but/actually）后的关键信息',
    '听句子选答案': '辨音专项（数字、日期、同音干扰）；一句两听复述训练',
}

# 学校名称粗判（剔除培训机构等非学校主体，与 graphic_report 口径一致）
_SCHOOL_PAT = re.compile(r"(中学|学校|学园|小学|中心|附中|分校|实验)")

def _is_school(name):
    return bool(_SCHOOL_PAT.search(str(name)))

# 作业明细中语义为数值、但可能混入 '-' 等占位符的列
_HW_NUM_COLS = ('作业得分率', '100%完成学生占比', '作答学生占比', '作答学生总数', '单次作业平均耗时/min')

def _coerce_hw_numeric(hw_df):
    for c in _HW_NUM_COLS:
        if c in hw_df.columns:
            hw_df[c] = pd.to_numeric(hw_df[c], errors='coerce')
    return hw_df

def parse_class_overview(file_obj):
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active
    rows = []
    headers = [c.value for c in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in row):
            rows.append(dict(zip(headers, row)))
    return pd.DataFrame(rows)

def parse_hw_details(file_obj):
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active
    rows = []
    headers = [c.value for c in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in row):
            rows.append(dict(zip(headers, row)))
    return pd.DataFrame(rows)

def parse_question_type(file_obj):
    """解析「听说模拟班级总体情况-题型」Excel"""
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active
    rows = []
    headers = [c.value for c in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in row):
            rows.append(dict(zip(headers, row)))
    df = pd.DataFrame(rows)
    # 统一列名：得分率转为小数→百分数
    if '得分率' in df.columns:
        df['得分率'] = pd.to_numeric(df['得分率'], errors='coerce').fillna(0)
    # 年级/班级统一为字符串
    for col in ['年级', '班级']:
        if col in df.columns:
            df[col] = df[col].astype(str).str.strip()
    return df

def _split_path(path):
    if pd.isna(path) or '-' not in str(path):
        return ('其他', '其他')
    parts = str(path).split('-')
    return parts[0], parts[1] if len(parts) > 1 else parts[0]

def _format_grade_class(grade, class_name):
    return f"{str(grade)}{str(class_name)}班"

def _grade_sort_key(grade):
    text = str(grade)
    order = {
        '一年级': 1, '二年级': 2, '三年级': 3, '四年级': 4, '五年级': 5, '六年级': 6,
        '七年级': 7, '八年级': 8, '九年级': 9,
        '初一': 7, '初二': 8, '初三': 9,
        '高一': 10, '高二': 11, '高三': 12,
        '高中一年级': 10, '高中二年级': 11, '高中三年级': 12,
    }
    if text in order:
        return (order[text], text)
    m = re.search(r'\d+', text)
    return (int(m.group()) if m else 99, text)

def _sort_grades(grades):
    return sorted([str(g) for g in grades if str(g) and str(g) != 'nan'], key=_grade_sort_key)

def _format_month_label(month):
    text = str(month)
    return re.sub(r'(\d{4})-(\d{2})', r'\1/\2', text)

def _are_consecutive_months(months, min_count=3):
    parsed = []
    for month in months:
        try:
            parsed.append(pd.Period(str(month), freq='M'))
        except Exception:
            continue
    if len(parsed) < min_count:
        return False
    parsed = sorted(set(parsed))
    run = 1
    for prev, cur in zip(parsed, parsed[1:]):
        if cur - prev == 1:
            run += 1
            if run >= min_count:
                return True
        else:
            run = 1
    return False

def _grade_color(grade, index=0):
    palette = ['#4C78A8', '#F58518', '#E45756', '#54A24B', '#B279A2', '#72B7B2', '#FF9DA6', '#9D755D']
    base = {'六年级': '#4C78A8', '七年级': '#F58518', '八年级': '#E45756', '九年级': '#54A24B',
            '高一': '#4C78A8', '高二': '#F58518', '高三': '#E45756'}
    return base.get(str(grade), palette[index % len(palette)])

def analyze_data(class_df, hw_df, qt_df=None):
    results = {}
    results['schools']    = int(class_df['学校名称'].nunique())
    results['classes']    = int(class_df['班级id'].nunique())
    results['total_students'] = int(class_df['总学生数'].sum())
    results['school_name'] = str(class_df['学校名称'].iloc[0]) if len(class_df) > 0 else '未知学校'
    results['province']  = str(class_df['省份'].iloc[0]) if '省份' in class_df.columns and len(class_df) > 0 else ''
    results['city']       = str(class_df['城市'].iloc[0]) if '城市' in class_df.columns and len(class_df) > 0 else ''
    stages = sorted([str(v).strip() for v in class_df.get('学段', pd.Series(dtype=str)).dropna().unique().tolist() if str(v).strip()])
    results['stages'] = stages
    results['stage_name'] = '、'.join(stages) if stages else '未标注学段'

    hw_df = _coerce_hw_numeric(hw_df.copy())
    hw_df['大类']     = hw_df['作业路径'].apply(lambda x: _split_path(x)[0])
    hw_df['小类']     = hw_df['作业路径'].apply(lambda x: _split_path(x)[1])
    hw_df['完整路径'] = hw_df['作业路径'].fillna('')
    hw_df['月份']     = pd.to_datetime(hw_df['作业开始日期'], errors='coerce').dt.to_period('M').astype(str)

    cat_counts = hw_df['大类'].value_counts().to_dict()
    results['category_counts'] = {k: int(v) for k, v in cat_counts.items()}
    results['total_hw']        = int(len(hw_df))
    results['category_pct']   = {k: round(v / results['total_hw'] * 100, 1) for k, v in cat_counts.items()}

    sub_raw = hw_df.groupby(['大类', '小类']).size()
    results['sub_counts'] = {str((c, s)): int(v) for (c, s), v in sub_raw.to_dict().items()}

    results['assign_count']    = int(class_df['布置作业次数'].sum())
    results['assign_total']    = int(class_df['布置作业份数'].sum())
    results['completion_rate']  = round(float(pd.to_numeric(class_df['作业完成率'], errors='coerce').mean()) * 100, 2)
    results['self_practice']    = int(class_df['自主练习次数'].sum())
    results['vocab_practice']   = int(class_df['词汇自主练习次数'].sum())
    results['score_rate_avg']  = round(float(pd.to_numeric(class_df['作业得分率'], errors='coerce').mean()) * 100, 2)

    monthly = hw_df.groupby('月份').size().to_dict()
    results['monthly_hw'] = {k: int(v) for k, v in sorted(monthly.items())}

    cat_monthly = hw_df.groupby(['月份', '大类']).size().unstack(fill_value=0)
    results['cat_monthly'] = {
        m: {str(k): int(v) for k, v in cat_monthly.loc[m].to_dict().items()}
        for m in results['monthly_hw'].keys()
    }

    actual_grades = _sort_grades(hw_df['年级'].dropna().unique().astype(str).tolist())
    results['actual_grades'] = actual_grades

    grade_monthly_hw_df = hw_df.groupby(['年级', '月份']).size().unstack(fill_value=0)
    results['grade_monthly_hw'] = {
        str(g): {str(c): int(v) for c, v in row.to_dict().items()}
        for g, row in grade_monthly_hw_df.iterrows()
    }

    # 听说模拟
    mock_hw = hw_df[hw_df['完整路径'].str.contains('模拟-', na=False)].copy()
    mock_hw['月份'] = pd.to_datetime(mock_hw['作业开始日期'], errors='coerce').dt.to_period('M').astype(str)

    monthly_score = mock_hw.groupby('月份')['作业得分率'].mean()
    results['mock_hw_score_monthly'] = {str(k): round(float(v)*100, 2) for k, v in sorted(monthly_score.to_dict().items())}
    monthly_count = mock_hw.groupby('月份').size()
    results['mock_hw_count_monthly'] = {str(k): int(v) for k, v in sorted(monthly_count.to_dict().items())}

    grade_score = mock_hw.groupby(['年级', '月份'])['作业得分率'].mean()
    results['mock_hw_grade_monthly'] = {}
    for (g, m), s in grade_score.to_dict().items():
        gs, ms = str(g), str(m)
        if gs not in results['mock_hw_grade_monthly']:
            results['mock_hw_grade_monthly'][gs] = {}
        results['mock_hw_grade_monthly'][gs][ms] = round(float(s)*100, 2)

    # Pearson相关性
    class_hw = mock_hw.groupby('班级id').agg(
        avg_score=('作业得分率', 'mean'), hw_count=('作业ID', 'count')
    ).reset_index()

    class_info = {}
    for _, row in class_df.iterrows():
        cid = row['班级id']
        class_info[cid] = {
            'vocab':   row.get('词汇自主练习次数', 0) or 0,
            'self_p':  row.get('自主练习次数', 0) or 0,
            'complete':row.get('作业完成率', 0) or 0,
        }

    def pearsonr(pairs):
        if len(pairs) < 3: return 0.0, len(pairs)
        n = len(pairs)
        mx = sum(p[0] for p in pairs) / n
        my = sum(p[1] for p in pairs) / n
        cov = sum((p[0]-mx)*(p[1]-my) for p in pairs) / n
        sx = math.sqrt(sum((p[0]-mx)**2 for p in pairs) / n)
        sy = math.sqrt(sum((p[1]-my)**2 for p in pairs) / n)
        return (cov/(sx*sy) if sx*sy else 0.0, n)

    pairs_vocab, pairs_complete, pairs_self = [], [], []
    for _, row in class_hw.iterrows():
        cid = row['班级id']
        avg = float(row['avg_score']) * 100
        if cid in class_info:
            vp = class_info[cid]['vocab']
            sp = class_info[cid]['self_p']
            cr = class_info[cid]['complete']
            if vp > 0: pairs_vocab.append((vp, avg))
            if sp > 0: pairs_self.append((sp, avg))
            if cr > 0: pairs_complete.append((cr, avg))

    r_vocab,    n_v = pearsonr(pairs_vocab)
    r_self,     n_s = pearsonr(pairs_self)
    r_complete, n_c = pearsonr(pairs_complete)
    results['corr_vocab']    = (round(r_vocab,    4), n_v)
    results['corr_self']     = (round(r_self,     4), n_s)
    results['corr_complete'] = (round(r_complete, 4), n_c)

    strong = []
    if abs(r_vocab)    >= 0.4: strong.append(('词汇自主练习次数', r_vocab,    n_v))
    if abs(r_complete) >= 0.4: strong.append(('作业完成率',       r_complete, n_c))
    if abs(r_self)     >= 0.4: strong.append(('自主练习次数',    r_self,     n_s))
    results['strong_corrs'] = strong

    # ── 学生分层数据（A/B/C三层，基于听说模拟得分）────────────
    student分层 = None
    if not mock_hw.empty and '作业得分率' in mock_hw.columns and '学生id' in mock_hw.columns:
        student_score = mock_hw.groupby('学生id')['作业得分率'].mean()
        if len(student_score) >= 3:
            p70 = student_score.quantile(0.70)
            p30 = student_score.quantile(0.30)
            a_s = student_score[student_score >= p70]
            b_s = student_score[(student_score < p70) & (student_score > p30)]
            c_s = student_score[student_score <= p30]
            total_s = len(student_score)
            student分层 = {
                'n_a': len(a_s), 'n_b': len(b_s), 'n_c': len(c_s),
                'pct_a': round(len(a_s)/total_s*100, 1) if total_s>0 else 0,
                'pct_b': round(len(b_s)/total_s*100, 1) if total_s>0 else 0,
                'pct_c': round(len(c_s)/total_s*100, 1) if total_s>0 else 0,
                'avg_a': round(a_s.mean()*100, 2) if len(a_s)>0 else 0,
                'avg_b': round(b_s.mean()*100, 2) if len(b_s)>0 else 0,
                'avg_c': round(c_s.mean()*100, 2) if len(c_s)>0 else 0,
                'gap':   round((a_s.mean()-c_s.mean())*100, 2) if (len(a_s)>0 and len(c_s)>0) else 0,
            }
    results['student分层'] = student分层

    # TOP5 班级：优先选择听说模拟连续3个月有布置的班级
    class_all_hw = hw_df.groupby(['班级id', '班级名称', '年级']).size().reset_index(name='all_hw_count')
    class_mock    = mock_hw.groupby(['班级id', '班级名称', '年级']).agg(
        avg_score=('作业得分率', 'mean'), mock_count=('作业ID', 'count')
    ).reset_index()
    mock_months_by_class = mock_hw.groupby('班级id')['月份'].apply(lambda s: sorted(set(s.astype(str)))).to_dict()
    class_all_hw['mock_consecutive_3m'] = class_all_hw['班级id'].map(lambda cid: _are_consecutive_months(mock_months_by_class.get(cid, []), 3))
    class_all_hw['mock_count_for_rank'] = class_all_hw['班级id'].map(
        lambda cid: int(class_mock.loc[class_mock['班级id'].astype(str) == str(cid), 'mock_count'].sum())
    )
    preferred_classes = class_all_hw[class_all_hw['mock_consecutive_3m']].copy()
    if preferred_classes.empty:
        preferred_classes = class_all_hw.copy()
    preferred_classes = preferred_classes.sort_values(
        ['mock_consecutive_3m', 'mock_count_for_rank', 'all_hw_count'],
        ascending=[False, False, False]
    )

    top5_all = preferred_classes.head(5)
    top5_list = []
    for _, row in top5_all.iterrows():
        cid = str(row['班级id'])
        mock_row = class_mock[class_mock['班级id'].astype(str) == cid]
        avg_s = float(mock_row['avg_score'].values[0]) * 100 if len(mock_row) > 0 else 0
        mc    = int(mock_row['mock_count'].values[0])         if len(mock_row) > 0 else 0
        top5_list.append({
            'class_id':     cid,
            'class_name':   str(row['班级名称']),
            'grade':        str(row['年级']),
            'display_name': _format_grade_class(row['年级'], row['班级名称']),
            'all_hw_count': int(row['all_hw_count']),
            'mock_count':   mc,
            'avg_score':    round(avg_s, 2),
            'mock_consecutive_3m': bool(row.get('mock_consecutive_3m', False)),
        })
    results['top_classes'] = top5_list

    if top5_list:
        top_cid = top5_list[0]['class_id']
        results['top_class_name']  = top5_list[0]['display_name']
        results['top_class_grade'] = top5_list[0]['grade']

        top_all_m = hw_df[hw_df['班级id'].astype(str) == top_cid].groupby('月份').size()
        results['top_class_all_monthly'] = {str(m): int(v) for m, v in top_all_m.to_dict().items()}

        top_mock_m = mock_hw[mock_hw['班级id'].astype(str) == top_cid].groupby('月份')['作业得分率'].agg(['mean','count'])
        results['top_class_mock_monthly'] = {
            str(m): {'score': round(float(v['mean'])*100,2), 'count': int(v['count'])}
            for m, v in top_mock_m.to_dict('index').items()
        }

    # 各班级布置作业次数（前10排行）
    ca = class_df[['班级id','班级名称','年级','布置作业次数','布置作业份数']].copy()
    ca = ca.sort_values('布置作业次数', ascending=False)
    results['class_assign_top10'] = [
        {'class_name': str(row['班级名称']), 'grade': str(row['年级']),
         'hw_times': int(row['布置作业次数']), 'hw_count': int(row['布置作业份数'])}
        for _, row in ca.head(10).iterrows()
    ]

    # 各年级汇总（布置作业次数、作业份数、平均完成率、平均得分率）
    grade_class_stats = {}
    for grade, grp in class_df.groupby('年级'):
        g = str(grade)
        grade_class_stats[g] = {
            'hw_times':   int(grp['布置作业次数'].sum()),
            'hw_count':    int(grp['布置作业份数'].sum()),
            'completion_rate': round(float(pd.to_numeric(grp['作业完成率'], errors='coerce').mean()) * 100, 2),
            'score_rate':  round(float(pd.to_numeric(grp['作业得分率'], errors='coerce').mean()) * 100, 2),
        }
    results['grade_class_stats'] = grade_class_stats

    grade_profiles = {}
    for grade, grp in class_df.groupby('年级'):
        g = str(grade)
        students = int(pd.to_numeric(grp['总学生数'], errors='coerce').fillna(0).sum())
        active_paid = int(pd.to_numeric(grp.get('未过期付费学生数', 0), errors='coerce').fillna(0).sum()) if '未过期付费学生数' in grp else 0
        parent_bind = int(pd.to_numeric(grp.get('绑定家长通学生数', 0), errors='coerce').fillna(0).sum()) if '绑定家长通学生数' in grp else 0
        grade_profiles[g] = {
            'classes': int(grp['班级id'].nunique()),
            'students': students,
            'active_paid': active_paid,
            'paid_rate': round(active_paid / students * 100, 2) if students else 0,
            'parent_bind_rate': round(parent_bind / students * 100, 2) if students else 0,
            'avg_time': round(float(pd.to_numeric(grp.get('作答平均耗时/min', 0), errors='coerce').mean()), 2) if '作答平均耗时/min' in grp else 0,
            'self_practice': int(pd.to_numeric(grp.get('自主练习次数', 0), errors='coerce').fillna(0).sum()) if '自主练习次数' in grp else 0,
            'vocab_practice': int(pd.to_numeric(grp.get('词汇自主练习次数', 0), errors='coerce').fillna(0).sum()) if '词汇自主练习次数' in grp else 0,
        }
    results['grade_profiles'] = grade_profiles

    grade_category_mix = {}
    grade_cat_counts = hw_df.groupby(['年级', '大类']).size().unstack(fill_value=0)
    for grade, row in grade_cat_counts.iterrows():
        total = int(row.sum())
        grade_category_mix[str(grade)] = {
            str(cat): {'count': int(cnt), 'pct': round(int(cnt) / total * 100, 1) if total else 0}
            for cat, cnt in row.to_dict().items()
        }
    results['grade_category_mix'] = grade_category_mix

    months = sorted(results.get('monthly_hw', {}).keys())
    results['month_range'] = f"{min(months)} 至 {max(months)}" if months else "N/A"

    # ── 题型分析（来自「听说模拟班级总体情况」Excel）────────────
    if qt_df is not None and not qt_df.empty and '题型名称' in qt_df.columns:
        qt = qt_df.copy()
        results['has_question_type'] = True

        # ① 全校各题型平均得分率（难度排序）
        qt_school = qt.groupby('题型名称')['得分率'].agg(['mean', 'std', 'count']).round(4)
        results['qt_total_records'] = int(len(qt))
        results['qt_type_count'] = int(qt['题型名称'].nunique())
        results['qt_school'] = {
            name: {'mean': round(v['mean']*100, 2), 'std': round(v['std']*100, 2), 'count': int(v['count'])}
            for name, v in qt_school.to_dict('index').items()
        }

        # ② 班级 × 题型 得分率矩阵
        if '班级' in qt.columns and '年级' in qt.columns:
            ct = qt.groupby(['班级', '年级'])['得分率'].mean().round(4)
            ct_dict = {}
            for (c, g), v in ct.to_dict().items():
                ct_dict[_format_grade_class(g, c)] = round(float(v)*100, 2)
            results['qt_class'] = ct_dict

        # ③ 年级 × 题型 得分率矩阵
        if '年级' in qt.columns:
            gt = qt.groupby(['年级', '题型名称'])['得分率'].mean().round(4)
            gt_dict = {}
            for (grade, qtype), value in gt.to_dict().items():
                g = str(grade)
                if g not in gt_dict:
                    gt_dict[g] = {}
                gt_dict[g][qtype] = round(float(value) * 100, 2)
            results['qt_grade'] = gt_dict
            results['qt_weak_by_grade'] = {
                grade: [
                    {'题型': qtype, '得分率': score}
                    for qtype, score in sorted(items.items(), key=lambda x: x[1])[:3]
                ]
                for grade, items in gt_dict.items()
            }

        # ④ 班级-题型 二维矩阵（横向各班对比，纵向各题型）
        if '班级' in qt.columns and '年级' in qt.columns:
            ct_matrix = qt.pivot_table(
                index=['班级', '年级'], columns='题型名称', values='得分率', aggfunc='mean'
            ).round(4)
            results['qt_matrix'] = {
                _format_grade_class(g, c): {q: round(float(ct_matrix.loc[(c, g), q])*100, 2)
                               if q in ct_matrix.columns and (c, g) in ct_matrix.index else None
                               for q in ct_matrix.columns}
                for c, g in ct_matrix.index
            }
            # 离均差（相对各班均值的偏离，正=强项，负=弱项）
            ct_mean = ct_matrix.mean(axis=1)
            ct_dev = (ct_matrix.sub(ct_mean, axis=0)).round(4)
            results['qt_deviation'] = {
                _format_grade_class(g, c): {q: round(float(ct_dev.loc[(c, g), q])*100, 2)
                               if q in ct_dev.columns and (c, g) in ct_dev.index else None
                               for q in ct_dev.columns}
                for c, g in ct_dev.index
            }

        # ⑤ 各班薄弱题型（按班内离均差识别，并为正文展示预先筛选重点项）
        if '班级' in qt.columns and '年级' in qt.columns:
            ct_mean2 = ct_matrix.mean(axis=1)
            ct_dev2 = ct_matrix.sub(ct_mean2, axis=0)
            weak_qt = ct_dev2.idxmin(axis=1)
            weak_score = ct_dev2.min(axis=1)
            results['qt_weak'] = {
                _format_grade_class(g, c): {'题型': weak_qt.loc[(c, g)], '离均差': round(float(weak_score.loc[(c, g)])*100, 2)}
                for c, g in weak_qt.index
            }
            # 正文仅展示差异最突出的重点班级，且总量不超过5个
            ranked_weak = []
            for c, g in weak_qt.index:
                diff = round(float(weak_score.loc[(c, g)]) * 100, 2)
                if diff < 0:
                    ranked_weak.append({
                        '班级': _format_grade_class(g, c),
                        '题型': weak_qt.loc[(c, g)],
                        '离均差': diff,
                        'abs_diff': abs(diff),
                    })
            ranked_weak = sorted(ranked_weak, key=lambda x: x['abs_diff'], reverse=True)
            results['qt_weak_top'] = ranked_weak[:5]

        # ⑥ 教师 × 题型 得分率（教师效能分析）
        if '教师' in qt.columns and '题型名称' in qt.columns:
            tc_qt = qt.groupby(['教师', '题型名称'])['得分率'].mean().unstack(fill_value=0).round(4)
            results['qt_teacher'] = {
                str(t): {q: round(float(v)*100, 2) for q, v in row.items() if v > 0}
                for t, row in tc_qt.iterrows()
            }
            # 教师综合得分率
            tc_mean = qt.groupby('教师')['得分率'].mean().round(4).sort_values(ascending=False)
            results['qt_teacher_rank'] = {
                str(t): round(float(v)*100, 2) for t, v in tc_mean.to_dict().items()
            }

        # ⑦ 高分率/低分率（题型两极分化分析）
        if '优秀率' in qt.columns and '低分率' in qt.columns:
            qt_hr = qt.groupby('题型名称')[['优秀率', '低分率']].mean().round(4)
            results['qt_hr_lr'] = {
                q: {'优秀率': round(float(r['优秀率'])*100, 2),
                    '低分率': round(float(r['低分率'])*100, 2)}
                for q, r in qt_hr.to_dict('index').items()
            }

    else:
        results['has_question_type'] = False

    return results

def _first_non_empty(df, column, default=''):
    if column not in df.columns or df.empty:
        return default
    values = [str(v).strip() for v in df[column].dropna().tolist() if str(v).strip()]
    return values[0] if values else default

def _pct(value):
    return round(float(value) * 100, 2) if pd.notna(value) else 0

def analyze_district_data(class_df, hw_df, qt_df=None):
    """区级报告分析：复用现有三份基础数据，额外按学校、教师、年级聚合。"""
    data = analyze_data(class_df, hw_df, qt_df)
    data['report_scope'] = 'district'
    data['region_name'] = _first_non_empty(class_df, '区县') or _first_non_empty(class_df, '城市') or '区域'
    data['school_name'] = data['region_name']

    class_df = class_df.copy()
    hw_df = _coerce_hw_numeric(hw_df.copy())
    hw_df['大类'] = hw_df['作业路径'].apply(lambda x: _split_path(x)[0])
    hw_df['月份'] = pd.to_datetime(hw_df['作业开始日期'], errors='coerce').dt.to_period('M').astype(str)
    hw_df['完整路径'] = hw_df['作业路径'].fillna('')
    mock_hw = hw_df[hw_df['完整路径'].str.contains('模拟-', na=False)].copy()

    district_avg_apc = data['assign_count'] / max(data['classes'], 1)
    data['district_avg_assign_per_class'] = round(district_avg_apc, 1)

    mock_count_by_school = mock_hw.groupby('学校名称').size().to_dict() if not mock_hw.empty else {}

    school_rows = []
    for school, grp in class_df.groupby('学校名称'):
        classes = int(grp['班级id'].nunique())
        assign = int(pd.to_numeric(grp['布置作业次数'], errors='coerce').fillna(0).sum())
        completion = _pct(pd.to_numeric(grp['作业完成率'], errors='coerce').mean())
        avg_assign = round(assign / max(classes, 1), 1)
        mock_count = int(mock_count_by_school.get(school, 0))
        school_rows.append({
            '学校': str(school),
            '班级数': classes,
            '学生数': int(pd.to_numeric(grp['总学生数'], errors='coerce').fillna(0).sum()),
            '布置次数': assign,
            '班均布置次数': avg_assign,
            '完成率': completion,
            '得分率': _pct(pd.to_numeric(grp['作业得分率'], errors='coerce').mean()),
            '班均有效布置指数': round(avg_assign * completion / 100, 1),
            '听说模拟次数': mock_count,
            '班均听说模拟': round(mock_count / max(classes, 1), 1),
        })
    school_rows = sorted(school_rows, key=lambda r: r['布置次数'], reverse=True)
    data['school_stats'] = school_rows

    active_pool = [r for r in school_rows if r['班均布置次数'] > district_avg_apc]
    top_n = max(1, math.ceil(len(school_rows) * 0.3)) if school_rows else 0
    data['district_top_school_count'] = top_n
    data['active_school_top'] = sorted(active_pool, key=lambda r: r['班均有效布置指数'], reverse=True)[:top_n]
    data['low_active_schools'] = sorted(
        [r for r in school_rows if r['班均布置次数'] <= district_avg_apc],
        key=lambda r: r['班均布置次数'],
        reverse=True
    )

    teacher_rows = []
    if '教师姓名' in class_df.columns:
        for (school, teacher), grp in class_df.groupby(['学校名称', '教师姓名']):
            classes = int(grp['班级id'].nunique())
            assign = int(pd.to_numeric(grp['布置作业次数'], errors='coerce').fillna(0).sum())
            completion = _pct(pd.to_numeric(grp['作业完成率'], errors='coerce').mean())
            avg_assign = round(assign / max(classes, 1), 1)
            teacher_rows.append({
                '学校': str(school),
                '教师': str(teacher),
                '班级数': classes,
                '布置次数': assign,
                '班均布置次数': avg_assign,
                '完成率': completion,
                '得分率': _pct(pd.to_numeric(grp['作业得分率'], errors='coerce').mean()),
                '班均有效布置指数': round(avg_assign * completion / 100, 1),
            })
    data['teacher_active_top'] = sorted(teacher_rows, key=lambda r: r['班均有效布置指数'], reverse=True)[:10]

    school_cat = hw_df.groupby(['学校名称', '大类']).size().unstack(fill_value=0)
    data['school_category_mix'] = {}
    for school, row in school_cat.iterrows():
        total = int(row.sum())
        data['school_category_mix'][str(school)] = {
            str(cat): {'count': int(cnt), 'pct': round(int(cnt) / total * 100, 1) if total else 0}
            for cat, cnt in row.to_dict().items()
        }

    case_rows = []
    if not mock_hw.empty:
        school_grade_hw = hw_df.groupby(['学校名称', '年级']).size()
        school_grade_mock = mock_hw.groupby(['学校名称', '年级']).agg(
            mock_count=('作业ID', 'count'),
            avg_score=('作业得分率', 'mean'),
        )
        school_grade_month = mock_hw.groupby(['学校名称', '年级', '月份'])['作业得分率'].mean()
        class_sg = class_df.groupby(['学校名称', '年级'])['作业完成率'].mean()
        for idx in school_grade_mock.index:
            monthly = school_grade_month.loc[idx].dropna() if idx in school_grade_month.index.droplevel('月份') else pd.Series(dtype=float)
            scores = [float(v) * 100 for v in monthly.tolist()]
            improve = round(max(scores) - min(scores), 2) if len(scores) >= 2 else 0
            case_rows.append({
                '学校': str(idx[0]),
                '年级': str(idx[1]),
                '总作业次数': int(school_grade_hw.get(idx, 0)),
                '听说模拟次数': int(school_grade_mock.loc[idx, 'mock_count']),
                '完成率': _pct(class_sg.get(idx, 0)),
                '提分率': improve,
                '平均得分率': _pct(school_grade_mock.loc[idx, 'avg_score']),
            })
    data['district_case_top5'] = sorted(
        case_rows,
        key=lambda r: (r['听说模拟次数'], r['完成率'], r['提分率']),
        reverse=True
    )[:5]

    # ── 教研增强（v2.7）：作答率 / 单次耗时 / 班级四象限 / 作业负担 ──
    if '作答学生占比' in hw_df.columns and pd.notna(hw_df['作答学生占比'].mean()):
        data['answer_rate'] = round(float(hw_df['作答学生占比'].mean()) * 100, 1)
    if '单次作业平均耗时/min' in hw_df.columns and pd.notna(hw_df['单次作业平均耗时/min'].mean()):
        data['avg_minutes'] = round(float(hw_df['单次作业平均耗时/min'].mean()), 1)
        sch_min = hw_df.groupby('学校名称')['单次作业平均耗时/min'].mean()
        data['heavy_minute_schools'] = [(str(s), round(float(v), 1))
                                        for s, v in sch_min.items() if pd.notna(v) and v > 20]

    # 班级四象限：以全区班级平均完成率/得分率为分界，布置≥5次的班级参与，剔除非学校主体
    if {'作答学生占比', '100%完成学生占比', '作业得分率'}.issubset(hw_df.columns):
        cq = hw_df.groupby(['学校名称', '班级id']).agg(
            班级名称=('班级名称', 'first'), 年级=('年级', 'first'),
            assigns=('作业ID', 'nunique'), comp=('100%完成学生占比', 'mean'),
            score=('作业得分率', 'mean'), ans=('作答学生占比', 'mean'),
        ).reset_index()
        cq = cq[cq['学校名称'].map(_is_school)]
        cq = cq[cq['assigns'] >= 5].dropna(subset=['comp', 'score'])
        if len(cq) >= 4:
            mc, ms = cq['comp'].mean(), cq['score'].mean()

            def _cls_label(r):
                cn, g = str(r['班级名称']), str(r['年级'])
                return cn if cn.startswith(g[:1]) and len(cn) > 2 else f"{g}{cn}班"

            def _cls_rows(df, n=5):
                return [{'学校': str(r['学校名称']), '班级': _cls_label(r),
                         '布置次数': int(r['assigns']),
                         '完成率': round(float(r['comp']) * 100, 1),
                         '得分率': round(float(r['score']) * 100, 1),
                         '作答率': round(float(r['ans']) * 100, 1) if pd.notna(r['ans']) else 0}
                        for _, r in df.head(n).iterrows()]

            q_hh = cq[(cq.comp >= mc) & (cq.score >= ms)]
            q_hl = cq[(cq.comp >= mc) & (cq.score < ms)]
            q_lh = cq[(cq.comp < mc) & (cq.score >= ms)]
            q_ll = cq[(cq.comp < mc) & (cq.score < ms)]
            data['class_quadrant'] = {
                'n': int(len(cq)),
                'mc': round(float(mc) * 100, 1), 'ms': round(float(ms) * 100, 1),
                'counts': [('双高（示范班级）', len(q_hh)),
                           ('高完成低得分（方法待改进）', len(q_hl)),
                           ('低完成高得分（督促待加强）', len(q_lh)),
                           ('双低（重点跟进）', len(q_ll))],
                'method': _cls_rows(q_hl.sort_values('comp', ascending=False)),
                'follow': _cls_rows(q_ll.sort_values('assigns', ascending=False)),
            }

    return data


# ════════════════════════════════════════════════════════════════════════
# 阶段性联考（可选入口）：解析联考 Word、模糊匹配校名、做“练—考”关联分析
# ════════════════════════════════════════════════════════════════════════
def _exam_quals(data):
    """根据练习数据所在区域，构造校名归一化时要剥离的行政区前缀。"""
    quals = set()
    for key in ('province', 'city', 'region_name'):
        v = str(data.get(key, '') or '').strip()
        if v:
            quals.add(v)
            quals.add(v.rstrip('省市区县'))
    quals.discard('')
    # 长度优先，避免“滨州”先于“滨州市”被剥离导致残留
    return sorted(quals, key=len, reverse=True)


def _exam_norm(name, quals):
    s = str(name)
    for q in quals:
        s = s.replace(q, '')
    return s.strip()


def _exam_short(name, quals):
    """图表用短名：去前缀后进一步精简过长校名。"""
    s = _exam_norm(name, quals)
    return s.replace('街道办事处', '').replace('中心学校', '中心')


def _fuzzy_match_schools(practice_names, exam_names, quals):
    """练习校名 → 联考校名 模糊匹配。
    优先级：去前缀完全相等 > 互相包含 > difflib 相似度≥0.6。"""
    matched, used = {}, set()
    for p in practice_names:
        np_ = _exam_norm(p, quals)
        hit = None
        for ex in exam_names:
            if _exam_norm(ex, quals) == np_:
                hit = ex
                break
        if hit is None:
            for ex in exam_names:
                ne = _exam_norm(ex, quals)
                if np_ and (np_ in ne or ne in np_):
                    hit = ex
                    break
        if hit is None:
            best, br = None, 0.0
            for ex in exam_names:
                r = _seq_ratio(_exam_norm(ex, quals), np_)
                if r > br:
                    br, best = r, ex
            if br >= 0.6:
                hit = best
        if hit is not None:
            matched[p] = hit
            used.add(hit)
    only_practice = [p for p in practice_names if p not in matched]
    only_exam = [ex for ex in exam_names if ex not in used]
    return matched, only_practice, only_exam


def _seq_ratio(a, b):
    import difflib
    return difflib.SequenceMatcher(None, a, b).ratio()


def _exam_stage_label(text):
    """从联考 Word 文本中识别阶段标签与可排序时间键。
    返回 (label, sort_key)。label 形如 “2026年1月·期末”。"""
    year = month = None
    m = re.search(r'(20\d{2})\s*年\s*0?(\d{1,2})\s*月', text)
    if m:
        year, month = int(m.group(1)), int(m.group(2))
    phase, phase_order = '', 5
    if '期中' in text:
        phase, phase_order = '期中', 1
    elif '期末' in text:
        phase, phase_order = '期末', 2
    elif '月考' in text:
        phase, phase_order = '月考', 0
    ymd = re.search(r'(20\d{2})\s*年\s*0?(\d{1,2})\s*月\s*0?(\d{1,2})\s*日', text)
    day = int(ymd.group(3)) if ymd else 0
    if year:
        label = f"{year}年{month}月" + (f"·{phase}" if phase else '')
        sort_key = (year, month or 0, phase_order, day)
    else:
        label = phase or '联考'
        sort_key = (0, 0, phase_order, day)
    short = (f"{str(year)[2:]}年{month}月" if year else '联考') + (phase if phase else '')
    return label, short, sort_key


def _parse_one_exam(path):
    """解析单份联考 Word：整体 / 各校平均分 / 各题型得分率。"""
    from docx import Document
    doc = Document(path)
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    label, short, sort_key = _exam_stage_label(full_text)

    full_mark = 30.0
    mm = re.search(r'满分为?\s*(\d+(?:\.\d+)?)\s*分', full_text)
    if mm:
        full_mark = float(mm.group(1))
    sr = re.search(r'整体得分率为?\s*(\d+(?:\.\d+)?)%', full_text)
    score_rate = float(sr.group(1)) if sr else None

    schools, overall = {}, None
    qtype = {}
    for tbl in doc.tables:
        header = [c.text.strip() for c in tbl.rows[0].cells]
        head_join = ''.join(header)
        # 各校成绩表：表头含“学校”和“平均分”
        if '学校' in head_join and '平均分' in head_join and len(tbl.columns) >= 4:
            cols = {h: i for i, h in enumerate(header)}
            def col(row, key, default=''):
                idx = cols.get(key)
                return row[idx] if idx is not None and idx < len(row) else default
            for r in tbl.rows[1:]:
                cells = [c.text.strip() for c in r.cells]
                name = cells[0]
                if not name:
                    continue
                rec = {
                    '平均分': _safe_float(col(cells, '平均分')),
                    '学生人数': _safe_int(col(cells, '学生人数')),
                    '实考人数': _safe_int(col(cells, '实考人数')),
                    '缺考率': col(cells, '缺考率'),
                }
                if name in ('全部', '合计', '总计'):
                    overall = rec
                else:
                    schools[name] = rec
        # 题型表：某列存在“得分率”等指标行
        col0 = [r.cells[0].text.strip() for r in tbl.rows]
        if '得分率' in col0 and ('平均分' in col0 or '满分率' in col0):
            metric_start = col0.index('平均分') if '平均分' in col0 else col0.index('得分率')
            name_row = tbl.rows[metric_start - 1] if metric_start >= 1 else tbl.rows[0]
            names = [c.text.strip() for c in name_row.cells][1:]
            for r in tbl.rows[metric_start:]:
                cells = [c.text.strip() for c in r.cells]
                metric = cells[0]
                if metric not in ('平均分', '得分率', '满分率', '优秀率', '低分率'):
                    continue
                for n, v in zip(names, cells[1:]):
                    if not n:
                        continue
                    qtype.setdefault(n, {})[metric] = v

    if overall is None and schools:
        avgs = [s['平均分'] for s in schools.values() if s['平均分'] is not None]
        overall = {'平均分': round(sum(avgs) / len(avgs), 2) if avgs else None,
                   '学生人数': sum(s['学生人数'] or 0 for s in schools.values()),
                   '实考人数': sum(s['实考人数'] or 0 for s in schools.values()),
                   '缺考率': ''}
    if score_rate is None and overall and overall.get('平均分') and full_mark:
        score_rate = round(overall['平均分'] / full_mark * 100, 1)

    return {
        'label': label, 'short': short, 'sort_key': sort_key,
        'full_mark': full_mark, 'score_rate': score_rate,
        'overall': overall or {}, 'schools': schools, 'qtype': qtype,
    }


def _safe_float(v):
    try:
        return float(str(v).strip().replace('%', ''))
    except Exception:
        return None


def _safe_int(v):
    try:
        return int(float(str(v).strip().replace(',', '')))
    except Exception:
        return 0


def parse_exam_reports(paths):
    """解析多份联考 Word 并按时间先后排序（识别文档中的年月与期中/期末）。"""
    stages = []
    for p in paths or []:
        try:
            stages.append(_parse_one_exam(p))
        except Exception:
            continue
    stages.sort(key=lambda s: s['sort_key'])
    return stages


def attach_exam_analysis(data, exam_stages):
    """把联考结果与练习画像按模糊匹配拼接，写入 data['exam']。仅区级使用。"""
    if not exam_stages:
        return data
    quals = _exam_quals(data)
    latest = exam_stages[-1]
    multi = len(exam_stages) > 1

    practice = {r['学校']: r for r in data.get('school_stats', [])}
    exam_names = list(latest['schools'].keys())
    matched, only_p, only_e = _fuzzy_match_schools(list(practice.keys()), exam_names, quals)

    link = []
    for p, ex in matched.items():
        s = practice[p]
        avg = latest['schools'][ex].get('平均分')
        if avg is None:
            continue
        link.append({
            '练习校名': p, '联考校名': ex, '短名': _exam_short(ex, quals),
            '练习得分率': s.get('得分率', 0),
            '班均听说模拟': s.get('班均听说模拟', 0),
            '完成率': s.get('完成率', 0),
            '联考平均分': avg,
            '实考人数': latest['schools'][ex].get('实考人数', 0),
        })
    link.sort(key=lambda r: r['联考平均分'], reverse=True)

    def _corr(xs, ys):
        pairs = [(x, y) for x, y in zip(xs, ys) if x is not None and y is not None]
        n = len(pairs)
        if n < 3:
            return 0.0
        mx = sum(p[0] for p in pairs) / n
        my = sum(p[1] for p in pairs) / n
        cov = sum((p[0]-mx)*(p[1]-my) for p in pairs)
        sx = math.sqrt(sum((p[0]-mx)**2 for p in pairs))
        sy = math.sqrt(sum((p[1]-my)**2 for p in pairs))
        return round(cov/(sx*sy), 3) if sx*sy else 0.0

    r_score = _corr([r['练习得分率'] for r in link], [r['联考平均分'] for r in link])
    r_freq = _corr([r['班均听说模拟'] for r in link], [r['联考平均分'] for r in link])

    # 题型：日常练习 vs 联考（最新一场）
    qt_school = data.get('qt_school', {})
    qt_compare = []
    for q in EXAM_QTYPES:
        prac_v = round(qt_school[q]['mean'], 1) if q in qt_school else None
        exam_v = _safe_float(latest['qtype'].get(q, {}).get('得分率'))
        if prac_v is not None or exam_v is not None:
            qt_compare.append({'题型': q, '练习': prac_v, '联考': exam_v})

    # 多阶段：整体趋势 + 各校阶段矩阵
    stage_labels = [s['short'] for s in exam_stages]
    overall_trend = [s['overall'].get('平均分') for s in exam_stages]
    school_stage = {}
    for ex_name in exam_names:
        series = []
        for st in exam_stages:
            series.append(st['schools'].get(ex_name, {}).get('平均分'))
        school_stage[_exam_short(ex_name, quals)] = series

    data['exam'] = {
        'multi': multi,
        'stages': [{'label': s['label'], 'short': s['short'], 'overall': s['overall'],
                    'score_rate': s['score_rate'], 'full_mark': s['full_mark'],
                    'school_count': len(s['schools'])} for s in exam_stages],
        'latest': {'label': latest['label'], 'overall': latest['overall'],
                   'score_rate': latest['score_rate'], 'full_mark': latest['full_mark'],
                   'school_count': len(latest['schools']),
                   'schools': latest['schools'], 'qtype': latest['qtype']},
        'full_mark': latest['full_mark'],
        'link': link, 'r_score': r_score, 'r_freq': r_freq,
        'only_practice': only_p, 'only_exam': only_e,
        'only_exam_scores': {_exam_short(e, quals): latest['schools'][e].get('平均分') for e in only_e},
        'qt_compare': qt_compare,
        'stage_labels': stage_labels, 'overall_trend': overall_trend,
        'school_stage': school_stage,
        '_quals': quals,
    }
    return data


def _build_province_policy(province, city):
    if province == '黑龙江省' and city == '哈尔滨市':
        return (
            f"哈尔滨市于2024年发布中考综合改革实施方案（试行），2024-2025年过渡期内，"
            f"英语听说考试采用人机对话形式，口语10分、听力20分，合计30分，2026年起全部计入中考总分。"
            f"省统一要求外语纳入口语、听力测试并计入总分。"
        )
    elif province == '黑龙江省':
        return (
            f"黑龙江省积极推进中考综合改革，外语科目纳入口语、听力测试，"
            f"英语听说考试采用人机对话形式，2026年起计入中考总分。"
        )
    elif city:
        return f"{city}市积极推进中考英语听说教育改革，人机对话测试已纳入日常教学训练体系。"
    else:
        return "各地正全面推进英语听说中考试点，人机对话测试已逐步纳入中考范围。"

CURRICULUM_STD = (
    "《义务教育课程方案和课程标准（2022年版）》明确提出，要培养学生核心素养，强调语言运用能力，尤其是听说能力的培养。"
    "英语听说训练是落实学科核心素养、提升学生综合语言运用能力的重要途径，日常朗读、跟读训练和真题训练是提升学生听说能力的有效手段。"
)

def generate_report_text(data):
    school        = data['school_name']
    months        = sorted(data.get('monthly_hw', {}).keys())
    mr            = data.get('month_range', 'N/A')
    total_hw     = data['total_hw']
    syn_pct       = data['category_pct'].get('同步', 0)
    mon_pct       = data['category_pct'].get('模拟', 0)
    sub_pct       = data['category_pct'].get('专项', 0)
    r_v, n_v      = data['corr_vocab']
    r_c, n_c      = data['corr_complete']
    r_s, n_s      = data['corr_self']
    strong        = data.get('strong_corrs', [])
    top           = data.get('top_classes', [])
    grade_profiles = data.get('grade_profiles', {})
    grade_category_mix = data.get('grade_category_mix', {})
    vocab_p       = data['vocab_practice']
    per_class_hw  = round(data['assign_count'] / max(data['classes'], 1), 1)
    tc_name       = data.get('top_class_name', '标杆班级')
    tc_grade      = data.get('top_class_grade', '')
    actual_grades = _sort_grades(data.get('actual_grades', []))

    # ── 辅助函数：构建相关性判定标签 ───────────────────────
    def corr_label(r):
        if   abs(r) >= 0.5: return "强正相关" if r > 0 else "强负相关"
        elif abs(r) >= 0.4: return "中等正相关" if r > 0 else "中等负相关"
        elif abs(r) >= 0.3: return "弱正相关" if r > 0 else "弱负相关"
        return "相关性弱"

    # ── 数据叙事辅助 ───────────────────────────────────────

    

    # ═══════════════════════════════════════════════════════════
    # 报告文字生成（generate_report_text）
    # 参考广州示范校报告写作规范：
    #   - 总领句先行（"在XXX方面…"）
    #   - 数据前有定性描述，数据后有趋势小结
    #   - 禁止无连接词的跳跃式表述
    #   - 对比参照：环比/同比方向须明确
    #   - 数字叙事：具体数字 > 笼统描述
    # ═══════════════════════════════════════════════════════════
    L = []
    L.append(f"# {school} 英语AI听说产品应用成效报告\n")
    L.append(f"**报告生成时间：{datetime.now().strftime('%Y年%m月%d日')}**\n")
    L.append(f"**模板版本：{REPORT_VERSION}**\n")
    L.append("\n")

    # ── 一、学校信息 ───────────────────────────────────────
    L.append("## 一、学校信息\n")
    L.append(f"　　{school}积极推进教育数字化转型，在{data.get('province', '黑龙江省')}英语听说训练改革持续推进的背景下，引入E听说AI听说训练系统，依托大数据与人工智能技术赋能英语听说训练变革。本报告聚焦系统在校内的应用过程、训练结构与阶段性成效，数据覆盖周期为{_format_month_label(mr)}。\n\n")
    L.append("| 项目 | 内容 |\n|------|------|\n")
    L.append(f"| 学校名称 | {school} |\n")
    L.append(f"| 所属省份 | {data.get('province', '黑龙江省')} |\n")
    L.append(f"| 所属城市 | {data.get('city', '哈尔滨市')} |\n")
    L.append(f"| 数据周期 | {_format_month_label(mr)} |\n")
    L.append("\n")

    # ── 二、激活/应用概况 ──────────────────────────────────
    L.append("## 二、激活/应用概况\n")
    L.append(f"　　在应用覆盖方面，{school}在{_format_month_label(mr)}期间共有{data['classes']}个班级、{data['total_students']}名学生纳入E听说产品应用。教师端累计布置作业{data['assign_count']}次，班均约{per_class_hw}次，说明系统已进入日常训练使用链条，而不是停留在一次性试用阶段。\n\n")
    L.append("**核心应用数据如下：**\n\n")
    L.append("| 指标 | 数值 |\n|------|------|\n")
    L.append(f"| 参与学校数 | {data['schools']}所 |\n")
    L.append(f"| 班级数（去重） | {data['classes']}个 |\n")
    L.append(f"| 激活学生总数 | {data['total_students']}人 |\n")
    L.append(f"| 布置作业次数（合计） | {data['assign_count']}次 |\n")
    L.append(f"| 布置作业份数（合计） | {data['assign_total']}份 |\n")
    L.append(f"| 作业完成率（均值） | {data['completion_rate']}% |\n")
    L.append(f"| 班级平均作业得分率 | {data['score_rate_avg']}% |\n")
    L.append("\n")

    avg_month_hw = round(total_hw / max(len(months), 1), 1)

    grade_stats = data.get('grade_class_stats', {})
    if grade_stats:
        L.append(f"**各年级应用画像：**\n\n")
        L.append("| 年级 | 班级数 | 学生数 | 布置作业次数 | 平均完成率 | 平均得分率 | 平均作答耗时 | 自主练习次数 |\n")
        L.append("|------|-------|-------|------------|----------|----------|------------|------------|\n")
        for grade in _sort_grades(grade_stats.keys()):
            g = grade_stats[grade]
            p = grade_profiles.get(grade, {})
            L.append(f"| {grade} | {p.get('classes', '—')}个 | {p.get('students', '—')}人 | {g['hw_times']}次 | {g['completion_rate']}% | {g['score_rate']}% | {p.get('avg_time', '—')}分钟 | {p.get('self_practice', '—')}次 |\n")
        L.append("\n")
        # 数据小结：找完成率最高和得分率最高的年级
        best_completion_grade = max(grade_stats, key=lambda g: grade_stats[g]['completion_rate'])
        best_score_grade = max(grade_stats, key=lambda g: grade_stats[g]['score_rate'])
        longest_time_grade = max(grade_profiles, key=lambda g: grade_profiles[g].get('avg_time', 0)) if grade_profiles else ''
        extra_time_text = f"；{longest_time_grade}平均作答耗时最长（{grade_profiles[longest_time_grade]['avg_time']}分钟）" if longest_time_grade else ""
        L.append(f"　　从各年级横向对比来看，{best_completion_grade}平均完成率最高（{grade_stats[best_completion_grade]['completion_rate']}%），{best_score_grade}平均得分率领先（{grade_stats[best_score_grade]['score_rate']}%）{extra_time_text}。结合布置次数、完成率、得分率和作答耗时，可以看出不同年级在训练投入与训练结果上的差异。\n\n")

    L.append("> 数据来源：班级数据总览、作业明细\n\n")

    # ── 三、应用情况分析 ────────────────────────────────────
    L.append("## 三、应用情况分析\n")

    L.append("### 3.1 训练内容/栏目介绍\n")
    L.append(f"　　系统应用覆盖四类训练场景，以「教材」日常开口训练为主体，辅以「真题」等训练形式，形成较为完整的听说训练支持体系。其中，「教材」训练侧重夯实发音基础与语感积累，「真题」训练侧重服务题型巩固、阶段性检测与考前演练。\n\n")
    L.append("| 大类 | 次数 | 占比 | 定位说明 |\n|------|------|------|----------|\n")
    cat_meta = {
        '同步':      '教材朗读/跟读等日常基础训练，帮助学生建立标准发音与语感',
        '专项':      '真题题型练习，针对性强化薄弱题型',
        '模拟':      '真题套卷训练，含区域精选/单元测试等，贴近真实考试场景',
        '课外拓展':  '趣味配音等拓展训练，提升学习兴趣与语用能力',
        '其他':      '其他内容',
    }
    for cat, cnt in sorted(data.get('category_counts', {}).items(), key=lambda x: -x[1]):
        pct_v = data['category_pct'].get(cat, 0)
        L.append(f"| **{cat}** | {cnt}次 | {pct_v}% | {cat_meta.get(cat,'')} |\n")
    L.append("\n")

    L.append("### 3.2 整体应用数据\n")
    practice_parts = []
    if data['self_practice'] >= 100:
        practice_parts.append(f"学生自主练习次数为**{data['self_practice']}次**")
    if vocab_p >= 100:
        practice_parts.append(f"词汇自主练习次数为**{vocab_p}次**")
    practice_text = "，" + "、".join(practice_parts) + "，体现出课堂任务之外的自主巩固基础" if practice_parts else ""
    L.append(f"　　在作业应用方面，本周期作业明细共记录**{total_hw}次**作业，月均约**{avg_month_hw}次**。与激活概况中的覆盖数据相比，本节重点关注应用深度：教师布置行为持续发生{practice_text}。\n\n")

    # 动态找峰值月及增长趋势
    peak_m = max(months, key=lambda m: data['monthly_hw'].get(m, 0)) if months else months[0] if months else ''
    peak_cnt = data['monthly_hw'].get(peak_m, 0)
    L.append("### 3.3 应用频次分析\n")
    L.append(f"　　在应用频次方面，{total_hw}次作业分布于{len(months)}个月份，整体呈现常态化推进态势。{_format_month_label(months[0]) if months else ''}至{_format_month_label(months[-1]) if months else ''}期间，{_format_month_label(peak_m)}作业量最高（{peak_cnt}次），与阶段复习安排基本一致，说明系统应用与学校教学节奏保持较好匹配。\n\n")
    grade_hw = data.get('grade_monthly_hw', {})
    if grade_hw:
        peak_month = max(months, key=lambda m: data['monthly_hw'][m])
        grade_candidates = _sort_grades(grade_hw.keys()) or actual_grades
        peak_grade = max(grade_candidates, key=lambda g: grade_hw.get(g, {}).get(peak_month, 0))
        grade_span = "、".join(grade_candidates)
        L.append(f"　　从各年级横向对比来看，本周期覆盖{grade_span}，{_format_month_label(peak_month)}作业量最高，{peak_grade}在当月作业量最大，表明该年级在当月应用节奏中最为活跃。\n\n")

    L.append("### 3.4 应用方式分析\n")
    L.append(f"　　从作业内容结构来看，「同步」训练（课文朗读、跟读等）是学生日常接触最多的形式，合计占比高达**{syn_pct}%**，构成学生每日开口说英语的基础；「专项」训练占比**{sub_pct}%**，主要用于针对性强化重点题型；「模拟」训练占比**{mon_pct}%**，直接服务阶段检测与听说考试备考。整体呈现“日常打基础 + 专项补短板 + 模拟促实战”的组合模式，符合循序渐进的教学与备考规律。\n\n")
    cat_monthly = data.get('cat_monthly', {})
    if cat_monthly:
        peak_cat_month = max(cat_monthly, key=lambda m: sum(cat_monthly.get(m, {}).values()))
        peak_cat_data = cat_monthly.get(peak_cat_month, {})
        top_cat = max(peak_cat_data, key=lambda c: peak_cat_data.get(c, 0)) if peak_cat_data else ''
        if top_cat:
            top_cat_cnt = peak_cat_data.get(top_cat, 0)
            total_cat_cnt = sum(peak_cat_data.values())
            top_cat_pct = round(top_cat_cnt / total_cat_cnt * 100, 1) if total_cat_cnt else 0
            L.append(f"　　按月份进一步拆解，{_format_month_label(peak_cat_month)}作业量最高（{total_cat_cnt}次），其中「{top_cat}」训练{top_cat_cnt}次，占当月作业量的{top_cat_pct}%。\n\n")
    if grade_category_mix:
        L.append("**各年级训练结构差异：**\n\n")
        L.append("| 年级 | 同步训练 | 专项训练 | 模拟训练 | 结构判断 |\n")
        L.append("|------|---------|---------|---------|---------|\n")
        for grade in _sort_grades(grade_category_mix.keys()):
            mix = grade_category_mix.get(grade, {})
            syn = mix.get('同步', {}).get('pct', 0)
            sub = mix.get('专项', {}).get('pct', 0)
            mock = mix.get('模拟', {}).get('pct', 0)
            if mock >= 60:
                label = "以模拟备考为主"
            elif syn >= 60:
                label = "以同步基础训练为主"
            elif sub >= 20:
                label = "专项补弱占比较高"
            else:
                label = "训练结构相对均衡"
            L.append(f"| {grade} | {syn}% | {sub}% | {mock}% | {label} |\n")
        L.append("\n")
        max_grade = None
        max_cat = None
        max_pct = -1
        for grade, mix in grade_category_mix.items():
            for cat, item in mix.items():
                if item.get('pct', 0) > max_pct:
                    max_grade, max_cat, max_pct = grade, cat, item.get('pct', 0)
        if max_grade and max_cat:
            L.append(f"　　从结构占比看，{max_grade}「{max_cat}」训练占比最高（{max_pct}%），是本周期年级训练结构中最突出的类目。\n\n")

    # 四、应用效果分析
    L.append("## 四、应用效果分析\n")

    L.append("### 4.1 成绩数据对比\n")
    grade_scores = data.get('mock_hw_grade_monthly', {})
    # 找最优回升案例（从最低点到最后一个月的最大涨幅）
    best_recovery = None
    best_gain = 0
    for grade, monthly in grade_scores.items():
        sm = sorted(monthly.items())
        if len(sm) >= 2:
            vals = [s for _, s in sm]
            low_idx = vals.index(min(vals))
            if low_idx < len(vals) - 1 and vals[-1] > vals[low_idx]:
                gain = vals[-1] - vals[low_idx]
                if gain > best_gain:
                    best_gain = gain
                    best_recovery = (grade, sm[low_idx][0], sm[low_idx][1], sm[-1][0], sm[-1][1])

    best_trend_text = None
    if best_recovery:
        g, lm, ls, ltm, lts = best_recovery
        gain = round(lts - ls, 2)
        best_trend_text = f"　　**{g}**听说模拟得分率从最低点{_format_month_label(lm)}的**{ls}%**逐步回升至{_format_month_label(ltm)}的**{lts}%**，整体提升**{gain}个百分点**"
    elif grade_scores:
        best_g = max(grade_scores.keys(), key=lambda g: len(grade_scores[g]))
        sm = sorted(grade_scores[best_g].items())
        best_trend_text = f"　　**{best_g}**听说模拟月均得分率走势为：{' → '.join([f'{_format_month_label(m)} {s}%' for m,s in sm])}"
    if best_trend_text:
        L.append(f"{best_trend_text}，具体数据如下：\n\n")

    L.append("**各年级听说模拟得分率趋势：**\n\n")
    # 交叉表：行=月份，列=年级
    all_grades_sorted = _sort_grades(grade_scores.keys())
    L.append("| 月份 | " + " | ".join(all_grades_sorted) + " |\n")
    L.append("|" + "|".join(["------"] * (len(all_grades_sorted)+1)) + "|\n")
    for m in months:
        vals = [f"{grade_scores[g][m]}%" if m in grade_scores.get(g, {}) else "—" for g in all_grades_sorted]
        L.append(f"| {_format_month_label(m)} | " + " | ".join(vals) + " |\n")
    L.append("\n")

    # ── 4.2 题型得分分析（可选，依赖「听说模拟班级总体情况」Excel）──
    if data.get('has_question_type'):
        L.append("### 4.2 题型得分分析\n")
        qt_school = data.get('qt_school', {})
        if qt_school:
            # 题型难度排序
            sorted_qt = sorted(qt_school.items(), key=lambda x: x[1]['mean'])
            hardest_qt  = sorted_qt[0][0]
            mid_qt      = sorted_qt[1][0] if len(sorted_qt) > 1 else ''
            easiest_qt  = sorted_qt[-1][0]
            qt_total = data.get('qt_total_records', sum(v.get('count', 0) for v in qt_school.values()))
            qt_type_count = data.get('qt_type_count', len(qt_school))
            L.append(f"基于本周期听说模拟题型数据，共分析**{qt_total}条**题型记录，覆盖**{qt_type_count}类**题型。全校各题型平均得分率如下：\n\n")
            L.append("| 题型 | 平均得分率 | 难度定位 |\n|------|-----------|----------|\n")
            qt_labels = {
                '对话或短文朗读': '基础题型，整体表现较好',
                '情景反应':       '中等难度题型，表现存在一定差异',
                '篇章复述':       '综合性较强题型，对能力要求较高',
                '朗读句子1':       '朗读基础题型，关注语音准确度',
                '朗读句子2':       '朗读基础题型，关注语音准确度',
                '朗读段落':        '篇章朗读题型，关注连贯性和语音语调',
                '情景提问':        '输出型题型，对信息提取与表达组织要求较高',
                '简述和回答':      '综合输出题型，对复述、概括和即时回答要求较高',
                '图片描述':        '半开放表达题型，关注观察与组织表达能力',
                '快速应答':        '即时反应题型，关注听懂问题后的快速表达',
                'SectionA':        '听说考试分项，需结合试卷结构解读',
                'SectionB':        '听说考试分项，需结合试卷结构解读',
            }
            for qt_name, info in sorted_qt:
                label = qt_labels.get(qt_name, '')
                L.append(f"| {qt_name} | {info['mean']}% | {label} |\n")
            L.append("\n")
            L.append(f"　　从全校横向对比情况看，**{easiest_qt}**得分率最高（{qt_school[easiest_qt]['mean']}%），反映出学生在该题型上的整体掌握情况相对较好；**{hardest_qt}**得分率最低（{qt_school[hardest_qt]['mean']}%），低分率为**{data.get('qt_hr_lr', {}).get(hardest_qt, {}).get('低分率', 'N/A')}%**，应作为下一阶段专项教学与复习提升的重点关注题型。\n\n")

            weak_by_grade = data.get('qt_weak_by_grade', {})
            if weak_by_grade:
                L.append("**各年级薄弱题型聚焦：**\n\n")
                L.append("| 年级 | 第一薄弱题型 | 第二薄弱题型 | 第三薄弱题型 |\n")
                L.append("|------|-------------|-------------|-------------|\n")
                for grade in _sort_grades(weak_by_grade.keys()):
                    items = weak_by_grade.get(grade, [])
                    cells = [f"{item['题型']}（{item['得分率']}%）" for item in items[:3]]
                    cells += ['—'] * (3 - len(cells))
                    L.append(f"| {grade} | {cells[0]} | {cells[1]} | {cells[2]} |\n")
                L.append("\n")
                L.append("　　按年级定位薄弱题型后，后续教研可从“全校共同短板”和“年级差异短板”两条线推进：共同短板适合统一设计专项训练，年级差异短板则更适合由备课组结合教学进度分层处理。\n\n")

        # 各班薄弱题型（仅展示差异最突出的重点项，总量不超过5个）
        qt_weak_top = data.get('qt_weak_top', [])
        if qt_weak_top:
            L.append("**各班薄弱题型诊断：**\n\n")
            weak_list = []
            for item in qt_weak_top:
                weak_list.append(f"- **{item['班级']}**：**{item['题型']}**相对班级平均水平偏弱，离均差为{item['离均差']}个百分点")

            if weak_list:
                L.append("\n".join(weak_list) + "\n")
                L.append("　　以上诊断聚焦差异最突出的重点班级与题型，便于学校优先开展针对性训练与教学改进。\n\n")

        # 教师效能
        qt_teacher = data.get('qt_teacher_rank', {})
        if qt_teacher:
            L.append("**教师效能排名（综合各题型平均得分率）：**\n\n")
            L.append("| 排名 | 教师 | 综合得分率 | 薄弱题型提示 |\n")
            L.append("|------|------|-----------|------------|\n")
            # 找教师对应的薄弱题型
            qt_t_full = data.get('qt_teacher', {})
            sorted_teachers = sorted(qt_teacher.items(), key=lambda x: -x[1])
            for rank, (t, score) in enumerate(sorted_teachers[:8], 1):
                t_scores = qt_t_full.get(t, {})
                weak_t = min(t_scores.items(), key=lambda x: x[1]) if t_scores else ('', 0)
                weak_str = f"{weak_t[0]}({weak_t[1]}%)" if weak_t[0] else '—'
                L.append(f"| {rank} | {t} | {score}% | {weak_str} |\n")
            L.append("\n")
            best_t = sorted_teachers[0][0] if sorted_teachers else ''
            worst_t = sorted_teachers[-1][0] if sorted_teachers else ''
            if best_t and worst_t and best_t != worst_t:
                best_score = sorted_teachers[0][1]
                worst_score = sorted_teachers[-1][1]
                diff = best_score - worst_score
            L.append(f"　　教师所带班级综合得分率中，{best_t}所带班级最高（{best_score}%），可作为校本教研、课堂观摩与经验交流的重点参考。\n\n")

    if strong:
        L.append("### 4.3 相关性分析\n")
        L.append("以班级为单位，分析各类学习行为与作业得分率之间的相关性（Pearson相关系数）：\n\n")
        L.append("| 分析维度 | 相关系数 | 样本量 | 强度判定 | 结论 |\n")
        L.append("|---------|---------|--------|---------|------|\n")
        for lbl, r, n in [
            ('词汇自主练习次数 vs 平均得分率', r_v, n_v),
            ('作业完成率 vs 平均得分率',        r_c, n_c),
            ('自主练习次数 vs 平均得分率',      r_s, n_s),
        ]:
            d    = corr_label(r)
            flag = " ✅" if abs(r) >= 0.4 else ""
            L.append(f"| {lbl} | {r:.4f} | {n}个班级 | {d}{flag} | {'正向关联' if r > 0.3 else '需进一步观察'} |\n")
        L.append("\n")
        L.append("**Pearson相关系数理论说明：**\n\n")
        L.append("| 系数范围 | 相关强度 | 统计含义 |\n")
        L.append("|---------|---------|---------|\n")
        L.append("| 0.7 ≤ |r| ≤ 1.0 | 强相关 | 两变量存在明显线性关系 |\n")
        L.append("| 0.4 ≤ |r| < 0.7 | 中等相关 | 两变量存在一定线性关系 |\n")
        L.append("| 0.2 ≤ |r| < 0.4 | 弱相关 | 两变量存在微弱线性关系 |\n")
        L.append("| 0 ≤ |r| < 0.2 | 几乎无相关 | 两变量线性关系较弱 |\n")
        L.append("\n")
        L.append("**强相关发现：**\n")
        for lbl, r, n in strong:
            direction = "正相关" if r > 0 else "负相关"
            meaning = "表明相关学习行为与学业表现之间具有较为明确的同向关系" if r > 0 else "提示该行为可能更多集中在薄弱班级或补救场景中，需要结合班级基础进一步解释"
            L.append(f"- **{lbl}**与得分率呈中等及以上{direction}（r={r:.4f}，n={n}），{meaning}，可为后续教学管理与过程评价提供参考。\n")
        L.append("\n")

    # 五、典型案例
    L.append("## 五、典型班级分析\n")
    if top:
        top0 = top[0]
        L.append(f"　　以**{tc_name}**作为标杆班级（数据周期内作业总量位居全校前列）：\n\n")
        L.append(f"- 该班共完成**{top0['all_hw_count']}次**作业（所有类目），其中听说模拟**{top0['mock_count']}次**\n")
        L.append(f"- 听说模拟平均得分率高达**{top0['avg_score']}%**，居全校前列\n\n")

        top_all_m = data.get('top_class_all_monthly', {})
        top_mock_m = data.get('top_class_mock_monthly', {})

        if top_all_m and top_mock_m:
            # 合并为一个表：三列
            L.append("| 月份 | 所有类目布置作业次数 | 听说模拟类目布置作业次数 | 得分率 |\n")
            L.append("|------|------------------|----------------------|--------|\n")
            all_months = sorted(set(top_all_m.keys()) | set(top_mock_m.keys()))
            for m in all_months:
                all_cnt = top_all_m.get(m, '—')
                mdata = top_mock_m.get(m, None)
                mock_cnt = mdata['count'] if mdata else '—'
                score    = f"{mdata['score']}%" if mdata else '—'
                L.append(f"| {m} | {all_cnt} | {mock_cnt} | {score} |\n")
            L.append("\n")

        L.append("**全校TOP5班级（按总作业量排名）：**\n\n")
        L.append("| 排名 | 班级 | 总作业次数 | 听说模拟次数 | 平均得分率 |\n")
        L.append("|------|------|----------|------------|--------|\n")
        for i, c in enumerate(top, 1):
            L.append(f"| {i} | {c['display_name']} | {c['all_hw_count']}次 | {c['mock_count']}次 | {c['avg_score']}% |\n")
        L.append("\n")

    # 六、总结与建议
    L.append("## 六、总结与建议\n")
    top_score = top[0]["avg_score"] if top else 0

    # ── 问题诊断（基于实际数据发现）────────────────────────
    issues = []   # 存储具体问题描述，供亮点和问题两节共用

    # 问题1：类别失衡
    if syn_pct >= 75:
        issues.append(f"同步训练占比偏高（{syn_pct}%），专项训练与模拟训练占比相对不足，训练结构仍有优化空间")
    if mon_pct < 10 and sub_pct < 10:
        issues.append(f"专项训练（{sub_pct}%）与模拟训练（{mon_pct}%）总体占比较低，阶段性巩固与实战演练力度有待加强")
    elif mon_pct < 10:
        issues.append(f"模拟训练占比仅{mon_pct}%，阶段性检验与考前演练支撑相对不足，训练链条仍需进一步完善")

    # 问题2：班级间不均衡
    if grade_stats:
        rates = [g['completion_rate'] for g in grade_stats.values()]
        if rates and max(rates) - min(rates) > 20:
            worst_g = min(grade_stats, key=lambda g: grade_stats[g]['completion_rate'])
            best_g  = max(grade_stats, key=lambda g: grade_stats[g]['completion_rate'])
            issues.append(f"年级间作业完成率存在较明显差异：{best_g}最高为{grade_stats[best_g]['completion_rate']}%，{worst_g}为{grade_stats[worst_g]['completion_rate']}%，相差{max(rates)-min(rates):.0f}个百分点，分层推进的均衡性仍需加强")

    # 问题3：自主练习薄弱
    if vocab_p < data['total_students'] * 2 and vocab_p >= 100:
        issues.append(f"词汇自主练习总量为{vocab_p}次，自主巩固训练频次仍有提升空间")
    if abs(r_v) < 0.2 and vocab_p > 0:
        issues.append(f"词汇自主练习与得分相关性较弱（r={r_v:.4f}），自主练习与学业表现之间的转化效应仍需进一步观察")

    # 问题4：整体完成率偏低
    if data['completion_rate'] < 70:
        issues.append(f"作业平均完成率为{data['completion_rate']}%，与较高质量落实要求相比仍有差距，学生持续性练习习惯仍需进一步培育")

    # ── 亮点提炼（从数据中找最大优势）───────────────────────
    highlights = []

    # 亮点1：覆盖率
    highlights.append(f"应用覆盖面较广：{data['classes']}个班级、{data['total_students']}名学生已纳入系统应用范围，作业完成率均值为{data['completion_rate']}%，常态化使用机制已基本建立")

    # 亮点2：最强相关指标
    corr_list = [('词汇自主练习', r_v, n_v), ('作业完成率', r_c, n_c), ('自主练习次数', r_s, n_s)]
    best_corr = max(corr_list, key=lambda x: abs(x[1]))
    bc_name, bc_r, bc_n = best_corr
    if abs(bc_r) >= 0.4:
        highlights.append(f"关键过程指标与结果指标关联较为明确：{bc_name}与得分呈中等正相关（r={bc_r:.4f}，n={bc_n}个班级），说明相关学习行为对学习成效具有一定支撑作用")
    elif abs(bc_r) >= 0.2:
        highlights.append(f"自主练习积累呈现一定积极作用：{bc_name}与得分呈弱正相关（r={bc_r:.4f}），持续练习对能力提升的长期价值值得进一步跟踪")

    # 亮点3：训练体系
    if syn_pct + mon_pct + sub_pct >= 85:
        highlights.append(f"训练结构较为完整：同步训练占比{syn_pct}%，专项训练占比{sub_pct}%，模拟训练占比{mon_pct}%，能够较好支撑日常教学、专项提升与阶段检测的衔接")
    elif syn_pct >= 50:
        highlights.append(f"日常训练基础较为扎实：同步训练占比{syn_pct}%，朗读跟读等常规训练已较好融入课堂教学过程，有助于巩固发音基础与语感培养")

    # 亮点4：标杆班级
    if top and top[0]['avg_score'] >= 75:
        highlights.append(f"典型班级示范作用较为突出：{tc_name}平均得分率为{top[0]['avg_score']}%，位居全校前列，可为同年级教学改进提供参考样本")

    # ── 题型数据增强亮点（可选）──────────────────────────────
    if data.get('has_question_type'):
        qt_school = data.get('qt_school', {})
        qt_teacher = data.get('qt_teacher_rank', {})
        if qt_school:
            hardest = min(qt_school.items(), key=lambda x: x[1]['mean'])
            easiest = max(qt_school.items(), key=lambda x: x[1]['mean'])
            highlights.append(f"题型表现差异较为清晰：{easiest[0]}得分率最高（{easiest[1]['mean']}%），整体表现相对稳定；{hardest[0]}得分率相对较低（{hardest[1]['mean']}%，低分率{data.get('qt_hr_lr', {}).get(hardest[0], {}).get('低分率', 'N/A')}%），可作为后续专项改进的重点方向")
        if qt_teacher:
            best_t = max(qt_teacher.items(), key=lambda x: x[1])
            worst_t = min(qt_teacher.items(), key=lambda x: x[1])
            if best_t[0] != worst_t[0]:
                highlights.append(f"教师间教学结果存在一定差异：{best_t[0]}所带班级综合得分率为{best_t[1]}%，较{worst_t[0]}（{worst_t[1]}%）高{best_t[1]-worst_t[1]:.1f}个百分点，相关经验可纳入校本教研交流范畴")

    # ── 建议推导（与问题一一对应）────────────────────────────
    suggestions = []
    action_rows = []
    if any('同步' in iss for iss in issues) or any('实战' in iss for iss in issues):
        action = "优化训练结构：适度提升专项训练与模拟训练频次（建议各占总量15%~20%），进一步完善日常训练与阶段检测相结合的实施体系"
        suggestions.append(action)
        action_rows.append(("训练结构不均衡", action, "专项/模拟占比、月度模拟作业次数"))
    if any('不均衡' in iss or '差异' in iss for iss in issues):
        action = "加强分层推进：重点关注完成率偏低班级，结合具体原因开展分类指导，推动年级内应用质量更加均衡"
        suggestions.append(action)
        action_rows.append(("年级或班级差异", action, "年级完成率差值、低完成率班级数量"))
    if any('自主' in iss for iss in issues):
        action = "强化自主学习引导：通过过程性评价与正向激励相结合的方式，提升词汇自主练习覆盖率，逐步培育学生自主巩固习惯"
        suggestions.append(action)
        action_rows.append(("自主练习不足", action, "词汇自主练习次数、参与学生覆盖率"))
    if any('完成率' in iss for iss in issues):
        action = "强化作业过程管理：建立作业完成率监测与预警机制，对持续偏低班级开展原因分析与定向改进"
        suggestions.append(action)
        action_rows.append(("完成率偏低", action, "作业完成率、连续低完成率班级名单"))
    # 保底建议（必定有）
    if len(suggestions) == 0:
        action = "深化数据应用：持续跟踪核心指标变化，定期复盘教学策略，将数据分析结果更好融入校本教研与教学改进过程"
        suggestions.append(action)
        action_rows.append(("持续优化", action, "月度作业量、完成率、得分率、训练结构占比"))

    # ── 输出 6.1 主要亮点 ──────────────────────────────────
    L.append("### 6.1 主要亮点\n\n")
    for i, h in enumerate(highlights, 1):
        L.append(f"**亮点{i}：{h}**\n\n")

    # ── 输出 6.2 问题与不足 ───────────────────────────────
    L.append("### 6.2 问题与不足\n\n")
    if issues:
        for issue in issues:
            L.append(f"- **{issue}**\n\n")
    else:
        L.append("本期应用中未发现突出结构性问题，各项指标总体平稳。\n\n")

    # ── 输出 6.3 下阶段计划 ────────────────────────────────
    L.append("### 6.3 下阶段计划\n\n")
    L.append("| 对应问题 | 行动建议 |\n")
    L.append("|---------|---------|\n")
    for problem, action, metric in action_rows:
        L.append(f"| {problem} | {action} |\n")
    L.append("| 听说专项提升 | 适度增加听说模拟套卷训练频次，引导学生结合答题录音开展自我诊断与纠音练习，提升训练的针对性与实效性 |\n")
    L.append("\n")

    return ''.join(L)


def generate_district_report_text(data):
    region = data.get('region_name', '区域')
    months = sorted(data.get('monthly_hw', {}).keys())
    mr = _format_month_label(data.get('month_range', 'N/A'))
    schools = data.get('schools', 0)
    classes = data.get('classes', 0)
    students = data.get('total_students', 0)
    total_hw = data.get('total_hw', 0)
    completion = data.get('completion_rate', 0)
    avg_apc = data.get('district_avg_assign_per_class', 0)

    L = []
    L.append(f"# {region}英语AI听说训练区级应用成效报告\n")
    L.append(f"**报告生成时间：{datetime.now().strftime('%Y年%m月%d日')}**\n")
    L.append(f"**模板版本：{REPORT_VERSION} 区级标准模板**\n\n")

    L.append("## 一、平台使用概况\n")
    L.append(f"　　本报告基于{region}{mr}期间的班级数据总览、作业明细和听说模拟题型数据生成，重点从区级整体应用、学校推进差异、教师使用画像、学生训练表现和教研建议等维度，为后续区域教研与过程管理提供参考。\n\n")
    L.append("| 指标 | 数值 |\n|------|------|\n")
    L.append(f"| 覆盖学校 | {schools}所 |\n")
    L.append(f"| 覆盖班级 | {classes}个 |\n")
    L.append(f"| 覆盖学生 | {students:,}人 |\n")
    L.append(f"| 作业明细记录 | {total_hw:,}条 |\n")
    L.append(f"| 平均完成率 | {completion}% |\n")
    L.append(f"| 班均布置次数 | {avg_apc}次/班 |\n")
    if data.get('answer_rate') is not None:
        L.append(f"| 平均作答率 | {data['answer_rate']}% |\n")
    if data.get('avg_minutes') is not None:
        L.append(f"| 单次作业平均耗时 | {data['avg_minutes']}分钟 |\n")
    L.append("\n")
    L.append(f"　　从全区整体看，统计周期内共覆盖{schools}所学校、{classes}个班级和{students:,}名学生，作业明细共记录{total_hw:,}条，平均完成率为{completion}%。下一步分析将从学校推进差异、教师应用方式、学生训练表现和题型成效等方面展开。\n\n")

    L.append("## 二、平台使用情况\n")
    L.append("### 2.1 全区月度应用趋势\n")
    if months:
        peak_m = max(months, key=lambda m: data.get('monthly_hw', {}).get(m, 0))
        L.append(f"　　全区作业应用覆盖{_format_month_label(months[0])}至{_format_month_label(months[-1])}，其中{_format_month_label(peak_m)}作业量最高，为{data['monthly_hw'][peak_m]:,}条。月度变化可用于观察区域整体推进节奏和阶段训练重点。\n\n")
    grade_hw = data.get('grade_monthly_hw', {})
    # 仅 1 个年级时，年级分布语句（及图2）无对比意义，整体省略
    if grade_hw and len(grade_hw) > 1:
        grade_totals = {g: sum(v.values()) for g, v in grade_hw.items()}
        top_grade = max(grade_totals, key=grade_totals.get)
        L.append(f"　　从年级分布看，{top_grade}作业量最高，为{grade_totals[top_grade]:,}条。年级间差异提示与年级任务难度、学校推进节奏等有关。\n\n")

    L.append("## 三、学校及教师使用画像\n")
    L.append("### 3.1 学校应用画像\n")
    L.append(f"　　考虑学校班级规模差异，本报告采用“班均有效布置指数”观察学校应用推进质量。该指标先筛选班均布置次数高于全区平均值（{avg_apc}次/班）的学校，再按“班均布置次数×平均完成率”排序，不纳入得分率，得分率仅作为结果参考。\n\n")
    top_schools = data.get('active_school_top', [])
    if top_schools:
        L.append("**班均有效布置指数Top学校（约占学校总数30%）：**\n\n")
        L.append("| 学校 | 班级数 | 学生数 | 布置次数 | 班均布置次数 | 完成率 | 班均有效布置指数 |\n|------|------|------|------|------|------|------|\n")
        for row in top_schools:
            L.append(f"| {row['学校']} | {row['班级数']} | {row['学生数']:,} | {row['布置次数']} | {row['班均布置次数']} | {row['完成率']}% | {row['班均有效布置指数']} |\n")
        L.append("\n")

    low_schools = data.get('low_active_schools', [])
    L.append(f"　　低活跃待跟进学校共{len(low_schools)}所。该列表基于班均布置次数低于或等于全区平均值生成，并按照班均布置次数由高到低排序，可供后续区级过程跟进参考。\n\n")
    if low_schools:
        L.append("| 学校 | 班级数 | 学生数 | 布置次数 | 班均布置次数 | 完成率 |\n|------|------|------|------|------|------|\n")
        for row in low_schools:
            L.append(f"| {row['学校']} | {row['班级数']} | {row['学生数']:,} | {row['布置次数']} | {row['班均布置次数']} | {row['完成率']}% |\n")
        L.append("\n")

    quad = data.get('class_quadrant')
    if quad:
        L.append("### 3.2 班级应用四象限对比\n")
        L.append(f"　　以全区班级平均完成率（{quad['mc']}%）与平均得分率（{quad['ms']}%）为分界，"
                 f"对布置5次及以上的{quad['n']}个班级做四象限定位：高完成低得分提示训练方法待改进，"
                 "低完成提示督促机制待加强，可供教研组按班施策。\n\n")
        L.append("| 象限 | 班级数 | 占比 |\n|------|------|------|\n")
        for name, cnt in quad['counts']:
            L.append(f"| {name} | {cnt} | {round(cnt / max(quad['n'], 1) * 100, 1)}% |\n")
        L.append("\n")
        if quad['method']:
            L.append("**“高完成低得分”典型班级（方法待改进，建议加强作业讲评与错题归因）：**\n\n")
            L.append("| 学校 | 班级 | 布置次数 | 完成率 | 得分率 | 作答率 |\n|------|------|------|------|------|------|\n")
            for r in quad['method']:
                L.append(f"| {r['学校']} | {r['班级']} | {r['布置次数']} | {r['完成率']}% | {r['得分率']}% | {r['作答率']}% |\n")
            L.append("\n")
        if quad['follow']:
            L.append("**“双低”典型班级（重点跟进，先核查作答覆盖与督促机制）：**\n\n")
            L.append("| 学校 | 班级 | 布置次数 | 完成率 | 得分率 | 作答率 |\n|------|------|------|------|------|------|\n")
            for r in quad['follow']:
                L.append(f"| {r['学校']} | {r['班级']} | {r['布置次数']} | {r['完成率']}% | {r['得分率']}% | {r['作答率']}% |\n")
            L.append("\n")
        L.append("### 3.3 教师使用画像\n")
    else:
        L.append("### 3.2 教师使用画像\n")
    teacher_top = data.get('teacher_active_top', [])
    if teacher_top:
        L.append("　　教师层面同样参考班均布置次数和完成率形成应用画像，得分率作为结果表现展示，不参与排序规则。\n\n")
        L.append("| 学校 | 教师 | 班级数 | 布置次数 | 班均布置次数 | 完成率 | 得分率 |\n|------|------|------|------|------|------|------|\n")
        for row in teacher_top:
            L.append(f"| {row['学校']} | {row['教师']} | {row['班级数']} | {row['布置次数']} | {row['班均布置次数']} | {row['完成率']}% | {row['得分率']}% |\n")
        L.append("\n")

    L.append("## 四、学生使用画像\n")
    grade_stats = data.get('grade_class_stats', {})
    grade_profiles = data.get('grade_profiles', {})
    if grade_stats:
        L.append("| 年级 | 班级数 | 学生数 | 布置次数 | 完成率 | 得分率 | 自主练习次数 |\n|------|------|------|------|------|------|------|\n")
        for grade in _sort_grades(grade_stats.keys()):
            g = grade_stats[grade]
            p = grade_profiles.get(grade, {})
            L.append(f"| {grade} | {p.get('classes', '—')} | {p.get('students', 0):,} | {g['hw_times']} | {g['completion_rate']}% | {g['score_rate']}% | {p.get('self_practice', '—')} |\n")
        L.append("\n")
        if len(grade_stats) > 1:
            best_completion = max(grade_stats, key=lambda g: grade_stats[g]['completion_rate'])
            L.append(f"　　年级间差异提示与年级任务难度、学校推进节奏等有关。其中{best_completion}完成率相对较高，可作为观察年级推进节奏的参考样本。\n\n")

    if data.get('avg_minutes') is not None:
        heavy = data.get('heavy_minute_schools') or []
        if heavy:
            heavy_txt = '、'.join(f"{s}（{v}分钟）" for s, v in heavy)
            L.append(f"　　从作业负担看，全区单次作业平均耗时{data['avg_minutes']}分钟。"
                     f"其中{heavy_txt}超过20分钟，建议将长套题拆分为两次短训布置，控制单次时长。\n\n")
        else:
            L.append(f"　　从作业负担看，全区单次作业平均耗时{data['avg_minutes']}分钟，"
                     "各校单次耗时均在20分钟以内，负担总体可控。\n\n")

    L.append("## 五、训练内容与阶段成效分析\n")
    L.append("### 5.1 各校训练结构\n")
    cat_counts = data.get('category_counts', {})
    if cat_counts:
        total = sum(cat_counts.values())
        top_cat = max(cat_counts, key=cat_counts.get)
        L.append(f"　　从全区训练栏目结构看，「{top_cat}」训练量最高，共{cat_counts[top_cat]:,}次，占比{round(cat_counts[top_cat] / max(total, 1) * 100, 1)}%。各校各栏目训练结构见图5，可用于观察学校训练内容是否过于集中。\n\n")
        L.append("| 栏目 | 次数 | 占比 |\n|------|------|------|\n")
        for cat, cnt in sorted(cat_counts.items(), key=lambda x: -x[1]):
            L.append(f"| {cat} | {cnt:,} | {data.get('category_pct', {}).get(cat, 0)}% |\n")
        L.append("\n")

    L.append("### 5.2 典型学校年级分析\n")
    case_top = data.get('district_case_top5', [])
    if case_top:
        sample = case_top[0]
        L.append(f"　　以{sample['学校']}{sample['年级']}为例，该学校年级总作业次数为{sample['总作业次数']}次，听说模拟次数为{sample['听说模拟次数']}次，统计周期内提分率为{sample['提分率']}个百分点。下表按“听说模拟次数→完成率→提分率”的顺序筛选全区Top5学校年级样本，其中提分率为统计周期内听说模拟月均最高得分率减最低得分率。\n\n")
        L.append("| 排名 | 学校 | 年级 | 总作业次数 | 听说模拟次数 | 提分率 |\n|------|------|------|------|------|------|\n")
        for idx, row in enumerate(case_top, 1):
            L.append(f"| {idx} | {row['学校']} | {row['年级']} | {row['总作业次数']} | {row['听说模拟次数']} | {row['提分率']}个百分点 |\n")
        L.append("\n")

    if data.get('has_question_type') and data.get('qt_school'):
        L.append("### 5.3 听说模拟题型表现\n")
        L.append("结合全区听说模拟题型，以下是各题型的得分率展示。\n\n")
        L.append("| 题型 | 平均得分率 | 样本数 |\n|------|------|------|\n")
        qt_school = data.get('qt_school', {})
        for qtype, info in sorted(qt_school.items(), key=lambda x: x[1]['mean']):
            L.append(f"| {qtype} | {info['mean']}% | {info['count']} |\n")
        L.append("\n")
        weak_by_grade = data.get('qt_weak_by_grade', {})
        if weak_by_grade:
            parts = []
            for grade in _sort_grades(weak_by_grade.keys()):
                items = weak_by_grade.get(grade, [])[:3]
                if items:
                    parts.append(f"{grade}相对薄弱题型为" + "、".join([f"{i['题型']}（{i['得分率']}%）" for i in items]))
            if parts:
                L.append("　　从各年级各题型看，" + "；".join(parts) + "。低于60%的题型在图7中以红色标识，建议作为后续专题教研和专项训练的优先关注对象。\n\n")
        # 薄弱题型 → 课堂训练动作（仅取样本数≥100 的题型，零星题型不具教研行动价值）
        qt_pool = [(q, info) for q, info in qt_school.items() if info.get('count', 0) >= 100]
        qt_weak3 = sorted(qt_pool, key=lambda x: x[1]['mean'])[:3]
        if qt_weak3:
            L.append("**薄弱题型对应课堂训练动作（通用建议）：**\n\n")
            L.append("| 薄弱题型 | 得分率 | 建议课堂动作 |\n|------|------|------|\n")
            for q, info in qt_weak3:
                L.append(f"| {q} | {info['mean']}% | {QT_ACTIONS.get(q, '专项讲评＋限时训练')} |\n")
            L.append("\n")

    exam = data.get('exam')
    # 有联考数据时插入“六、阶段性联考成绩对比”，后续小节顺延为七、八
    if exam:
        _append_exam_section(L, data, exam, region)
        n_research, n_summary = '七', '八'
    else:
        n_research, n_summary = '六', '七'

    L.append(f"## {n_research}、教研数据板块\n")
    L.append("　　建议区级教研侧将以上数据作为过程观察依据，重点关注学校班均应用深度、教师布置节奏、年级题型短板和典型学校年级样本变化，形成“数据观察—问题定位—专项训练—阶段复盘”的持续改进链条。\n\n")

    L.append(f"## {n_summary}、总结与建议\n")
    if exam:
        # 将联考相关结论融入总结，避免与第六章重复
        bench = exam['link'][0] if exam.get('link') else None
        bench_txt = f"，{bench['短名']}等学校表现突出" if bench else ''
        L.append(f"　　总体来看，{region}已形成较扎实的英语AI听说训练基础，且日常训练成效在联考中得到正向印证"
                 f"（练习得分率与联考平均分相关系数 r={exam['r_score']}{bench_txt}）。"
                 "后续建议：①提炼头部学校“以练促考”经验并在区域内推广；"
                 "②围绕联考薄弱题型开展专项训练；"
                 "③推动尚未接入或训练频次偏低的学校补齐日常练习；"
                 "④以“练—测—评—改”闭环持续跟踪成效，服务区域教研决策。\n\n")
    else:
        quad = data.get('class_quadrant')
        low_n = len(data.get('low_active_schools', []))
        _qt_pool = [(q, i) for q, i in (data.get('qt_school') or {}).items() if i.get('count', 0) >= 100]
        weak_qt = sorted(_qt_pool, key=lambda x: x[1]['mean'])[0][0] if _qt_pool else ''
        L.append(f"　　总体来看，{region}已形成一定规模的英语AI听说训练数据基础，下阶段建议按“谁—何时—做什么—怎么验收”落实：\n\n")
        if quad:
            L.append("　　一是区教研员牵头，第2周前组织“双高”示范班级所在校做1次区级经验分享，重点萃取作业讲评与督促做法，以分享材料归档为验收；\n\n")
            L.append("　　二是各校教研组长于第3周前，对本校“双低”与“高完成低得分”班级各听评课1次，形成按班改进清单，下期报告复核象限迁移情况；\n\n")
        else:
            L.append("　　一是区教研员牵头，组织班均有效布置指数领先学校做1次区级经验分享，重点萃取作业布置与督促做法，以分享材料归档为验收；\n\n")
            L.append("　　二是各校教研组长对完成率偏低班级开展听评课与作业讲评检查，形成按班改进清单，下期报告复核变化；\n\n")
        L.append(f"　　三是对{low_n}所低活跃学校，由区管理员按月核对账号启用与布置量，设定“班均布置次数达到全区均线”的阶段目标，下期报告验收达标率；\n\n")
        if weak_qt:
            L.append(f"　　四是各年级备课组围绕「{weak_qt}」等薄弱题型，按课堂训练动作建议开展4周专项训练，以下期题型得分率提升3个百分点以上为验收线。\n\n")
        else:
            L.append("　　四是围绕训练结构优化与高频资源教研持续推进，使区级应用分析更好服务区域教研决策。\n\n")
    return ''.join(L)


def _append_exam_section(L, data, exam, region):
    """六、阶段性联考成绩对比（联考成绩 + “练—考”关联，合二为一）。"""
    latest = exam['latest']
    ov = latest.get('overall', {})
    fm = exam.get('full_mark', 30)
    stages = exam.get('stages', [])
    multi = exam.get('multi')

    L.append("## 六、阶段性联考成绩对比\n")

    # 6.1 联考整体情况
    L.append("### 6.1 联考整体情况\n")
    avg = ov.get('平均分')
    parts = [f"{latest['label']}共 {latest['school_count']} 所学校参加，满分 {fm:.0f} 分"]
    if ov.get('学生人数'):
        parts.append(f"报名 {ov['学生人数']:,} 人、实考 {ov.get('实考人数', 0):,} 人")
    if avg is not None:
        parts.append(f"全区平均分 {avg} 分")
    if latest.get('score_rate'):
        parts.append(f"整体得分率约 {latest['score_rate']}%")
    L.append("　　" + "，".join(parts) + "。\n\n")
    if multi:
        labels = [s['short'] for s in stages]
        trend = [s['overall'].get('平均分') for s in stages]
        valid = [(l, v) for l, v in zip(labels, trend) if v is not None]
        if len(valid) >= 2:
            delta = valid[-1][1] - valid[0][1]
            trend_word = '上升' if delta > 0 else ('回落' if delta < 0 else '基本持平')
            L.append(f"　　从 {valid[0][0]} 到 {valid[-1][0]}，全区整体均分由 {valid[0][1]} 分{trend_word}"
                     f"至 {valid[-1][1]} 分（变化 {delta:+.2f} 分），与同期日常听说训练的推进节奏基本对应。\n\n")
        L.append("| 阶段 | 参考学校 | 实考人数 | 平均分 | 得分率 |\n|------|------|------|------|------|\n")
        for s in stages:
            so = s['overall']
            L.append(f"| {s['label']} | {s['school_count']} | {so.get('实考人数', 0):,} | "
                     f"{so.get('平均分', '—')} | {s['score_rate'] if s['score_rate'] else '—'}% |\n")
        L.append("\n")

    # 6.2 各校联考成绩对比
    L.append("### 6.2 各校联考成绩对比\n")
    link = exam.get('link', [])
    sch = latest.get('schools', {})
    if avg is not None and sch:
        ranked = sorted([(n, v.get('平均分')) for n, v in sch.items() if v.get('平均分') is not None],
                        key=lambda x: x[1], reverse=True)
        if ranked:
            quals = exam.get('_quals', [])
            top_n, top_v = ranked[0]
            low_n, low_v = ranked[-1]
            L.append(f"　　全区平均分 {avg} 分，平均分最高的是{_exam_norm(top_n, quals)}（{top_v} 分），"
                     f"最低的是{_exam_norm(low_n, quals)}（{low_v} 分）。图中绿色为高于全区均线、红色为低于均线"
                     "，灰色为暂无平台练习数据的学校。\n\n")
    if multi and exam.get('school_stage'):
        # 提分明显的学校
        gains = []
        for name, series in exam['school_stage'].items():
            vals = [v for v in series if v is not None]
            if len(vals) >= 2:
                gains.append((name, vals[-1] - vals[0]))
        gains.sort(key=lambda x: x[1], reverse=True)
        up = [g for g in gains if g[1] > 0][:3]
        if up:
            L.append("　　跨阶段提分明显的学校：" + "、".join(f"{n}（{d:+.2f}分）" for n, d in up) + "。\n\n")

    # 6.3 各题型联考表现
    L.append("### 6.3 各题型联考表现\n")
    qt = latest.get('qtype', {})
    qnames = [q for q in EXAM_QTYPES if q in qt]
    if qnames:
        metrics = ['平均分', '得分率', '满分率', '优秀率', '低分率']
        L.append("| 题型 | " + " | ".join(qnames) + " |\n")
        L.append("|------|" + "------|" * len(qnames) + "\n")
        for m in metrics:
            row = [qt.get(q, {}).get(m, '—') for q in qnames]
            L.append(f"| {m} | " + " | ".join(row) + " |\n")
        L.append("\n")
        rates = [(q, _safe_float(qt[q].get('得分率'))) for q in qnames if _safe_float(qt[q].get('得分率')) is not None]
        if rates:
            hi = max(rates, key=lambda x: x[1])
            lo = min(rates, key=lambda x: x[1])
            L.append(f"　　得分率最高的是「{hi[0]}」（{hi[1]}%），最低的是「{lo[0]}」（{lo[1]}%），"
                     "与日常练习中的强项、弱项基本一致。\n\n")

    # 6.4 训练与联考关联
    L.append("### 6.4 训练与联考关联\n")
    if link and len(link) >= 3:
        L.append(f"　　将各校日常听说练习得分率与联考平均分按校名模糊匹配后分析，二者呈正相关，"
                 f"相关系数 r = {exam['r_score']}：日常练习越扎实的学校，联考平均分越高，"
                 "日常训练的投入在联考中得到正向回报。\n\n")
        head = link[:3]
        bench = "、".join(f"{r['短名']}（练习得分率{r['练习得分率']:.0f}%、联考{r['联考平均分']}分）" for r in head)
        L.append(f"　　正向标杆：{bench} 等学校，日常练习扎实、听说模拟投入充分，联考平均分稳居全区前列，"
                 "可作为区域教研观摩与经验提炼对象。\n\n")
    qtc = exam.get('qt_compare', [])
    cons = [c for c in qtc if c['练习'] is not None and c['联考'] is not None]
    if cons:
        strong = max(cons, key=lambda c: c['联考'])
        weak = min(cons, key=lambda c: c['联考'])
        L.append(f"　　题型层面，日常练习与联考走势一致：「{strong['题型']}」在练、考两端均为相对强项，"
                 f"「{weak['题型']}」在两端均为相对弱项（日常练习{weak['练习']}%、联考{weak['联考']}%），"
                 "宜作为下一阶段专项训练重点。\n\n")
    only_e = exam.get('only_exam', [])
    if only_e:
        quals = exam.get('_quals', [])
        scores = exam.get('only_exam_scores', {})
        lines = "、".join(f"{_exam_short(s, quals)}（联考{scores.get(_exam_short(s, quals), '—')}分）" for s in only_e)
        L.append(f"　　参加联考但在平台基本无日常练习数据的学校有：{lines}，其联考成绩多位于全区偏后位置，"
                 "建议优先纳入平台日常听说训练，建立常态化练习节奏，先把过程补齐，再谈提升。\n\n")
    low = sorted([r for r in link if avg is not None and r['联考平均分'] < avg],
                 key=lambda r: r['班均听说模拟'])[:3]
    if low:
        low_txt = "、".join(f"{r['短名']}（班均听说模拟{r['班均听说模拟']:.1f}、联考{r['联考平均分']}分）" for r in low)
        L.append(f"　　已接入平台、但联考低于全区均线的学校中，听说模拟投入相对偏低的有：{low_txt}。"
                 f"考虑到听说模拟次数与联考成绩的相关性（r={exam['r_freq']}）弱于练习得分率（r={exam['r_score']}），"
                 "提升重点不在“多布置”，而在“练扎实、练到位”。\n\n")


def make_charts(data):
    charts = {}
    FONT_FAMILY = 'Noto Sans CJK SC, Microsoft YaHei, SimHei, Arial Unicode MS, sans-serif'
    BASE_FONT_SIZE = 19
    LABEL_FONT_SIZE = 20
    BAR_LABEL_FONT_SIZE = 18
    LEGEND_FONT_SIZE = 18
    CC = {'同步': '#4C78A8', '专项': '#F58518', '模拟': '#E45756', '课外拓展': '#72D7B8', '其他': '#9CA3AF'}
    preferred_cats = ['同步', '专项', '模拟', '课外拓展', '其他']
    all_cats = set(data.get('category_counts', {}).keys())
    for month_data in data.get('cat_monthly', {}).values():
        all_cats.update(month_data.keys())
    for mix in data.get('grade_category_mix', {}).values():
        all_cats.update(mix.keys())
    cats = [c for c in preferred_cats if c in all_cats] + sorted(c for c in all_cats if c not in preferred_cats)
    months = sorted(data.get('monthly_hw', {}).keys())
    month_labels = [_format_month_label(m) for m in months]

    totals = [data['monthly_hw'].get(m, 0) for m in months]
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=month_labels, y=totals, mode='lines+markers+text',
        line=dict(color='#2E86AB', width=3),
        marker=dict(size=10, color='#2E86AB'),
        fill='tozeroy', fillcolor='rgba(46,134,171,0.1)',
        text=totals, textposition='top center', textfont=dict(size=LABEL_FONT_SIZE, family=FONT_FAMILY, color='#1f2937'),
        name='作业总量', hovertemplate='%{x}<br>作业量：%{y}次<extra></extra>'
    ))
    fig.update_layout(
        title=None,
        xaxis_title='月份', yaxis_title='作业次数',
        height=460, template='plotly_white', hovermode='x unified',
        font=dict(size=BASE_FONT_SIZE, color='#1f2937', family=FONT_FAMILY),
        margin=dict(l=70, r=40, t=90, b=70),
        xaxis=dict(tickangle=0, automargin=True),
        yaxis=dict(automargin=True)
    )
    charts['monthly_line'] = fig

    grade_hw = data.get('grade_monthly_hw', {})
    fig2 = go.Figure()
    for idx, grade in enumerate(_sort_grades(grade_hw.keys())):
        gd = grade_hw.get(grade, {})
        y = [gd.get(m, 0) for m in months]
        fig2.add_trace(go.Scatter(
            name=grade, x=month_labels, y=y,
            mode='lines+markers+text', line=dict(width=2.5),
            marker=dict(size=7, color=_grade_color(grade, idx)),
            text=y, textposition='top center', textfont=dict(size=LABEL_FONT_SIZE, family=FONT_FAMILY, color='#1f2937'),
            cliponaxis=False
        ))
    fig2.update_layout(
        title=None,
        xaxis_title='月份', yaxis_title='作业次数',
        height=440, template='plotly_white',
        font=dict(size=BASE_FONT_SIZE, color='#1f2937', family=FONT_FAMILY),
        margin=dict(l=70, r=40, t=90, b=70),
        xaxis=dict(automargin=True),
        yaxis=dict(automargin=True)
    )
    charts['grade_monthly_line'] = fig2

    fig3 = go.Figure()
    for cat in cats:
        y = [data.get('cat_monthly', {}).get(m, {}).get(cat, 0) for m in months]
        fig3.add_trace(go.Bar(
            name=cat, x=month_labels, y=y, marker_color=CC.get(cat, '#999'),
            text=[v if v > 0 else '' for v in y], textposition='inside', textfont=dict(size=BAR_LABEL_FONT_SIZE, color='#1f2937', family=FONT_FAMILY), cliponaxis=False
        ))
    fig3.update_layout(
        barmode='stack',
        title=None,
        xaxis_title='月份', yaxis_title='作业次数',
        height=460, template='plotly_white',
        font=dict(size=BASE_FONT_SIZE, color='#1f2937', family=FONT_FAMILY),
        legend=dict(
            orientation='v',
            x=1.02, xanchor='left',
            y=1, yanchor='top',
            font=dict(size=LEGEND_FONT_SIZE, color='#1f2937', family=FONT_FAMILY)
        ),
        margin=dict(l=70, r=120, t=40, b=70),
        xaxis=dict(automargin=True),
        yaxis=dict(automargin=True)
    )
    charts['cat_stacked'] = fig3

    grade_mix = data.get('grade_category_mix', {})
    fig4 = go.Figure()
    grade_labels = _sort_grades(grade_mix.keys())
    for cat in cats:
        values = [grade_mix.get(grade, {}).get(cat, {}).get('pct', 0) for grade in grade_labels]
        fig4.add_trace(go.Bar(
            name=cat,
            x=grade_labels,
            y=values,
            marker_color=CC.get(cat, '#999'),
            text=[f"{v}%" if v > 0 else '' for v in values],
            textposition='inside',
            textfont=dict(size=BAR_LABEL_FONT_SIZE, family=FONT_FAMILY, color='#1f2937'),
            hovertemplate='%{x}<br>%{fullData.name}：%{y}%<extra></extra>',
            cliponaxis=False
        ))
    fig4.update_layout(
        barmode='stack',
        title=None,
        height=460, template='plotly_white',
        font=dict(size=BASE_FONT_SIZE, color='#1f2937', family=FONT_FAMILY),
        xaxis_title='年级', yaxis_title='训练占比',
        yaxis=dict(range=[0, 100], ticksuffix='%', automargin=True),
        legend=dict(
            orientation='v',
            x=1.02, xanchor='left',
            y=0.95, yanchor='top',
            font=dict(size=LEGEND_FONT_SIZE, color='#1f2937', family=FONT_FAMILY)
        ),
        margin=dict(l=70, r=140, t=40, b=70)
    )
    charts['cat_grade_pct'] = fig4

    grade_scores = data.get('mock_hw_grade_monthly', {})
    if grade_scores:
        fig6 = go.Figure()
        for idx, grade in enumerate(_sort_grades(grade_scores.keys())):
            gm = sorted(grade_scores[grade].items())
            xs = [_format_month_label(m) for m, s in gm]
            ys = [s for m, s in gm]
            fig6.add_trace(go.Scatter(
                name=grade, x=xs, y=ys,
                mode='lines+markers+text', line=dict(width=2.5),
                marker=dict(size=7, color=_grade_color(grade, idx)),
                text=[f"{v}%" for v in ys], textposition='top center', textfont=dict(size=LABEL_FONT_SIZE, family=FONT_FAMILY, color='#1f2937'),
                cliponaxis=False
            ))
        fig6.update_layout(
            title=None,
            xaxis_title='月份', yaxis_title='得分率（%）',
            height=430, template='plotly_white',
            font=dict(size=BASE_FONT_SIZE, color='#1f2937', family=FONT_FAMILY),
            yaxis=dict(range=[0, 100], automargin=True),
            margin=dict(l=70, r=40, t=90, b=70),
            xaxis=dict(automargin=True)
        )
        charts['grade_score'] = fig6

    top_all_m = data.get('top_class_all_monthly', {})
    top_mock_m = data.get('top_class_mock_monthly', {})
    if top_all_m and top_mock_m:
        all_months = sorted(set(top_all_m.keys()) | set(top_mock_m.keys()))
        all_month_labels = [_format_month_label(m) for m in all_months]
        sc_t  = [top_mock_m.get(m, {}).get('score', None) for m in all_months]
        ct_t  = [top_all_m.get(m, 0) for m in all_months]
        sc_t_fmt = [f"{s}%" if s is not None else '—' for s in sc_t]

        fig7 = make_subplots(specs=[[{"secondary_y": True}]])
        fig7.add_trace(go.Scatter(
            x=all_month_labels, y=sc_t, name='听说模拟得分率',
            mode='lines+markers+text',
            line=dict(color='#4C78A8', width=2.5), marker=dict(size=8),
            text=sc_t_fmt, textposition='top center', textfont=dict(size=LABEL_FONT_SIZE, family=FONT_FAMILY, color='#1f2937'),
            yaxis='y'
        ))
        fig7.add_trace(go.Bar(
            x=all_month_labels, y=ct_t, name='所有类目布置次数',
            opacity=0.35, marker_color='#F58518', yaxis='y2',
            text=ct_t, textposition='outside', textfont=dict(size=LABEL_FONT_SIZE, family=FONT_FAMILY, color='#1f2937')
        ))
        fig7.update_layout(
            title=None,
            template='plotly_white', height=430,
            font=dict(size=BASE_FONT_SIZE, color='#1f2937', family=FONT_FAMILY),
            legend=dict(
                orientation='v',
                x=1.02, xanchor='left',
                y=1, yanchor='top',
                font=dict(size=LEGEND_FONT_SIZE, color='#1f2937', family=FONT_FAMILY)
            ),
            hovermode='x unified',
            margin=dict(l=70, r=120, t=40, b=70),
            xaxis=dict(automargin=True),
            yaxis=dict(automargin=True)
        )
        fig7.update_layout(yaxis2=dict(title_text='布置次数', overlaying='y', side='right'))
        fig7.update_yaxes(title_text='得分率（%）', range=[0, 100])
        charts['top_class_trend'] = fig7

    return charts


def make_district_charts(data):
    charts = make_charts(data)
    # 仅 1 个年级时，图2（各年级月度作业量趋势）无对比意义，随正文语句一并省略
    if len(data.get('grade_monthly_hw', {}) or {}) <= 1:
        charts.pop('grade_monthly_line', None)
    FONT_FAMILY = 'Noto Sans CJK SC, Microsoft YaHei, SimHei, Arial Unicode MS, sans-serif'
    CC = {'同步': '#4C78A8', '专项': '#F58518', '模拟': '#E45756', '课外拓展': '#72D7B8', '其他': '#9CA3AF'}
    cats = ['同步', '专项', '模拟', '课外拓展', '其他']
    for cat in data.get('category_counts', {}).keys():
        if cat not in cats:
            cats.append(cat)

    # 图3：区域内全部学校的班均有效布置指数总列表（非仅 Top 学校）
    active = sorted(data.get('school_stats') or data.get('active_school_top', []),
                    key=lambda r: r['班均有效布置指数'], reverse=True)
    if active:
        fig = go.Figure(go.Bar(
            x=[r['班均有效布置指数'] for r in active[::-1]],
            y=[r['学校'] for r in active[::-1]],
            orientation='h',
            marker_color='#4C78A8',
            text=[r['班均有效布置指数'] for r in active[::-1]],
            textposition='outside',
            textfont=dict(size=18, family=FONT_FAMILY, color='#1f2937'),
            hovertemplate='%{y}<br>班均有效布置指数：%{x}<extra></extra>',
        ))
        fig.update_layout(
            title=None, template='plotly_white', height=max(420, 34 * len(active) + 120),
            font=dict(size=18, family=FONT_FAMILY, color='#1f2937'),
            xaxis_title='班均有效布置指数', yaxis_title='学校',
            margin=dict(l=260, r=80, t=40, b=70)
        )
        charts['district_school_active'] = fig

    school_mix = data.get('school_category_mix', {})
    if school_mix:
        schools = list(school_mix.keys())
        fig = go.Figure()
        for cat in cats:
            values = [school_mix.get(school, {}).get(cat, {}).get('pct', 0) for school in schools]
            fig.add_trace(go.Bar(
                name=cat,
                x=values,
                y=schools,
                orientation='h',
                marker_color=CC.get(cat, '#9CA3AF'),
                text=[f"{v}%" if v >= 8 else '' for v in values],
                textposition='inside',
                textfont=dict(size=18, family=FONT_FAMILY, color='#1f2937'),
                hovertemplate='%{y}<br>%{fullData.name}：%{x}%<extra></extra>',
            ))
        fig.update_layout(
            barmode='stack',
            title=None, template='plotly_white',
            height=max(620, 34 * len(schools) + 150),
            font=dict(size=18, family=FONT_FAMILY, color='#1f2937'),
            xaxis=dict(title='训练占比', range=[0, 100], ticksuffix='%'),
            yaxis=dict(title='学校', automargin=True),
            legend=dict(orientation='h', x=0.5, xanchor='center', y=1.08, yanchor='bottom', font=dict(size=18, family=FONT_FAMILY)),
            margin=dict(l=280, r=70, t=80, b=70)
        )
        charts['district_school_category_pct'] = fig

    case_top = data.get('district_case_top5', [])
    if case_top:
        labels = [f"{r['学校']} {r['年级']}" for r in case_top]
        fig = make_subplots(specs=[[{"secondary_y": True}]])
        fig.add_trace(go.Bar(
            x=labels,
            y=[r['听说模拟次数'] for r in case_top],
            name='听说模拟次数',
            marker_color='#4C78A8',
            text=[r['听说模拟次数'] for r in case_top],
            textposition='outside',
            textfont=dict(size=18, family=FONT_FAMILY, color='#1f2937')
        ), secondary_y=False)
        fig.add_trace(go.Scatter(
            x=labels,
            y=[r['提分率'] for r in case_top],
            name='提分率',
            mode='lines+markers+text',
            line=dict(color='#E45756', width=2.5),
            marker=dict(size=8),
            text=[f"{r['提分率']}" for r in case_top],
            textposition='top center',
            textfont=dict(size=18, family=FONT_FAMILY, color='#1f2937')
        ), secondary_y=True)
        fig.update_layout(
            title=None, template='plotly_white', height=520,
            font=dict(size=18, family=FONT_FAMILY, color='#1f2937'),
            legend=dict(orientation='h', x=0.5, xanchor='center', y=1.08, yanchor='bottom', font=dict(size=18, family=FONT_FAMILY)),
            margin=dict(l=70, r=80, t=80, b=150),
            xaxis=dict(tickangle=25, automargin=True)
        )
        fig.update_yaxes(title_text='听说模拟次数', secondary_y=False)
        fig.update_yaxes(title_text='提分率（百分点）', secondary_y=True)
        charts['district_case_grade'] = fig

    qt_grade = data.get('qt_grade', {})
    if qt_grade:
        grades = _sort_grades(qt_grade.keys())
        qtypes = sorted({q for items in qt_grade.values() for q in items.keys()})
        z = [[qt_grade.get(grade, {}).get(q, None) for grade in grades] for q in qtypes]
        text = [[f"{v}%" if v is not None else "" for v in row] for row in z]
        fig = go.Figure(go.Heatmap(
            z=z,
            x=grades,
            y=qtypes,
            text=text,
            texttemplate='%{text}',
            colorscale=[[0, '#ef4444'], [0.6, '#ef4444'], [0.6001, '#bfdbfe'], [1, '#2563eb']],
            zmin=0,
            zmax=100,
            colorbar=dict(title='得分率'),
            hovertemplate='%{x}<br>%{y}：%{z}%<extra></extra>',
        ))
        fig.update_layout(
            title=None, template='plotly_white',
            height=max(640, 34 * len(qtypes) + 150),
            font=dict(size=18, family=FONT_FAMILY, color='#1f2937'),
            xaxis_title='年级', yaxis_title='题型',
            margin=dict(l=180, r=120, t=40, b=80)
        )
        charts['district_qt_grade'] = fig

    _make_exam_charts(charts, data, FONT_FAMILY)
    return charts


def _make_exam_charts(charts, data, FONT_FAMILY):
    """阶段性联考相关图表（仅在上传联考报告时生成）。"""
    exam = data.get('exam')
    if not exam:
        return
    fm = exam.get('full_mark', 30)
    latest = exam.get('latest', {})
    avg = latest.get('overall', {}).get('平均分')
    multi = exam.get('multi')
    quals = exam.get('_quals', [])

    # 全区整体均分阶段趋势（多场）
    if multi:
        labels = exam.get('stage_labels', [])
        trend = exam.get('overall_trend', [])
        pts = [(l, v) for l, v in zip(labels, trend) if v is not None]
        if len(pts) >= 2:
            fig = go.Figure(go.Scatter(
                x=[p[0] for p in pts], y=[p[1] for p in pts], mode='lines+markers+text',
                line=dict(color='#4C78A8', width=3), marker=dict(size=12),
                text=[f"{p[1]}" for p in pts], textposition='top center',
                textfont=dict(size=18, family=FONT_FAMILY)))
            fig.update_layout(title=None, template='plotly_white', height=460,
                              font=dict(size=18, family=FONT_FAMILY, color='#1f2937'),
                              xaxis_title='联考阶段', yaxis_title=f'全区平均分（满分{fm:.0f}）',
                              margin=dict(l=80, r=60, t=40, b=70))
            charts['exam_trend'] = fig

    # 各校联考平均分（单场=彩色条形；多场=分组柱状）
    school_stage = exam.get('school_stage', {})
    stage_labels = exam.get('stage_labels', [])
    if multi and school_stage and len(stage_labels) >= 2:
        # 按最新阶段降序排列学校
        order = sorted(school_stage.keys(),
                       key=lambda n: (school_stage[n][-1] is None, -(school_stage[n][-1] or 0)))
        fig = go.Figure()
        palette = ['#9CA3AF', '#72B7B2', '#F58518', '#4C78A8', '#E45756']
        for si, lab in enumerate(stage_labels):
            fig.add_trace(go.Bar(
                name=lab, x=order, y=[school_stage[n][si] for n in order],
                marker_color=palette[si % len(palette)],
                text=[f"{school_stage[n][si]:.1f}" if school_stage[n][si] is not None else '' for n in order],
                textposition='outside', textfont=dict(size=13, family=FONT_FAMILY)))
        if avg is not None:
            fig.add_hline(y=avg, line_dash='dash', line_color='#6b7280',
                          annotation_text=f'最新全区平均 {avg}', annotation_font_size=14)
        fig.update_layout(barmode='group', title=None, template='plotly_white',
                          height=max(520, 30 * len(order) + 180),
                          font=dict(size=16, family=FONT_FAMILY, color='#1f2937'),
                          yaxis=dict(title=f'联考平均分（满分{fm:.0f}）'),
                          xaxis=dict(tickangle=30, automargin=True),
                          legend=dict(orientation='h', x=0.5, xanchor='center', y=1.06, yanchor='bottom',
                                      font=dict(size=15, family=FONT_FAMILY)),
                          margin=dict(l=70, r=50, t=80, b=140))
        charts['exam_school'] = fig
    elif latest.get('schools'):
        only_e = set(exam.get('only_exam', []))
        rows = sorted([(n, v.get('平均分')) for n, v in latest['schools'].items() if v.get('平均分') is not None],
                      key=lambda x: x[1])
        names, vals, colors = [], [], []
        for n, v in rows:
            is_np = n in only_e
            names.append(_exam_short(n, quals) + ('（无平台练习）' if is_np else ''))
            vals.append(v)
            colors.append('#9CA3AF' if is_np else ('#54A24B' if (avg is not None and v >= avg) else '#E45756'))
        fig = go.Figure(go.Bar(x=vals, y=names, orientation='h', marker_color=colors,
                               text=[f"{v:.1f}" for v in vals], textposition='outside',
                               textfont=dict(size=16, family=FONT_FAMILY)))
        if avg is not None:
            fig.add_vline(x=avg, line_dash='dash', line_color='#6b7280',
                          annotation_text=f'全区平均 {avg}', annotation_font_size=14)
        fig.update_layout(title=None, template='plotly_white', height=max(420, 34 * len(names) + 120),
                          font=dict(size=16, family=FONT_FAMILY, color='#1f2937'),
                          xaxis_title=f'联考平均分（满分{fm:.0f}）', yaxis_title='学校',
                          margin=dict(l=320, r=90, t=40, b=60))
        charts['exam_school'] = fig

    # 各题型：日常练习 vs 联考
    qtc = [c for c in exam.get('qt_compare', []) if c['练习'] is not None or c['联考'] is not None]
    if qtc:
        cats = [c['题型'] for c in qtc]
        fig = go.Figure()
        fig.add_trace(go.Bar(name='日常练习得分率', x=cats, y=[c['练习'] for c in qtc],
                             marker_color='#4C78A8',
                             text=[f"{c['练习']:.0f}%" if c['练习'] is not None else '' for c in qtc],
                             textposition='outside', textfont=dict(size=14, family=FONT_FAMILY)))
        fig.add_trace(go.Bar(name='联考得分率', x=cats, y=[c['联考'] for c in qtc],
                             marker_color='#F58518',
                             text=[f"{c['联考']:.0f}%" if c['联考'] is not None else '' for c in qtc],
                             textposition='outside', textfont=dict(size=14, family=FONT_FAMILY)))
        fig.update_layout(barmode='group', title=None, template='plotly_white', height=500,
                          font=dict(size=16, family=FONT_FAMILY, color='#1f2937'),
                          yaxis=dict(title='得分率', range=[0, 100], ticksuffix='%'),
                          xaxis=dict(tickangle=20, automargin=True),
                          legend=dict(orientation='h', x=0.5, xanchor='center', y=1.08, yanchor='bottom',
                                      font=dict(size=15, family=FONT_FAMILY)),
                          margin=dict(l=70, r=40, t=70, b=110))
        charts['exam_qt'] = fig

    # 练习得分率 × 联考平均分 散点
    link = exam.get('link', [])
    if len(link) >= 3:
        xs = [r['练习得分率'] for r in link]
        ys = [r['联考平均分'] for r in link]
        fig = go.Figure(go.Scatter(
            x=xs, y=ys, mode='markers+text', showlegend=False,
            text=[r['短名'] for r in link], textposition='top center',
            textfont=dict(size=13, family=FONT_FAMILY),
            marker=dict(size=14, color='#4C78A8', line=dict(width=1, color='#1f2937'))))
        n = len(xs)
        mx, my = sum(xs)/n, sum(ys)/n
        denom = sum((x-mx)**2 for x in xs)
        if denom:
            slope = sum((x-mx)*(y-my) for x, y in zip(xs, ys)) / denom
            intercept = my - slope*mx
            x0, x1 = min(xs), max(xs)
            fig.add_trace(go.Scatter(x=[x0, x1], y=[slope*x0+intercept, slope*x1+intercept],
                                     mode='lines', showlegend=False,
                                     line=dict(color='#E45756', width=2.5, dash='dash')))
        fig.update_layout(title=None, template='plotly_white', height=520, showlegend=False,
                          font=dict(size=16, family=FONT_FAMILY, color='#1f2937'),
                          xaxis_title='日常听说练习得分率', yaxis_title=f'联考平均分（满分{fm:.0f}）',
                          margin=dict(l=80, r=60, t=60, b=70),
                          annotations=[dict(x=0.02, y=0.98, xref='paper', yref='paper',
                                            text=f"相关系数 r = {exam['r_score']}", showarrow=False,
                                            font=dict(size=18, color='#E45756'), align='left')])
        charts['exam_scatter'] = fig


def export_to_docx(report_md: str, charts: dict = None, data: dict = None) -> tuple:
    """导出为公文格式Word

    格式要求：
    - 标题：方正小标宋简体，二号(22pt)，居中
    - 一级标题：黑体，三号(16pt)
    - 二级标题：楷体_GB2312，三号(16pt)
    - 正文：仿宋_GB2312，三号(16pt)，首行缩进2字符
    - 行间距：固定值31磅
    - 页边距：上3.7cm、下3.5cm、左2.8cm、右2.6cm
    - 表格：无边框线，宋体五号(10.5pt)，列宽紧凑
    - 图表：无间隔，居中，宽500px×高250px
    """
    data = data or {}
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None, "python-docx未安装"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # ── 辅助函数 ──────────────────────────────────────────────

    def set_font(run, fname, fsize, bold=False):
        run.font.name = fname
        run.font.size = Pt(fsize)
        run.font.bold = bold
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
        except Exception:
            pass

    def para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=False, space_before=0, space_after=0,
                  line_spacing=31):
        para.alignment = align
        pf = para.paragraph_format
        if first_indent:
            pf.first_line_indent = Pt(32)
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = Pt(line_spacing)

    def add_para(text, fname='仿宋_GB2312', fsize=16, bold=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=False, space_before=0, space_after=0):
        p = doc.add_paragraph()
        para_fmt(p, align, first_indent, space_before, space_after, 31)
        r = p.add_run(text)
        set_font(r, fname, fsize, bold)
        return p

    def remove_table_borders(tbl):
        """去掉表格所有边框线"""
        tblPr = tbl._tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl._tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for btype in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{btype}')
            b.set(qn('w:val'), 'none')
            tblBorders.append(b)
        tblPr.append(tblBorders)

    def add_border_table(headers, rows_data):
        """有边框线表格：黑体五号(10.5pt)，紧凑列宽，防御性列数对齐"""
        # 统一列数：取所有行的最大列数，不足者在末尾补空字符串
        all_rows_raw = [headers] + list(rows_data)
        ncol = max(len(row) for row in all_rows_raw) if all_rows_raw else 1
        def pad_row(row):
            return list(row) + [''] * (ncol - len(row))
        headers_padded = pad_row(headers)
        rows_padded    = [pad_row(r) for r in rows_data]

        tbl = doc.add_table(rows=1+len(rows_padded), cols=ncol)
        tbl.style = 'Table Grid'
        FNAME = '宋体'; FSIZE = 10.5

        def fill_cell(cell, text, center=True, bold=False):
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, FNAME, FSIZE, bold)

        for ci, h in enumerate(headers_padded):
            fill_cell(tbl.rows[0].cells[ci], h, center=True, bold=True)
        for ri, row in enumerate(rows_padded):
            for ci, val in enumerate(row):
                fill_cell(tbl.rows[ri+1].cells[ci], str(val), center=True, bold=False)

        # 紧凑列宽（基于填入后的表格实际内容计算）
        all_content = [headers_padded] + rows_padded
        col_widths = []
        for ci in range(ncol):
            max_len = max(len(str(row[ci])) for row in all_content)
            width_cm = max(1.5, min(max_len * 0.5 + 0.4, 10))
            col_widths.append(Cm(width_cm))
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = col_widths[ci]

        # 居中：整表水平居中（通过段落对齐实现）
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        return tbl

    def add_chart_image(key, title, width=Cm(16.2), height=Cm(9.0), legend_text=None):
        """以无间隔居中图片方式插入图表"""
        if not charts or key not in charts:
            return
        fig = charts[key]
        try:
            img_bytes = fig.to_image(format='png', width=1280, height=720, scale=2)
        except Exception:
            note = doc.add_paragraph()
            para_fmt(note, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_before=0, space_after=6, line_spacing=24)
            note_run = note.add_run(f"【{title}未导出：当前环境缺少图表导出依赖 kaleido】")
            set_font(note_run, '宋体', 10.5, False)
            return
        img_io = BytesIO(img_bytes)
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_before=0, space_after=0, line_spacing=0)
        run = p.add_run()
        run.add_picture(img_io, width=width, height=height)
        # 图注
        cap = doc.add_paragraph()
        para_fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_before=0, space_after=2, line_spacing=31)
        cap_run = cap.add_run(title)
        set_font(cap_run, '宋体', 10.5, False)


    def add_chart_data_table(title, headers, rows):
        """为导出文档补充原生中文数据表，避免静态图中文丢失"""
        if not rows:
            return
        note = doc.add_paragraph()
        para_fmt(note, align=WD_ALIGN_PARAGRAPH.LEFT, first_indent=False, space_before=0, space_after=3, line_spacing=24)
        note_run = note.add_run(f"{title}（文字版）")
        set_font(note_run, '宋体', 10.5, True)
        add_border_table(headers, rows)


    # ── Markdown解析与Word构建 ────────────────────────────────
    lines = report_md.split('\n')
    i = 0
    active_section = None       # 当前节编号如'四、'
    section_had_table = False   # 当前节是否已渲染过表格
    pending_charts = {}         # 当前节待插入图表 {key: caption}
    pending_subsection_charts = []

    CHART_MAP = {
        '三、': {},
        '四、': {},
        '五、': {
            'top_class_trend': '图6  标杆班级月度作业量与得分率组合图',
        },
    }

    SUBSECTION_CHART_MAP = {
        '3.3 应用频次分析': [
            ('monthly_line', '图1  月度作业总量趋势', Cm(16.2), Cm(9.0)),
            ('grade_monthly_line', '图2  各年级月度作业量趋势', Cm(16.2), Cm(9.0)),
        ],
        '3.4 应用方式分析': [
            ('cat_stacked', '图3  各月各类作业量分布', Cm(16.2), Cm(9.0)),
            ('cat_grade_pct', '图4  各年级各栏目训练占比', Cm(16.4), Cm(9.4)),
        ],
        '4.1 成绩数据对比': [
            ('grade_score', '图5  各年级听说模拟得分率趋势', Cm(16.2), Cm(9.0)),
        ],
        '2.1 全区月度应用趋势': [
            ('monthly_line', '图1  全区月度作业量趋势', Cm(16.2), Cm(9.0)),
            ('grade_monthly_line', '图2  各年级月度作业量趋势', Cm(16.2), Cm(9.0)),
        ],
        '3.1 学校应用画像': [
            ('district_school_active', '图3  各校班均有效布置指数', Cm(16.2), Cm(9.0)),
        ],
        '5.1 各校训练结构': [
            ('district_school_category_pct', '图5  各校各栏目训练结构（二维条形图-百分比堆积条形图）', Cm(16.4), Cm(11.0)),
        ],
        '5.2 典型学校年级分析': [
            ('district_case_grade', '图6  典型学校年级听说模拟次数与提分率', Cm(16.2), Cm(9.0)),
        ],
        '5.3 听说模拟题型表现': [
            ('district_qt_grade', '图7  各年级各题型平均得分率展示', Cm(16.2), Cm(11.0)),
        ],
        '6.1 联考整体情况': [
            ('exam_trend', '图8  全区联考整体均分阶段趋势', Cm(15.0), Cm(8.0)),
        ],
        '6.2 各校联考成绩对比': [
            ('exam_school', '图9  各校联考平均分对比', Cm(16.2), Cm(10.5)),
        ],
        '6.3 各题型联考表现': [
            ('exam_qt', '图10  各题型“日常练习 vs 联考”得分率对照', Cm(16.2), Cm(9.0)),
        ],
        '6.4 训练与联考关联': [
            ('exam_scatter', '图11  日常练习得分率与联考平均分关联', Cm(15.0), Cm(9.0)),
        ],
    }

    def flush_section_charts():
        """将当前节所有待插图表插入文档，并标记已处理"""
        global section_had_table
        if pending_charts:
            # 图表与上文（通常是表格）之间空一行
            gap = doc.add_paragraph()
            para_fmt(gap, space_before=0, space_after=0, line_spacing=0)
        for key, caption in list(pending_charts.items()):
            if charts and key in charts:
                add_chart_image(key, caption, width=Cm(13), height=Cm(6.5))
        pending_charts.clear()
        section_had_table = True   # 标记已处理，防止重复插入

    def flush_subsection_charts():
        """将小节图表放在对应分析文字和表格之后，避免图先于解释出现。"""
        nonlocal pending_subsection_charts
        if not pending_subsection_charts:
            return
        gap = doc.add_paragraph()
        para_fmt(gap, space_before=0, space_after=0, line_spacing=0)
        for key, caption, width, height in pending_subsection_charts:
            if charts and key in charts:
                add_chart_image(key, caption, width=width, height=height)
        pending_subsection_charts = []

    while i < len(lines):
        line = lines[i].strip()

        # 跳过注释行和页脚
        if not line or line.startswith('>') or line.startswith('*数据') or line.startswith('*报告') or line.startswith('*生成时间'):
            i += 1; continue

        # ── 主标题 ────────────────────────────────────────────
        if line.startswith('# ') and '成效报告' in line:
            title_text = line.replace('# ', '').replace('**', '').strip()
            add_para(title_text, '方正小标宋简体', 22, True,
                     WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
            i += 1; continue

        # ── 一级标题（## xxx）───────────────────────────────
        if line.startswith('## '):
            section_text = line.replace('## ', '').strip()
            flush_subsection_charts()
            # 先加载新课明的pending，再把上一节的pending flush出去
            # 这样图表出现在前一节末尾 + 新节内容之前
            new_pending = dict(CHART_MAP.get(section_text[:2], {}))
            if active_section and active_section in CHART_MAP and pending_charts:
                flush_section_charts()
            active_section = section_text[:2]
            section_had_table = False
            pending_charts = new_pending
            add_para(section_text, '黑体', 16, True,
                     WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
            i += 1; continue

        # ── 二级标题（### xxx）─────────────────────────────
        if line.startswith('### '):
            sub_text = line.replace('### ', '').strip()
            flush_subsection_charts()
            add_para(sub_text, '楷体_GB2312', 16, True,
                     WD_ALIGN_PARAGRAPH.LEFT, first_indent=False,
                     space_before=6, space_after=3)
            pending_subsection_charts = list(SUBSECTION_CHART_MAP.get(sub_text, []))
            i += 1; continue

        # ── 段落（处理内联加粗）────────────────────────────────
        if line and not line.startswith('|') and not line.startswith('- ') and not line.startswith('```'):
            # 用 finditer 构建段落segments，避免 re.split 产生的空档问题
            segments = []   # [(text, bold), ...]
            last_end = 0
            for m in re.finditer(r'\*\*(.+?)\*\*', line):
                if m.start() > last_end:
                    segments.append((line[last_end:m.start()], False))
                segments.append((m.group(1), True))
                last_end = m.end()
            if last_end < len(line):
                segments.append((line[last_end:], False))

            if any(b for _, b in segments):
                p = doc.add_paragraph()
                para_fmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, space_before=0, space_after=3, line_spacing=31)
                for text, bold in segments:
                    r = p.add_run(text)
                    set_font(r, '仿宋_GB2312', 16, bold)
            else:
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line).strip()
                if clean:
                    add_para(clean, '仿宋_GB2312', 16, False,
                             WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, space_before=0, space_after=3)
            i += 1; continue

        # ── 列表项（处理加粗）────────────────────────────────
        if line.startswith('- '):
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line).lstrip('- ')
            p = doc.add_paragraph(style='List Bullet')
            para_fmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=False, space_before=0, space_after=2, line_spacing=31)
            parts = re.split(r'(\*\*(.+?)\*\*)', clean)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    set_font(r, '仿宋_GB2312', 16, True)
                elif part:
                    r = p.add_run(part)
                    set_font(r, '仿宋_GB2312', 16, False)
            i += 1; continue

        # ── 表格（有边框，宋体五号，紧凑列宽）────────────────
        if line.startswith('|') and '---' not in line:
            rows_data = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                if '---' not in lines[j]:
                    cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c.strip()) for c in lines[j].strip().split('|')[1:-1]]
                    rows_data.append(cells)
                j += 1
            if rows_data:
                add_border_table(rows_data[0], rows_data[1:])
                section_had_table = True
                # 该节所有图表在最后一个表格之后立即插入
                if pending_charts:
                    flush_section_charts()
            i = j; continue

        i += 1

    # 最后一节如有剩余图表，插入末尾
    flush_subsection_charts()
    if active_section and active_section in CHART_MAP and pending_charts:
        flush_section_charts()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, None
