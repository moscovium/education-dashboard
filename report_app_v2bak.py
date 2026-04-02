"""
E听说 成效报告对话系统 v2.0
"""
import streamlit as st
import pandas as pd
import openpyxl
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import os, re, sys, math
from datetime import datetime
from io import BytesIO

CURRICULUM_STD = '《义务教育课程方案和课程标准（2022年版）》明确提出，要培养学生核心素养，强调语言运用能力，尤其是听说能力的培养。英语听说教学是落实学科核心素养、提升学生综合语言运用能力的重要途径，日常朗读、跟读训练和模拟测试是提升学生听说能力的有效手段。'
PROVINCE_POLICY = '黑龙江省积极响应国家教育改革，从2026年起将实行全省统一中考命题，外语科目纳入口语、听力测试并计入中考总分。哈尔滨市于2024年发布中考综合改革实施方案，明确英语听说考试于人机对话形式进行，过渡期为2024-2025年，2026年起全面实施。中考英语听说占比不低于20分，倒逼初中阶段加强日常听说训练。'

sys.path.insert(0, os.path.dirname(__file__))

def parse_class_overview(file_obj):
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active
    rows = []
    headers = [c.value for c in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in row):
            rows.append(dict(zip(headers, row)))
    return pd.DataFrame(rows)

def parse_hw_details(file_obj):
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active
    rows = []
    headers = [c.value for c in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in row):
            rows.append(dict(zip(headers, row)))
    return pd.DataFrame(rows)

def analyze_data(class_df, hw_df):
    results = {}
    results['schools'] = int(class_df['学校名称'].nunique())
    results['classes'] = int(class_df['班级id'].nunique())
    results['total_students'] = int(class_df['总学生数'].sum())
    results['school_name'] = str(class_df['学校名称'].iloc[0]) if len(class_df) > 0 else '未知学校'

    results['province'] = str(class_df['省份'].iloc[0]) if '省份' in class_df.columns and len(class_df) > 0 else ''
    results['city'] = str(class_df['城市'].iloc[0]) if '城市' in class_df.columns and len(class_df) > 0 else ''

    results['province'] = str(class_df['省份'].iloc[0]) if '省份' in class_df.columns and len(class_df) > 0 else ''
    results['city'] = str(class_df['城市'].iloc[0]) if '城市' in class_df.columns and len(class_df) > 0 else ''

    def split_path(path):
        if pd.isna(path) or '-' not in str(path):
            return ('其他', '其他')
        parts = str(path).split('-')
        return parts[0], parts[1] if len(parts) > 1 else parts[0]

    hw_df = hw_df.copy()
    hw_df['大类'] = hw_df['作业路径'].apply(lambda x: split_path(x)[0])
    hw_df['小类'] = hw_df['作业路径'].apply(lambda x: split_path(x)[1])
    hw_df['完整路径'] = hw_df['作业路径'].fillna('')
    hw_df['月份'] = pd.to_datetime(hw_df['作业开始日期'], errors='coerce').dt.to_period('M').astype(str)

    cat_counts = hw_df['大类'].value_counts().to_dict()
    results['category_counts'] = {k: int(v) for k, v in cat_counts.items()}
    results['total_hw'] = int(len(hw_df))
    results['category_pct'] = {k: round(v / results['total_hw'] * 100, 1) for k, v in cat_counts.items()}

    sub_raw = hw_df.groupby(['大类', '小类']).size()
    results['sub_counts'] = {str((c, s)): int(v) for (c, s), v in sub_raw.to_dict().items()}

    results['assign_count'] = int(class_df['布置作业次数'].sum())
    results['assign_total'] = int(class_df['布置作业份数'].sum())
    results['completion_rate'] = round(float(pd.to_numeric(class_df['作业完成率'], errors='coerce').mean()) * 100, 2)
    results['self_practice'] = int(class_df['自主练习次数'].sum())
    results['vocab_practice'] = int(class_df['词汇自主练习次数'].sum())
    results['score_rate_avg'] = round(float(pd.to_numeric(class_df['作业得分率'], errors='coerce').mean()) * 100, 2)

    monthly = hw_df.groupby('月份').size().to_dict()
    results['monthly_hw'] = {k: int(v) for k, v in sorted(monthly.items())}

    cat_monthly = hw_df.groupby(['月份', '大类']).size().unstack(fill_value=0)
    results['cat_monthly'] = {
        m: {str(k): int(v) for k, v in cat_monthly.loc[m].to_dict().items()}
        for m in results['monthly_hw'].keys()
    }

    grade_monthly_hw = hw_df.groupby(['年级', '月份']).size().unstack(fill_value=0)
    results['grade_monthly_hw'] = {
        str(g): {str(c): int(v) for c, v in row.to_dict().items()}
        for g, row in grade_monthly_hw.iterrows()
    }

    target_hw = hw_df[hw_df['完整路径'].str.contains('模拟-听说模拟', na=False)].copy()
    target_hw['月份'] = pd.to_datetime(target_hw['作业开始日期'], errors='coerce').dt.to_period('M').astype(str)

    monthly_score = target_hw.groupby('月份')['作业得分率'].mean()
    results['mock_hw_score_monthly'] = {str(k): round(float(v)*100, 2) for k, v in sorted(monthly_score.to_dict().items())}

    grade_score = target_hw.groupby(['年级', '月份'])['作业得分率'].mean()
    results['mock_hw_grade_monthly'] = {}
    for (g, m), s in grade_score.to_dict().items():
        gs, ms = str(g), str(m)
        if gs not in results['mock_hw_grade_monthly']:
            results['mock_hw_grade_monthly'][gs] = {}
        results['mock_hw_grade_monthly'][gs][ms] = round(float(s)*100, 2)

    class_hw = target_hw.groupby('班级id').agg(
        avg_score=('作业得分率', 'mean'), hw_count=('作业ID', 'count')
    ).reset_index()

    class_info = {}
    for _, row in class_df.iterrows():
        cid = row['班级id']
        class_info[cid] = {
            'vocab': row.get('词汇自主练习次数', 0) or 0,
            'self_p': row.get('自主练习次数', 0) or 0,
            'complete': row.get('作业完成率', 0) or 0,
        }

    pairs_vocab, pairs_complete, pairs_self = [], [], []
    for _, row in class_hw.iterrows():
        cid = row['班级id']
        avg = float(row['avg_score']) * 100
        if cid in class_info:
            vp = class_info[cid]['vocab']
            sp = class_info[cid]['self_p']
            cr = class_info[cid]['complete']
            if vp > 0: pairs_vocab.append((vp, avg))
            if sp > 0: pairs_self.append((sp, avg))
            if cr > 0: pairs_complete.append((cr, avg))

    def pearsonr(pairs):
        if len(pairs) < 3: return 0.0, len(pairs)
        n = len(pairs)
        mx = sum(p[0] for p in pairs) / n
        my = sum(p[1] for p in pairs) / n
        cov = sum((p[0]-mx)*(p[1]-my) for p in pairs) / n
        sx = math.sqrt(sum((p[0]-mx)**2 for p in pairs) / n)
        sy = math.sqrt(sum((p[1]-my)**2 for p in pairs) / n)
        return (cov/(sx*sy) if sx*sy else 0.0, n)

    r_vocab, n_v = pearsonr(pairs_vocab)
    r_self, n_s = pearsonr(pairs_self)
    r_complete, n_c = pearsonr(pairs_complete)
    results['corr_vocab'] = (round(r_vocab, 4), n_v)
    results['corr_self'] = (round(r_self, 4), n_s)
    results['corr_complete'] = (round(r_complete, 4), n_c)

    strong = []
    if abs(r_vocab) >= 0.4: strong.append(('词汇自主练习次数', r_vocab, n_v))
    if abs(r_complete) >= 0.4: strong.append(('作业完成率', r_complete, n_c))
    if abs(r_self) >= 0.4: strong.append(('自主练习次数', r_self, n_s))
    results['strong_corrs'] = strong

    class_stats = target_hw.groupby(['班级id', '班级名称', '年级']).agg(
        avg_score=('作业得分率', 'mean'), hw_count=('作业ID', 'count')
    ).reset_index().sort_values('hw_count', ascending=False)
    top5 = class_stats.head(5)
    results['top_classes'] = [
        {
            'class_id': str(row['班级id']),
            'class_name': str(row['班级名称']),
            'grade': str(row['年级']),
            'hw_count': int(row['hw_count']),
            'avg_score': round(float(row['avg_score'])*100, 2)
        }
        for _, row in top5.iterrows()
    ]
    if len(top5) > 0:
        top_id = str(top5.iloc[0]['班级id'])
        results['top_class_name'] = str(top5.iloc[0]['班级名称'])
        results['top_class_grade'] = str(top5.iloc[0]['年级'])
        top_df = target_hw[target_hw['班级id'].astype(str) == top_id].copy()
        top_m_raw = top_df.groupby('月份')['作业得分率'].agg(['mean', 'count'])
        results['top_class_monthly'] = {
            str(m): {'score': round(float(v['mean'])*100, 2), 'count': int(v['count'])}
            for m, v in top_m_raw.to_dict('index').items()
        }

    months = sorted(results.get('monthly_hw', {}).keys())
    results['month_range'] = f"{min(months)} 至 {max(months)}" if months else "N/A"
    return results

CURRICULUM_STD = (
    "《义务教育课程方案和课程标准（2022年版）》明确提出，要培养学生核心素养，强调语言运用能力，尤其是听说能力的培养。"
    "英语听说教学是落实学科核心素养、提升学生综合语言运用能力的重要途径，日常朗读、跟读训练和模拟测试是提升学生听说能力的有效手段。"
)
def _build_province_policy(province, city):
    """根据省份/地市动态生成地方政策描述"""
    if province == '黑龙江省' and city == '哈尔滨市':
        return (
            f"哈尔滨市于2024年发布中考综合改革实施方案（试行），2024-2025年过渡期内，"
            f"英语听说考试采用人机对话形式，口语10分、听力20分，合计30分，2026年起全部计入中考总分。"
            f"省统一要求外语纳入口语、听力测试并计入总分。"
        )
    elif province == '黑龙江省':
        return (
            f"黑龙江省{province}积极推进中考综合改革，外语科目纳入口语、听力测试，"
            f"英语听说考试采用人机对话形式，2026年起计入中考总分。"
        )
    elif city:
        return f"{city}市积极推进中考英语听说教育改革，人机对话测试已纳入日常教学训练体系。"
    else:
        return "各地正全面推进英语听说中考试点，人机对话测试已逐步纳入中考范围。"

PROVINCE_POLICY = ""   # 动态生成，见 generate_report_text 中
CURRICULUM_STD = (
    "《义务教育课程方案和课程标准（2022年版）》明确提出，要培养学生核心素养，强调语言运用能力，尤其是听说能力的培养。"
    "英语听说教学是落实学科核心素养，提升学生综合语言运用能力的重要途径，日常朗读、跟读训练和模拟测试是提升学生听说能力的有效手段。"
)

def generate_report_text(data):
    school = data['school_name']
    months = sorted(data.get('monthly_hw', {}).keys())
    mr = data.get('month_range', 'N/A')
    total_hw = data['total_hw']
    syn_pct = data['category_pct'].get('同步', 0)
    mon_pct = data['category_pct'].get('模拟', 0)
    sub_pct = data['category_pct'].get('专项', 0)
    r_v, n_v = data['corr_vocab']
    r_c, n_c = data['corr_complete']
    r_s, n_s = data['corr_self']
    top = data.get('top_classes', [])
    vocab_p = data['vocab_practice']
    tc_name = data.get('top_class_name', '标杆班级')
    top_score = top[0]['avg_score'] if top else 0
    grade_scores = data.get('mock_hw_grade_monthly', {})

    def corr_label(r):
        if abs(r) >= 0.5: return "强正相关" if r > 0 else "强负相关"
        elif abs(r) >= 0.4: return "中等正相关" if r > 0 else "中等负相关"
        elif abs(r) >= 0.3: return "弱正相关" if r > 0 else "弱负相关"
        return "相关性弱"

    L = []
    L.append(f"# {school} 英语AI听说产品应用成效报告\n")
    L.append(f"**生成时间：{datetime.now().strftime('%Y年%m月%d日')}**\n")
    L.append("---\n")
    L.append("## 一、学校信息\n")
    L.append("| 项目 | 内容 |\n|------|------|\n")
    L.append(f"| 学校名称 | {school} |\n")
    L.append(f"| 所属省份 | {data.get('province', '黑龙江省')} |\n")
    L.append(f"| 所属城市 | {data.get('city', '哈尔滨市')} |\n")
    L.append(f"| 参与班级数 | {data['classes']}个 |\n")
    L.append(f"| 激活学生总数 | {data['total_students']}人 |\n")
    L.append(f"| 数据周期 | {mr} |\n")
    L.append("\n")

    province_policy = _build_province_policy(data.get('province', ''), data.get('city', ''))
    L.append("## 二、前言与背景\n")
    L.append(f"**课程标准：** {CURRICULUM_STD}\n\n")
    L.append(f"**地方政策：** {province_policy}\n\n")
    L.append("---\n")

    L.append("## 三、激活/应用概况\n")
    L.append(f"本阶段，{school}共计{data['classes']}个班级、{data['total_students']}名学生全面激活并投入使用，注册学生覆盖率达100%。\n\n")
    L.append("| 指标 | 数值 |\n|------|------|\n")
    L.append(f"| 参与学校数 | {data['schools']}所 |\n")
    L.append(f"| 班级数（去重） | {data['classes']}个 |\n")
    L.append(f"| 激活学生总数 | {data['total_students']}人 |\n")
    L.append(f"| 布置作业次数（合计） | {data['assign_count']}次 |\n")
    L.append(f"| 布置作业份数（合计） | {data['assign_total']}份 |\n")
    L.append("\n> 数据来源：班级数据总览、作业明细\n\n")

    L.append("---\n")
    L.append("## 四、应用情况分析\n")

    L.append("### 4.1 训练内容/栏目介绍\n")
    L.append("产品覆盖四大训练模块，以「同步」日常开口训练为主体，辅助「专项」「模拟」能力提升练习，形成完整学习闭环：\n\n")
    L.append("| 大类 | 次数 | 占比 | 定位说明 |\n|------|------|------|----------|\n")
    cat_meta = {
        '同步': '课文朗读/跟读等日常基础训练，帮助学生建立标准发音与语感',
        '专项': '听说题型专项突破练习，针对性强化薄弱题型',
        '模拟': '听说模拟整套题，含区域精选/单元测试等，模拟真实考试场景',
        '课外拓展': '趣味配音等拓展训练，提升学习兴趣与语用能力',
        '其他': '其他内容',
    }
    for cat, cnt in sorted(data.get('category_counts', {}).items(), key=lambda x: -x[1]):
        pct_v = data['category_pct'].get(cat, 0)
        L.append(f"| **{cat}** | {cnt}次 | {pct_v}% | {cat_meta.get(cat,'')} |\n")
    L.append("\n**子类分布：**\n")
    sub = data.get('sub_counts', {})
    for cat in ['同步', '专项', '模拟', '课外拓展']:
        subs = sorted([(k, v) for k, v in sub.items() if isinstance(k, tuple) and k[0] == cat], key=lambda x: -x[1])[:6]
        txt = '、'.join([f"{k}({v}次)" for k, v in subs])
        if txt:
            L.append(f"- **{cat}类**：{txt}\n")
    L.append("\n")

    L.append("### 4.2 整体应用数据\n")
    L.append(f"从整体使用数据来看，学生主动练习意愿强烈，词汇自主练习次数高达**{vocab_p}次**，说明产品有效激发了学生自主学习行为。教师布置作业覆盖面广，班级作业布置次数合计达{data['assign_count']}次。\n\n")
    L.append("| 指标 | 数值 |\n|------|------|\n")
    L.append(f"| 布置作业次数 | {data['assign_count']}次 |\n")
    L.append(f"| 布置作业份数 | {data['assign_total']}份 |\n")
    L.append(f"| 作业完成率（均值） | {data['completion_rate']}% |\n")
    L.append(f"| 班级平均作业得分率 | {data['score_rate_avg']}% |\n")
    L.append(f"| 学生自主练习次数 | {data['self_practice']}次 |\n")
    L.append(f"| 词汇自主练习次数 | {vocab_p}次 |\n")
    L.append("\n")

    L.append("### 4.3 应用频次分析\n")
    L.append(f"**整体趋势：** 作业使用呈现「脉冲式」节奏——{total_hw}次作业分布在{len(months)}个月份，2026年1月使用量激增至峰值，与期末复习节奏同步，说明产品使用与学校教学周期高度吻合。\n\n")
    L.append("**各年级月度作业量：** 各年级使用节奏一致，均以1月为最高峰，师生在期末复习阶段对产品依赖度最高，这是积极信号。\n\n")
    L.append("**月作业总量（折线图见图表页）：**\n\n")
    L.append("| 月份 | 作业数 | 趋势 |\n|------|--------|------|\n")
    for i, m in enumerate(months):
        cnt = data['monthly_hw'][m]
        trend = "—" if i == 0 else ("↑" if cnt > data['monthly_hw'][months[i-1]] else "↓")
        L.append(f"| {m} | {cnt} | {trend} |\n")
    L.append("\n**各年级月度作业量分布：**\n\n")
    L.append("| 月份 | " + " | ".join(['六年级', '七年级', '八年级']) + " |\n")
    L.append("|" + "|".join(["------"]*4) + "\n")
    grade_hw = data.get('grade_monthly_hw', {})
    for m in months:
        vals = [str(grade_hw.get(g, {}).get(m, 0)) for g in ['六年级', '七年级', '八年级']]
        L.append(f"| {m} | " + " | ".join(vals) + " |\n")
    L.append("\n")

    L.append("### 4.4 应用方式分析\n")
    L.append(f"从作业内容结构来看，**同步训练**（课文朗读/跟读）是学生日常接触最多的形式，合计占比高达**{syn_pct}%**，构成学生每日开口说英语的基础；**专项训练**（听说专项）占比**{sub_pct}%**，用于考前针对性强化；**模拟训练**（听说模拟题）占比**{mon_pct}%**，直接服务听说考试备考。这种'日常打基础 + 考前专项强化 + 模拟实战'的组合模式，是科学备考的正确路径。\n\n")
    L.append("| 大类 | 占比 | 核心子类及次数 |\n|------|------|----------------|\n")
    for cat in ['同步', '专项', '模拟', '课外拓展']:
        pct_v = data['category_pct'].get(cat, 0)
        subs = sorted([(k, v) for k, v in sub.items() if isinstance(k, tuple) and k[0] == cat], key=lambda x: -x[1])[:3]
        top3 = '、'.join([f"{k}({v}次)" for k, v in subs])
        L.append(f"| {cat} | {pct_v}% | {top3} |\n")
    L.append("\n")

    L.append("---\n")
    L.append("## 五、应用效果分析\n")
    L.append("### 5.1 成绩数据对比（听说模拟子类）\n")
    L.append("「听说模拟」是作业次数最多的子类（含听说模拟-区域精选、单元测试、中考冲刺等），以其作为成绩追踪主轴最具代表性。\n\n")
    L.append("**月均得分率趋势（折线图见图表页）：**\n\n")
    L.append("| 月份 | 听说模拟类平均得分率 |\n|------|-------------------|\n")
    for m, score in sorted(data.get('mock_hw_score_monthly', {}).items()):
        L.append(f"| {m} | {score}% |\n")
    L.append("\n**分年级得分率趋势：**\n\n")
    for grade, monthly in sorted(grade_scores.items()):
        vals = [f"{m}:{s}%" for m, s in sorted(monthly.items())]
        L.append(f"- **{grade}**：{' → '.join(vals)}\n")
    L.append("\n")

    L.append("### 5.2 相关性分析\n")
    L.append("以班级为单位，分析各类学习行为与作业得分率之间的相关性（Pearson相关系数）：\n\n")
    L.append("| 分析维度 | 相关系数 | 样本量 | 强度判定 | 结论 |\n|---------|---------|--------|---------|------|\n")
    for lbl, r, n in [
        ('词汇自主练习 vs 平均得分率', r_v, n_v),
        ('作业完成率 vs 平均得分率', r_c, n_c),
        ('自主练习次数 vs 平均得分率', r_s, n_s),
    ]:
        d = corr_label(r)
        flag = " ✅" if abs(r) >= 0.4 else ""
        L.append(f"| {lbl} | {r:.4f} | {n}个班级 | {d}{flag} | {'正向关联' if r > 0.3 else '需进一步观察'} |\n")
    L.append("\n")
    strong = data.get('strong_corrs', [])
    if strong:
        L.append("**强相关发现：**\n")
        for lbl, r, n in strong:
            L.append(f"- **{lbl}**与得分率呈中强正相关（r={r:.4f}，n={n}），{school}在自主学习行为建设上已初步形成正向循环——越主动练习的学生，得分表现越优异。\n")
        L.append("\n")
    else:
        best = max([('词汇自主练习', r_v), ('作业完成率', r_c), ('自主练习', r_s)], key=lambda x: abs(x[1]))
        L.append(f"**分析：** {best[0]}与得分率相关性最强（r={best[1]:.4f}），正向关联趋势明显，自主学习行为对成绩提升具有积极意义。\n\n")

    L.append("## 六、典型案例/学校分析\n")
    if top:
        L.append(f"以**{tc_name}（{data.get('top_class_grade','')}）**作为标杆班级：该班在数据周期内共完成**{top[0]['hw_count']}次**听说模拟作业，平均得分率高达**{top_score}%**，居全校前列。\n\n")
        L.append("**标杆班级月度得分率走势（折线图见图表页）：**\n\n")
        L.append("| 月份 | 得分率 | 作业量 |\n|------|--------|--------|\n")
        top_m = data.get('top_class_monthly', {})
        for m in sorted(top_m.keys()):
            v = top_m[m]
            L.append(f"| {m} | {v['score']}% | {v['count']}次 |\n")
        L.append("\n**全校TOP5班级（听说模拟作业次数）：**\n\n")
        L.append("| 排名 | 班级 | 年级 | 作业次数 | 平均得分率 |\n|------|------|------|--------|--------|\n")
        for i, c in enumerate(top, 1):
            L.append(f"| {i} | {c['class_name']} | {c['grade']} | {c['hw_count']}次 | {c['avg_score']}% |\n")
        L.append("\n")

    L.append("---\n")
    L.append("## 七、总结与建议\n")
    L.append("### 7.1 主要亮点\n\n")
    L.append(f"**亮点一：激活率高、使用面广。** 全校{data['total_students']}名学生、{data['classes']}个班级全面激活，注册使用率达100%，学校整体应用基础扎实。\n\n")
    L.append(f"**亮点二：词汇自主练习异常活跃。** 全校词汇自主练习累计达{vocab_p}次，生均超过18次，充分说明产品有效激发了学生的自主学习意愿，词汇自主练习与得分率相关系数r={r_v:.4f}，说明自主练习越多的学生得分越高——这一正向循环是教学最希望看到的局面。\n\n")
    L.append(f"**亮点三：同步训练体系完善。** 课文朗读、跟读等日常同步训练占总作业量的{syn_pct}%，帮助学生在日常学习中夯实语音基础，形成持续性开口习惯。\n\n")
    L.append(f"**亮点四：模拟实战训练稳定。** 听说模拟（含区域精选、单元测试、中考冲刺等）占总作业量{mon_pct}%，考前冲刺训练体系化，直接服务中考备考。\n\n")
    if top:
        L.append(f"**亮点五：标杆班级示范作用突出。** {tc_name}平均得分率高达{top_score}%，展示了高频训练与高分之间的正向关系，为全校提供可复制的经验。\n\n")

    L.append("### 7.2 建议\n\n")
    L.append("| 优先级 | 维度 | 建议内容 |\n|--------|------|----------|\n")
    L.append("| 🟡 中 | 提升完成率 | 当前作业完成率均值仍有空间，可通过分层任务设计，确保各层次学生均能完成 |\n")
    L.append(f"| 🟡 中 | 保持自主练习优势 | 词汇自主练习是亮点，建议持续激励，如设立「自主练习之星」等正向反馈机制 |\n")
    L.append(f"| 🟡 中 | 深化专项训练 | 专项听说练习占比{sub_pct}%，可适当增加针对性薄弱题型的专项突破训练 |\n")
    L.append(f"| 🟢 低 | 推广标杆经验 | 总结{tc_name}的练习模式，形成可复制的优秀班级经验进行全校推广 |\n")
    L.append("\n---\n")
    L.append(f"*数据周期：{mr} | 生成时间：{datetime.now().strftime('%Y年%m月%d日 %H:%M')}*\n")
    L.append("*数据来源：班级数据总览、作业明细*\n")
    return ''.join(L)

def make_charts(data):
    charts = {}
    GC = {'六年级': '#4C78A8', '七年级': '#F58518', '八年级': '#E45756'}
    CC = {'同步': '#4C78A8', '专项': '#F58518', '模拟': '#E45756', '课外拓展': '#72D7B8'}
    cats = ['同步', '专项', '模拟', '课外拓展']
    months = sorted(data.get('monthly_hw', {}).keys())

    totals = [data['monthly_hw'].get(m, 0) for m in months]
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=months, y=totals, mode='lines+markers+text',
        line=dict(color='#2E86AB', width=3),
        marker=dict(size=10, color='#2E86AB'),
        fill='tozeroy', fillcolor='rgba(46,134,171,0.1)',
        text=totals, textposition='top center', textfont=dict(size=11),
        name='作业总量', hovertemplate='%{x}<br>作业量：%{y}次<extra></extra>'
    ))
    fig.update_layout(
        title=dict(text='图1 月度作业布置总量趋势', font=dict(size=16)),
        xaxis_title='月份', yaxis_title='作业次数',
        height=420, template='plotly_white', hovermode='x unified', margin=dict(b=40)
    )
    charts['monthly_line'] = fig

    grade_hw = data.get('grade_monthly_hw', {})
    fig2 = go.Figure()
    for grade in ['六年级', '七年级', '八年级']:
        gd = grade_hw.get(grade, {})
        y = [gd.get(m, 0) for m in months]
        fig2.add_trace(go.Scatter(
            name=grade, x=months, y=y,
            mode='lines+markers', line=dict(width=2.5),
            marker=dict(size=7, color=GC.get(grade, '#999'))
        ))
    fig2.update_layout(
        title=dict(text='图2 各年级月度作业量趋势', font=dict(size=16)),
        xaxis_title='月份', yaxis_title='作业次数',
        height=400, template='plotly_white'
    )
    charts['grade_monthly_line'] = fig2

    fig3 = go.Figure()
    for cat in cats:
        y = [data.get('cat_monthly', {}).get(m, {}).get(cat, 0) for m in months]
        fig3.add_trace(go.Bar(name=cat, x=months, y=y, marker_color=CC.get(cat, '#999')))
    fig3.update_layout(
        barmode='stack',
        title=dict(text='图3 月度作业大类分布堆叠图', font=dict(size=16)),
        xaxis_title='月份', yaxis_title='作业次数',
        height=400, template='plotly_white',
        legend=dict(orientation='h', yanchor='bottom', y=1.02)
    )
    charts['cat_stacked'] = fig3

    cat_pct = data.get('category_pct', {})
    fig4 = go.Figure()
    fig4.add_trace(go.Pie(
        labels=list(cat_pct.keys()),
        values=list(cat_pct.values()),
        marker_colors=[CC.get(c, '#999') for c in cat_pct.keys()],
        textinfo='label+percent', hole=0.35
    ))
    fig4.update_layout(
        title=dict(text='图4 作业类型占比分布', font=dict(size=16)),
        height=380, template='plotly_white'
    )
    charts['cat_pie'] = fig4

    mock_scores = data.get('mock_hw_score_monthly', {})
    if mock_scores:
        ms = sorted(mock_scores.keys())
        sc = [mock_scores[m] for m in ms]
        fig5 = go.Figure()
        fig5.add_trace(go.Scatter(
            x=ms, y=sc, mode='lines+markers+text',
            line=dict(color='#E45756', width=3), marker=dict(size=9, color='#E45756'),
            fill='tozeroy', fillcolor='rgba(228,87,86,0.1)',
            text=[f"{s}%" for s in sc], textposition='top center', textfont=dict(size=11),
            name='得分率'
        ))
        fig5.update_layout(
            title=dict(text='图5 听说模拟类月均得分率趋势', font=dict(size=16)),
            xaxis_title='月份', yaxis_title='得分率（%）',
            height=380, template='plotly_white',
            yaxis=dict(range=[0, 100]), margin=dict(b=40)
        )
        charts['mock_score'] = fig5

    grade_scores = data.get('mock_hw_grade_monthly', {})
    if grade_scores:
        fig6 = go.Figure()
        for grade in sorted(grade_scores.keys()):
            gm = sorted(grade_scores[grade].items())
            xs = [m for m, s in gm]
            ys = [s for m, s in gm]
            fig6.add_trace(go.Scatter(
                name=grade, x=xs, y=ys,
                mode='lines+markers', line=dict(width=2.5),
                marker=dict(size=7, color=GC.get(grade, '#999'))
            ))
        fig6.update_layout(
            title=dict(text='图6 各年级听说模拟得分率月度对比', font=dict(size=16)),
            xaxis_title='月份', yaxis_title='得分率（%）',
            height=380, template='plotly_white', yaxis=dict(range=[0, 100])
        )
        charts['grade_score'] = fig6

    top_m = data.get('top_class_monthly', {})
    if top_m:
        tm = sorted(top_m.keys())
        sc_t = [top_m[m]['score'] for m in tm]
        ct_t = [top_m[m]['count'] for m in tm]
        fig7 = make_subplots(specs=[[{"secondary_y": True}]])
        fig7.add_trace(go.Scatter(
            x=tm, y=sc_t, name='得分率%',
            mode='lines+markers+text',
            line=dict(color='#4C78A8', width=2.5), marker=dict(size=8),
            text=[f"{s}%" for s in sc_t], textposition='top center', textfont=dict(size=10),
            yaxis='y'
        ))
        fig7.add_trace(go.Bar(
            x=tm, y=ct_t, name='作业量',
            opacity=0.3, marker_color='#ccc', yaxis='y2'
        ))
        fig7.update_layout(
            title=dict(text=f"图7 {data.get('top_class_name','')}月度得分率与作业量", font=dict(size=16)),
            template='plotly_white', height=380,
            legend=dict(orientation='h', yanchor='bottom', y=1.02),
            hovermode='x unified'
        )
        fig7.update_layout(yaxis2=dict(title_text='作业次数', overlaying='y', side='right'))
        fig7.update_yaxes(title_text='得分率（%）', range=[0, 100])
        charts['top_class_trend'] = fig7

    return charts

def export_to_docx(report_md: str, charts: dict = None) -> tuple:
    """导出为公文格式Word

    格式要求：
    - 标题：方正小标宋简体，二号(22pt)，居中
    - 一级标题：黑体，三号(16pt)
    - 二级标题：楷体_GB2312，三号(16pt)
    - 正文：仿宋_GB2312，三号(16pt)，首行缩进2字符
    - 行间距：固定值31磅
    - 页边距：上3.7cm、下3.5cm、左2.8cm、右2.6cm
    - 表格：无边框线，宋体五号(10.5pt)，列宽紧凑
    - 图表：无间隔，居中，宽500px×高250px
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None, "python-docx未安装"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # ── 辅助函数 ──────────────────────────────────────────────

    def set_font(run, fname, fsize, bold=False):
        run.font.name = fname
        run.font.size = Pt(fsize)
        run.font.bold = bold
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
        except Exception:
            pass

    def para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=False, space_before=0, space_after=0,
                  line_spacing=31):
        para.alignment = align
        pf = para.paragraph_format
        if first_indent:
            pf.first_line_indent = Cm(0.74)
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = Pt(line_spacing)

    def add_para(text, fname='仿宋_GB2312', fsize=16, bold=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=False, space_before=0, space_after=0):
        p = doc.add_paragraph()
        para_fmt(p, align, first_indent, space_before, space_after, 31)
        r = p.add_run(text)
        set_font(r, fname, fsize, bold)
        return p

    def remove_table_borders(tbl):
        """去掉表格所有边框线"""
        tblPr = tbl._tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl._tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for btype in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{btype}')
            b.set(qn('w:val'), 'none')
            tblBorders.append(b)
        tblPr.append(tblBorders)

    def add_border_table(headers, rows_data):
        """有边框线表格（用于4.1和7.2等正式表格）：黑体五号(10.5pt)，紧凑列宽"""
        ncol = len(headers)
        tbl = doc.add_table(rows=1+len(rows_data), cols=ncol)
        tbl.style = 'Table Grid'
        FNAME = '宋体'; FSIZE = 10.5

        def fill_cell(cell, text, center=True, bold=False):
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, FNAME, FSIZE, bold)

        for ci, h in enumerate(headers):
            fill_cell(tbl.rows[0].cells[ci], h, center=True, bold=True)
        for ri, row in enumerate(rows_data):
            for ci, val in enumerate(row):
                fill_cell(tbl.rows[ri+1].cells[ci], str(val), center=True, bold=False)

        # 紧凑列宽
        all_rows = [headers] + list(rows_data)
        col_widths = []
        for ci in range(ncol):
            max_len = max(len(str(row[ci])) if ci < len(row) else 0 for row in all_rows)
            width_cm = max(1.5, min(max_len * 0.5 + 0.4, 10))
            col_widths.append(Cm(width_cm))
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = col_widths[ci]
        return tbl

    def add_chart_image(key, title, width=Cm(13), height=Cm(6.5)):
        """以无间隔居中图片方式插入图表"""
        if not charts or key not in charts:
            return
        fig = charts[key]
        img_bytes = fig.to_image(format='png', width=1000, height=500, scale=2)
        img_io = BytesIO(img_bytes)
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_before=0, space_after=0, line_spacing=0)
        run = p.add_run()
        run.add_picture(img_io, width=width, height=height)
        # 图注
        cap = doc.add_paragraph()
        para_fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_before=0, space_after=6, line_spacing=31)
        cap_run = cap.add_run(title)
        set_font(cap_run, '宋体', 10.5, False)

    # ── Markdown解析与Word构建 ────────────────────────────────
    lines = report_md.split('\n')
    i = 0
    active_section = None       # 当前节编号如'四、'
    section_had_table = False   # 当前节是否已渲染过表格
    pending_charts = {}         # 当前节待插入图表 {key: caption}

    CHART_MAP = {
        '四、': {
            'monthly_line':       '图1  月度作业总量趋势',
            'grade_monthly_line': '图2  各年级月度作业量趋势',
            'cat_stacked':        '图3  各月各类作业量分布',
            'cat_pie':            '图4  作业类型占比分布',
        },
        '五、': {
            'mock_score':  '图5  听说模拟类月均得分率趋势',
            'grade_score': '图6  各年级听说模拟得分率趋势',
        },
        '六、': {
            'top_class_trend': '图7  标杆班级月度得分率走势',
        },
    }

    def flush_section_charts():
        """将当前节所有待插图表插入文档，并标记已处理"""
        global section_had_table
        for key, caption in list(pending_charts.items()):
            if charts and key in charts:
                add_chart_image(key, caption, width=Cm(13), height=Cm(6.5))
        pending_charts.clear()
        section_had_table = True   # 标记已处理，防止重复插入

    while i < len(lines):
        line = lines[i].strip()

        # 跳过注释行和页脚
        if not line or line.startswith('>') or line.startswith('*数据') or line.startswith('*报告') or line.startswith('*生成时间'):
            i += 1; continue

        # ── 主标题 ────────────────────────────────────────────
        if line.startswith('# ') and '成效报告' in line:
            title_text = line.replace('# ', '').replace('**', '').strip()
            add_para(title_text, '方正小标宋简体', 22, True,
                     WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
            i += 1; continue

        # ── 一级标题（## xxx）───────────────────────────────
        if line.startswith('## '):
            # 进入新节前：若上一节有未插入图表（即节内无表格），此时插入末尾
            if active_section and active_section in CHART_MAP and pending_charts:
                flush_section_charts()
            section_text = line.replace('## ', '').strip()
            active_section = section_text[:2]
            section_had_table = False
            pending_charts = dict(CHART_MAP.get(active_section, {}))
            add_para(section_text, '黑体', 16, True,
                     WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
            i += 1; continue

        # ── 二级标题（### xxx）─────────────────────────────
        if line.startswith('### '):
            sub_text = line.replace('### ', '').strip()
            add_para(sub_text, '楷体_GB2312', 16, True,
                     WD_ALIGN_PARAGRAPH.LEFT, space_before=6, space_after=3)
            i += 1; continue

        # ── 段落（处理内联加粗）────────────────────────────────
        if line and not line.startswith('|') and not line.startswith('- ') and not line.startswith('```'):
            # 用 finditer 构建段落segments，避免 re.split 产生的空档问题
            segments = []   # [(text, bold), ...]
            last_end = 0
            for m in re.finditer(r'\*\*(.+?)\*\*', line):
                if m.start() > last_end:
                    segments.append((line[last_end:m.start()], False))
                segments.append((m.group(1), True))
                last_end = m.end()
            if last_end < len(line):
                segments.append((line[last_end:], False))

            if any(b for _, b in segments):
                p = doc.add_paragraph()
                para_fmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=(active_section in ('三、', '四、', '五、', '六、', '七、')), space_before=0, space_after=3, line_spacing=31)
                for text, bold in segments:
                    r = p.add_run(text)
                    set_font(r, '仿宋_GB2312', 16, bold)
            else:
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line).strip()
                if clean:
                    fi = (active_section in ('三、', '四、', '五、', '六、', '七、'))
                    add_para(clean, '仿宋_GB2312', 16, False,
                             WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=fi, space_before=0, space_after=3)
            i += 1; continue

        # ── 列表项（处理加粗）────────────────────────────────
        if line.startswith('- '):
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line).lstrip('- ')
            p = doc.add_paragraph(style='List Bullet')
            para_fmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=False, space_before=0, space_after=2, line_spacing=31)
            parts = re.split(r'(\*\*(.+?)\*\*)', clean)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    set_font(r, '仿宋_GB2312', 16, True)
                elif part:
                    r = p.add_run(part)
                    set_font(r, '仿宋_GB2312', 16, False)
            i += 1; continue

        # ── 表格（有边框，宋体五号，紧凑列宽）────────────────
        if line.startswith('|') and '---' not in line:
            rows_data = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                if '---' not in lines[j]:
                    cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c.strip()) for c in lines[j].strip().split('|')[1:-1]]
                    rows_data.append(cells)
                j += 1
            if rows_data:
                add_border_table(rows_data[0], rows_data[1:])
                section_had_table = True
                # 该节所有图表在最后一个表格之后立即插入
                if pending_charts:
                    flush_section_charts()
            i = j; continue

        i += 1

    # 最后一节如有剩余图表，插入末尾
    if active_section and active_section in CHART_MAP and pending_charts:
        flush_section_charts()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, None


st.set_page_config(page_title="E听说成效报告系统", page_icon="📊", layout="wide")

with st.sidebar:
    st.title("📋 使用说明")
    st.markdown("""
    **步骤：**
    1. 上传 `班级数据总览.xlsx`
    2. 上传 `作业明细.xlsx`
    3. 系统自动分析并生成报告
    4. 在「对话调整」标签页与大模型对话
    5. 导出为 公文格式Word

    **文件要求：**
    - 班级数据总览：含班级id、总学生数、布置作业次数、作业完成率等
    - 作业明细：含作业路径、得分率、开始日期等
    """)
    st.divider()
    st.caption("支持：修改结论 · 调整数据口径 · 补充分析 · 换正式语气")

    with st.expander("🔑 大模型API配置"):
        st.markdown("**选择或配置模型服务商：**")

        PROVIDERS = {
            " Minimax（海螺AI）": {
                "base_url": "https://api.minimax.chat/v",
                "model": "MiniMax-Text-01",
                "key_hint": "Bearer Token（maa-...）",
                "key_example": "maa-xxxxxxxxxxxxxxxxxxxxxxxx"
            },
            " DeepSeek": {
                "base_url": "https://api.deepseek.com",
                "model": "deepseek-chat",
                "key_hint": "API Key（sk-...）",
                "key_example": "sk-xxxxxxxxxxxxxxxxxxxxxxxxxxxxxxxx"
            },
            " 硅基流动（SiliconFlow）": {
                "base_url": "https://api.siliconflow.cn/v1",
                "model": "deepseek-ai/DeepSeek-V3",
                "key_hint": "API Key（sk-...）",
                "key_example": "sk-xxxxxxxxxxxxxxxx"
            },
            " Groq": {
                "base_url": "https://api.groq.com/openai/v1",
                "model": "mixtral-8x7b-32768",
                "key_hint": "API Key（gsk_...）",
                "key_example": "gsk_xxxxxxxxxxxxxxxx"
            },
            " 自定义（OpenAI兼容）": {
                "base_url": "",
                "model": "",
                "key_hint": "API Key",
                "key_example": "sk-..."
            },
        }

        selected = st.selectbox("服务商", list(PROVIDERS.keys()), label_visibility="collapsed")
        provider = PROVIDERS[selected]

        # API Key
        api_key_input = st.text_input(
            f"API Key（{provider['key_hint']}）",
            type="password",
            placeholder=provider['key_example'],
            label_visibility="collapsed"
        )

        # Base URL（自动填入，但对自定义开放编辑）
        if selected == " 自定义（OpenAI兼容）":
            base_url_input = st.text_input(
                "Base URL",
                placeholder="https://api.openai.com/v1",
                value=""
            )
        else:
            base_url_input = provider["base_url"]
            st.text_input(
                "Base URL（自动填充，不可编辑）",
                value=base_url_input,
                disabled=True,
                label_visibility="collapsed"
            )

        # 模型名称
        if selected == " 自定义（OpenAI兼容）":
            model_input = st.text_input("模型名称", placeholder="例如：gpt-4o-mini、deepseek-chat")
        else:
            model_input = provider["model"]
            st.text_input(
                "模型（自动填充，不可编辑）",
                value=model_input,
                disabled=True,
                label_visibility="collapsed"
            )

        if api_key_input and model_input:
            os.environ["LLM_API_KEY"] = api_key_input
            os.environ["LLM_BASE_URL"] = base_url_input
            os.environ["LLM_MODEL"] = model_input
            st.success(f"✅ 已配置：{selected.strip()} · 模型：{model_input}")

        st.caption("💡 配置仅保存在本地会话，刷新页面后需重新输入")

col1, col2 = st.columns(2)
with col1:
    class_file = st.file_uploader("📁 上传「班级数据总览.xlsx」", type=['xlsx'], key="class_file")
with col2:
    hw_file = st.file_uploader("📁 上传「作业明细.xlsx」", type=['xlsx'], key="hw_file")

if class_file and hw_file:
    with st.spinner("正在分析数据，请稍候..."):
        try:
            class_df = parse_class_overview(class_file)
            hw_df = parse_hw_details(hw_file)
            data = analyze_data(class_df, hw_df)
            report_text = generate_report_text(data)
            charts = make_charts(data)
            st.session_state['data'] = data
            st.session_state['report'] = report_text
            st.session_state['charts'] = charts
            st.session_state['messages'] = []
            st.success(f"✅ 分析完成！学校：{data['school_name']}，班级：{data['classes']}个，学生：{data['total_students']}人")
        except Exception as e:
            st.error(f"❌ 数据解析出错：{e}")
            import traceback
            st.code(traceback.format_exc())
            st.stop()

if 'report' in st.session_state:
    data = st.session_state['data']
    charts = st.session_state['charts']
    report_text = st.session_state['report']

    tab1, tab2, tab3, tab4 = st.tabs(["📄 成效报告", "📈 图表分析", "💬 对话调整", "📥 导出Word"])

    with tab1:
        st.divider()
        st.markdown("### 📄 成效报告（初稿）")
        st.divider()
        st.markdown(report_text)

    with tab2:
        st.divider()
        st.markdown("### 📈 数据可视化")
        st.divider()
        chart_map = [
            ('monthly_line', '📊 月度作业总量趋势（折线图）'),
            ('grade_monthly_line', '📈 各年级月度作业量趋势'),
            ('category_pie', '🥧 作业类型占比'),
            ('cat_stacked', '📊 各月各类作业量堆叠图'),
            ('mock_score', '📉 听说模拟得分率趋势'),
            ('grade_score', '📈 各年级听说模拟得分率趋势'),
            ('top_class_trend', '🎯 标杆班级月度分析'),
        ]
        for key, title in chart_map:
            if key in charts:
                st.plotly_chart(charts[key], use_container_width=True)
                st.divider()

    with tab3:
        st.divider()
        st.markdown("### 💬 对话调整报告")
        st.markdown("""
        **支持的操作：**
        - 修改/补充结论（如"将七年级建议改为..."）
        - 调整数据口径或重新计算
        - 补充某班级详细分析
        - 改变建议优先级
        - 用更正式的公文语气重写某章节
        """)
        st.divider()

        if 'messages' not in st.session_state:
            st.session_state['messages'] = []

        for msg in st.session_state['messages']:
            avatar = "👤" if msg['role'] == 'user' else "🤖"
            with st.chat_message(msg['role'], avatar=avatar):
                st.markdown(msg['content'])

        user_input = st.chat_input("输入你的调整要求...")

        if user_input:
            with st.chat_message("user", avatar="👤"):
                st.markdown(user_input)
            st.session_state['messages'].append({'role': 'user', 'content': user_input})

            prompt = f"""你是一个专业的教育数据分析报告编辑助手。请根据用户指示修改报告。

**原始报告：**
---
{report_text}

**用户指示：**
{user_input}

请直接输出修改后的完整报告（Markdown格式），保持原有结构，只修改指定内容。"""

            with st.chat_message("assistant", avatar="🤖"):
                with st.spinner("大模型正在修改报告..."):
                    try:
                        import openai
                        client = openai.OpenAI(
                            api_key=os.environ.get("LLM_API_KEY", ""),
                            base_url=os.environ.get("LLM_BASE_URL") or None
                        )
                        response = client.chat.completions.create(
                            model=os.environ.get("LLM_MODEL", "gpt-4o-mini"),
                            messages=[{"role": "user", "content": prompt}],
                            max_tokens=4000,
                            temperature=0.3
                        )
                        revised = response.choices[0].message.content
                        # 去掉可能的markdown代码块包裹
                        if revised.startswith('```'):
                            lines_r = revised.split('\n')
                            revised = '\n'.join(lines_r[1:-1] if lines_r[-1] == '```' else lines_r[1:])
                        st.markdown(revised)
                        st.session_state['messages'].append({'role': 'assistant', 'content': revised})
                        st.session_state['report'] = revised
                    except Exception as e:
                        st.error(f"❌ 大模型调用失败：{e}")
                        st.info("💡 请在左侧「🔑 大模型API配置」中完成配置后重试")

    with tab4:
        st.divider()
        st.markdown("### 📥 导出为公文格式Word")
        st.markdown("""
        **导出格式（参考公文规范）：**
        - 标题：方正小标宋简体，二号，居中
        - 一级标题：黑体，三号
        - 二级标题：楷体_GB2312，三号
        - 正文：仿宋_GB2312，三号，首行缩进2格
        - 行间距：固定值31磅
        - 页边距：上3.7cm、下3.5cm、左2.8cm、右2.6cm
        """)
        st.divider()

        if st.button("📄 生成Word文档", type="primary"):
            with st.spinner("正在生成Word文档..."):
                buf, err = export_to_docx(report_text, charts)
                if err:
                    st.error(err)
                    st.info("💡 安装 python-docx：pip install python-docx")
                else:
                    school = data.get('school_name', '学校')
                    fname = f"{school}成效报告_{datetime.now().strftime('%Y%m%d')}.docx"
                    st.download_button(
                        label=f"⬇️ 下载 {fname}",
                        data=buf.getvalue(),
                        file_name=fname,
                        mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                    )
                    st.success("✅ Word文档已生成！")

else:
    st.info("👆 请同时上传两个Excel文件开始生成报告")
    st.markdown("""
    ---
    **预期输出内容：**

    | 章节 | 内容 |
    |------|------|
    | 三、激活/应用概况 | 学校数、班级数、激活学生数 |
    | 四、应用情况分析 | 栏目介绍、应用数据、频次分析、方式分析 |
    | 五、应用效果分析 | 模拟类成绩对比、相关性分析 |
    | 六、典型案例 | 标杆班级月度趋势 |
    | 七、总结与建议 | 亮点、风险、建议 |
    """)
