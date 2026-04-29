"""
E听说 成效报告对话系统 v2.0
"""
import pandas as pd
import openpyxl
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import os, re, sys, math
from datetime import datetime
from io import BytesIO

sys.path.insert(0, os.path.dirname(__file__))

def parse_class_overview(file_obj):
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active
    rows = []
    headers = [c.value for c in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in row):
            rows.append(dict(zip(headers, row)))
    return pd.DataFrame(rows)

def parse_hw_details(file_obj):
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active
    rows = []
    headers = [c.value for c in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in row):
            rows.append(dict(zip(headers, row)))
    return pd.DataFrame(rows)

def parse_question_type(file_obj):
    """解析「听说模拟班级总体情况-题型」Excel"""
    wb = openpyxl.load_workbook(file_obj, data_only=True)
    ws = wb.active
    rows = []
    headers = [c.value for c in ws[1]]
    for row in ws.iter_rows(min_row=2, values_only=True):
        if any(v is not None for v in row):
            rows.append(dict(zip(headers, row)))
    df = pd.DataFrame(rows)
    # 统一列名：得分率转为小数→百分数
    if '得分率' in df.columns:
        df['得分率'] = pd.to_numeric(df['得分率'], errors='coerce').fillna(0)
    # 年级/班级统一为字符串
    for col in ['年级', '班级']:
        if col in df.columns:
            df[col] = df[col].astype(str).str.strip()
    return df

def _split_path(path):
    if pd.isna(path) or '-' not in str(path):
        return ('其他', '其他')
    parts = str(path).split('-')
    return parts[0], parts[1] if len(parts) > 1 else parts[0]

def _format_grade_class(grade, class_name):
    return f"{str(grade)}{str(class_name)}班"

def analyze_data(class_df, hw_df, qt_df=None):
    results = {}
    results['schools']    = int(class_df['学校名称'].nunique())
    results['classes']    = int(class_df['班级id'].nunique())
    results['total_students'] = int(class_df['总学生数'].sum())
    results['school_name'] = str(class_df['学校名称'].iloc[0]) if len(class_df) > 0 else '未知学校'
    results['province']  = str(class_df['省份'].iloc[0]) if '省份' in class_df.columns and len(class_df) > 0 else ''
    results['city']       = str(class_df['城市'].iloc[0]) if '城市' in class_df.columns and len(class_df) > 0 else ''

    hw_df = hw_df.copy()
    hw_df['大类']     = hw_df['作业路径'].apply(lambda x: _split_path(x)[0])
    hw_df['小类']     = hw_df['作业路径'].apply(lambda x: _split_path(x)[1])
    hw_df['完整路径'] = hw_df['作业路径'].fillna('')
    hw_df['月份']     = pd.to_datetime(hw_df['作业开始日期'], errors='coerce').dt.to_period('M').astype(str)

    cat_counts = hw_df['大类'].value_counts().to_dict()
    results['category_counts'] = {k: int(v) for k, v in cat_counts.items()}
    results['total_hw']        = int(len(hw_df))
    results['category_pct']   = {k: round(v / results['total_hw'] * 100, 1) for k, v in cat_counts.items()}

    sub_raw = hw_df.groupby(['大类', '小类']).size()
    results['sub_counts'] = {str((c, s)): int(v) for (c, s), v in sub_raw.to_dict().items()}

    results['assign_count']    = int(class_df['布置作业次数'].sum())
    results['assign_total']    = int(class_df['布置作业份数'].sum())
    results['completion_rate']  = round(float(pd.to_numeric(class_df['作业完成率'], errors='coerce').mean()) * 100, 2)
    results['self_practice']    = int(class_df['自主练习次数'].sum())
    results['vocab_practice']   = int(class_df['词汇自主练习次数'].sum())
    results['score_rate_avg']  = round(float(pd.to_numeric(class_df['作业得分率'], errors='coerce').mean()) * 100, 2)

    monthly = hw_df.groupby('月份').size().to_dict()
    results['monthly_hw'] = {k: int(v) for k, v in sorted(monthly.items())}

    cat_monthly = hw_df.groupby(['月份', '大类']).size().unstack(fill_value=0)
    results['cat_monthly'] = {
        m: {str(k): int(v) for k, v in cat_monthly.loc[m].to_dict().items()}
        for m in results['monthly_hw'].keys()
    }

    actual_grades = sorted(hw_df['年级'].dropna().unique().astype(str).tolist())
    results['actual_grades'] = actual_grades

    grade_monthly_hw_df = hw_df.groupby(['年级', '月份']).size().unstack(fill_value=0)
    results['grade_monthly_hw'] = {
        str(g): {str(c): int(v) for c, v in row.to_dict().items()}
        for g, row in grade_monthly_hw_df.iterrows()
    }

    # 听说模拟
    mock_hw = hw_df[hw_df['完整路径'].str.contains('模拟-', na=False)].copy()
    mock_hw['月份'] = pd.to_datetime(mock_hw['作业开始日期'], errors='coerce').dt.to_period('M').astype(str)

    monthly_score = mock_hw.groupby('月份')['作业得分率'].mean()
    results['mock_hw_score_monthly'] = {str(k): round(float(v)*100, 2) for k, v in sorted(monthly_score.to_dict().items())}
    monthly_count = mock_hw.groupby('月份').size()
    results['mock_hw_count_monthly'] = {str(k): int(v) for k, v in sorted(monthly_count.to_dict().items())}

    grade_score = mock_hw.groupby(['年级', '月份'])['作业得分率'].mean()
    results['mock_hw_grade_monthly'] = {}
    for (g, m), s in grade_score.to_dict().items():
        gs, ms = str(g), str(m)
        if gs not in results['mock_hw_grade_monthly']:
            results['mock_hw_grade_monthly'][gs] = {}
        results['mock_hw_grade_monthly'][gs][ms] = round(float(s)*100, 2)

    # Pearson相关性
    class_hw = mock_hw.groupby('班级id').agg(
        avg_score=('作业得分率', 'mean'), hw_count=('作业ID', 'count')
    ).reset_index()

    class_info = {}
    for _, row in class_df.iterrows():
        cid = row['班级id']
        class_info[cid] = {
            'vocab':   row.get('词汇自主练习次数', 0) or 0,
            'self_p':  row.get('自主练习次数', 0) or 0,
            'complete':row.get('作业完成率', 0) or 0,
        }

    def pearsonr(pairs):
        if len(pairs) < 3: return 0.0, len(pairs)
        n = len(pairs)
        mx = sum(p[0] for p in pairs) / n
        my = sum(p[1] for p in pairs) / n
        cov = sum((p[0]-mx)*(p[1]-my) for p in pairs) / n
        sx = math.sqrt(sum((p[0]-mx)**2 for p in pairs) / n)
        sy = math.sqrt(sum((p[1]-my)**2 for p in pairs) / n)
        return (cov/(sx*sy) if sx*sy else 0.0, n)

    pairs_vocab, pairs_complete, pairs_self = [], [], []
    for _, row in class_hw.iterrows():
        cid = row['班级id']
        avg = float(row['avg_score']) * 100
        if cid in class_info:
            vp = class_info[cid]['vocab']
            sp = class_info[cid]['self_p']
            cr = class_info[cid]['complete']
            if vp > 0: pairs_vocab.append((vp, avg))
            if sp > 0: pairs_self.append((sp, avg))
            if cr > 0: pairs_complete.append((cr, avg))

    r_vocab,    n_v = pearsonr(pairs_vocab)
    r_self,     n_s = pearsonr(pairs_self)
    r_complete, n_c = pearsonr(pairs_complete)
    results['corr_vocab']    = (round(r_vocab,    4), n_v)
    results['corr_self']     = (round(r_self,     4), n_s)
    results['corr_complete'] = (round(r_complete, 4), n_c)

    strong = []
    if abs(r_vocab)    >= 0.4: strong.append(('词汇自主练习次数', r_vocab,    n_v))
    if abs(r_complete) >= 0.4: strong.append(('作业完成率',       r_complete, n_c))
    if abs(r_self)     >= 0.4: strong.append(('自主练习次数',    r_self,     n_s))
    results['strong_corrs'] = strong

    # ── 学生分层数据（A/B/C三层，基于听说模拟得分）────────────
    student分层 = None
    if not mock_hw.empty and '作业得分率' in mock_hw.columns and '学生id' in mock_hw.columns:
        student_score = mock_hw.groupby('学生id')['作业得分率'].mean()
        if len(student_score) >= 3:
            p70 = student_score.quantile(0.70)
            p30 = student_score.quantile(0.30)
            a_s = student_score[student_score >= p70]
            b_s = student_score[(student_score < p70) & (student_score > p30)]
            c_s = student_score[student_score <= p30]
            total_s = len(student_score)
            student分层 = {
                'n_a': len(a_s), 'n_b': len(b_s), 'n_c': len(c_s),
                'pct_a': round(len(a_s)/total_s*100, 1) if total_s>0 else 0,
                'pct_b': round(len(b_s)/total_s*100, 1) if total_s>0 else 0,
                'pct_c': round(len(c_s)/total_s*100, 1) if total_s>0 else 0,
                'avg_a': round(a_s.mean()*100, 2) if len(a_s)>0 else 0,
                'avg_b': round(b_s.mean()*100, 2) if len(b_s)>0 else 0,
                'avg_c': round(c_s.mean()*100, 2) if len(c_s)>0 else 0,
                'gap':   round((a_s.mean()-c_s.mean())*100, 2) if (len(a_s)>0 and len(c_s)>0) else 0,
            }
    results['student分层'] = student分层

    # TOP5 班级（按所有作业次数）
    class_all_hw = hw_df.groupby(['班级id', '班级名称', '年级']).size().reset_index(name='all_hw_count')
    class_all_hw = class_all_hw.sort_values('all_hw_count', ascending=False)
    class_mock    = mock_hw.groupby(['班级id', '班级名称', '年级']).agg(
        avg_score=('作业得分率', 'mean'), mock_count=('作业ID', 'count')
    ).reset_index()

    top5_all = class_all_hw.head(5)
    top5_list = []
    for _, row in top5_all.iterrows():
        cid = str(row['班级id'])
        mock_row = class_mock[class_mock['班级id'].astype(str) == cid]
        avg_s = float(mock_row['avg_score'].values[0]) * 100 if len(mock_row) > 0 else 0
        mc    = int(mock_row['mock_count'].values[0])         if len(mock_row) > 0 else 0
        top5_list.append({
            'class_id':     cid,
            'class_name':   str(row['班级名称']),
            'grade':        str(row['年级']),
            'display_name': _format_grade_class(row['年级'], row['班级名称']),
            'all_hw_count': int(row['all_hw_count']),
            'mock_count':   mc,
            'avg_score':    round(avg_s, 2),
        })
    results['top_classes'] = top5_list

    if top5_list:
        top_cid = top5_list[0]['class_id']
        results['top_class_name']  = top5_list[0]['display_name']
        results['top_class_grade'] = top5_list[0]['grade']

        top_all_m = hw_df[hw_df['班级id'].astype(str) == top_cid].groupby('月份').size()
        results['top_class_all_monthly'] = {str(m): int(v) for m, v in top_all_m.to_dict().items()}

        top_mock_m = mock_hw[mock_hw['班级id'].astype(str) == top_cid].groupby('月份')['作业得分率'].agg(['mean','count'])
        results['top_class_mock_monthly'] = {
            str(m): {'score': round(float(v['mean'])*100,2), 'count': int(v['count'])}
            for m, v in top_mock_m.to_dict('index').items()
        }

    # 各班级布置作业次数（前10排行）
    ca = class_df[['班级id','班级名称','年级','布置作业次数','布置作业份数']].copy()
    ca = ca.sort_values('布置作业次数', ascending=False)
    results['class_assign_top10'] = [
        {'class_name': str(row['班级名称']), 'grade': str(row['年级']),
         'hw_times': int(row['布置作业次数']), 'hw_count': int(row['布置作业份数'])}
        for _, row in ca.head(10).iterrows()
    ]

    # 各年级汇总（布置作业次数、作业份数、平均完成率、平均得分率）
    grade_class_stats = {}
    for grade, grp in class_df.groupby('年级'):
        g = str(grade)
        grade_class_stats[g] = {
            'hw_times':   int(grp['布置作业次数'].sum()),
            'hw_count':    int(grp['布置作业份数'].sum()),
            'completion_rate': round(float(pd.to_numeric(grp['作业完成率'], errors='coerce').mean()) * 100, 2),
            'score_rate':  round(float(pd.to_numeric(grp['作业得分率'], errors='coerce').mean()) * 100, 2),
        }
    results['grade_class_stats'] = grade_class_stats

    months = sorted(results.get('monthly_hw', {}).keys())
    results['month_range'] = f"{min(months)} 至 {max(months)}" if months else "N/A"

    # ── 题型分析（来自「听说模拟班级总体情况」Excel）────────────
    if qt_df is not None and not qt_df.empty and '题型名称' in qt_df.columns:
        qt = qt_df.copy()
        results['has_question_type'] = True

        # ① 全校各题型平均得分率（难度排序）
        qt_school = qt.groupby('题型名称')['得分率'].agg(['mean', 'std', 'count']).round(4)
        results['qt_school'] = {
            name: {'mean': round(v['mean']*100, 2), 'std': round(v['std']*100, 2), 'count': int(v['count'])}
            for name, v in qt_school.to_dict('index').items()
        }

        # ② 班级 × 题型 得分率矩阵
        if '班级' in qt.columns and '年级' in qt.columns:
            ct = qt.groupby(['班级', '年级'])['得分率'].mean().round(4)
            ct_dict = {}
            for (c, g), v in ct.to_dict().items():
                ct_dict[_format_grade_class(g, c)] = round(float(v)*100, 2)
            results['qt_class'] = ct_dict

        # ③ 年级 × 题型 得分率矩阵
        if '年级' in qt.columns:
            gt = qt.groupby(['年级', '题型名称'])['得分率'].mean().round(4)
            gt_dict = {}
            for (grade, qtype), value in gt.to_dict().items():
                g = str(grade)
                if g not in gt_dict:
                    gt_dict[g] = {}
                gt_dict[g][qtype] = round(float(value) * 100, 2)
            results['qt_grade'] = gt_dict

        # ④ 班级-题型 二维矩阵（横向各班对比，纵向各题型）
        if '班级' in qt.columns and '年级' in qt.columns:
            ct_matrix = qt.pivot_table(
                index=['班级', '年级'], columns='题型名称', values='得分率', aggfunc='mean'
            ).round(4)
            results['qt_matrix'] = {
                _format_grade_class(g, c): {q: round(float(ct_matrix.loc[(c, g), q])*100, 2)
                               if q in ct_matrix.columns and (c, g) in ct_matrix.index else None
                               for q in ct_matrix.columns}
                for c, g in ct_matrix.index
            }
            # 离均差（相对各班均值的偏离，正=强项，负=弱项）
            ct_mean = ct_matrix.mean(axis=1)
            ct_dev = (ct_matrix.sub(ct_mean, axis=0)).round(4)
            results['qt_deviation'] = {
                _format_grade_class(g, c): {q: round(float(ct_dev.loc[(c, g), q])*100, 2)
                               if q in ct_dev.columns and (c, g) in ct_dev.index else None
                               for q in ct_dev.columns}
                for c, g in ct_dev.index
            }

        # ⑤ 各班薄弱题型（按班内离均差识别，并为正文展示预先筛选重点项）
        if '班级' in qt.columns and '年级' in qt.columns:
            ct_mean2 = ct_matrix.mean(axis=1)
            ct_dev2 = ct_matrix.sub(ct_mean2, axis=0)
            weak_qt = ct_dev2.idxmin(axis=1)
            weak_score = ct_dev2.min(axis=1)
            results['qt_weak'] = {
                _format_grade_class(g, c): {'题型': weak_qt.loc[(c, g)], '离均差': round(float(weak_score.loc[(c, g)])*100, 2)}
                for c, g in weak_qt.index
            }
            # 正文仅展示差异最突出的重点班级，且总量不超过5个
            ranked_weak = []
            for c, g in weak_qt.index:
                diff = round(float(weak_score.loc[(c, g)]) * 100, 2)
                if diff < 0:
                    ranked_weak.append({
                        '班级': _format_grade_class(g, c),
                        '题型': weak_qt.loc[(c, g)],
                        '离均差': diff,
                        'abs_diff': abs(diff),
                    })
            ranked_weak = sorted(ranked_weak, key=lambda x: x['abs_diff'], reverse=True)
            results['qt_weak_top'] = ranked_weak[:5]

        # ⑥ 教师 × 题型 得分率（教师效能分析）
        if '教师' in qt.columns and '题型名称' in qt.columns:
            tc_qt = qt.groupby(['教师', '题型名称'])['得分率'].mean().unstack(fill_value=0).round(4)
            results['qt_teacher'] = {
                str(t): {q: round(float(v)*100, 2) for q, v in row.items() if v > 0}
                for t, row in tc_qt.iterrows()
            }
            # 教师综合得分率
            tc_mean = qt.groupby('教师')['得分率'].mean().round(4).sort_values(ascending=False)
            results['qt_teacher_rank'] = {
                str(t): round(float(v)*100, 2) for t, v in tc_mean.to_dict().items()
            }

        # ⑦ 高分率/低分率（题型两极分化分析）
        if '优秀率' in qt.columns and '低分率' in qt.columns:
            qt_hr = qt.groupby('题型名称')[['优秀率', '低分率']].mean().round(4)
            results['qt_hr_lr'] = {
                q: {'优秀率': round(float(r['优秀率'])*100, 2),
                    '低分率': round(float(r['低分率'])*100, 2)}
                for q, r in qt_hr.to_dict('index').items()
            }

    else:
        results['has_question_type'] = False

    return results

def _build_province_policy(province, city):
    if province == '黑龙江省' and city == '哈尔滨市':
        return (
            f"哈尔滨市于2024年发布中考综合改革实施方案（试行），2024-2025年过渡期内，"
            f"英语听说考试采用人机对话形式，口语10分、听力20分，合计30分，2026年起全部计入中考总分。"
            f"省统一要求外语纳入口语、听力测试并计入总分。"
        )
    elif province == '黑龙江省':
        return (
            f"黑龙江省积极推进中考综合改革，外语科目纳入口语、听力测试，"
            f"英语听说考试采用人机对话形式，2026年起计入中考总分。"
        )
    elif city:
        return f"{city}市积极推进中考英语听说教育改革，人机对话测试已纳入日常教学训练体系。"
    else:
        return "各地正全面推进英语听说中考试点，人机对话测试已逐步纳入中考范围。"

CURRICULUM_STD = (
    "《义务教育课程方案和课程标准（2022年版）》明确提出，要培养学生核心素养，强调语言运用能力，尤其是听说能力的培养。"
    "英语听说教学是落实学科核心素养、提升学生综合语言运用能力的重要途径，日常朗读、跟读训练和模拟测试是提升学生听说能力的有效手段。"
)

def generate_report_text(data):
    school        = data['school_name']
    months        = sorted(data.get('monthly_hw', {}).keys())
    mr            = data.get('month_range', 'N/A')
    total_hw     = data['total_hw']
    syn_pct       = data['category_pct'].get('同步', 0)
    mon_pct       = data['category_pct'].get('模拟', 0)
    sub_pct       = data['category_pct'].get('专项', 0)
    r_v, n_v      = data['corr_vocab']
    r_c, n_c      = data['corr_complete']
    r_s, n_s      = data['corr_self']
    strong        = data.get('strong_corrs', [])
    top           = data.get('top_classes', [])
    vocab_p       = data['vocab_practice']
    per_student  = round(vocab_p / max(data["total_students"], 1), 1)
    tc_name       = data.get('top_class_name', '标杆班级')
    tc_grade      = data.get('top_class_grade', '')
    actual_grades = data.get('actual_grades', ['六年级', '七年级', '八年级'])

    # ── 辅助函数：构建相关性判定标签 ───────────────────────
    def corr_label(r):
        if   abs(r) >= 0.5: return "强正相关" if r > 0 else "强负相关"
        elif abs(r) >= 0.4: return "中等正相关" if r > 0 else "中等负相关"
        elif abs(r) >= 0.3: return "弱正相关" if r > 0 else "弱负相关"
        return "相关性弱"

    # ── 数据叙事辅助 ───────────────────────────────────────

    

    # ═══════════════════════════════════════════════════════════
    # 报告文字生成（generate_report_text）
    # 参考广州示范校报告写作规范：
    #   - 总领句先行（"在XXX方面…"）
    #   - 数据前有定性描述，数据后有趋势小结
    #   - 禁止无连接词的跳跃式表述
    #   - 对比参照：环比/同比方向须明确
    #   - 数字叙事：具体数字 > 笼统描述
    # ═══════════════════════════════════════════════════════════
    L = []
    L.append(f"# {school} 英语AI听说产品应用成效报告\n")
    L.append(f"**报告生成时间：{datetime.now().strftime('%Y年%m月%d日')}**\n")
    L.append("\n")

    # ── 一、学校信息 ───────────────────────────────────────
    L.append("## 一、学校信息\n")
    L.append(f"　　{school}积极推进教育数字化转型，在{data.get('province', '黑龙江省')}全面推进英语听说教学改革的背景下，学校引入E听说AI听说教学系统，依托大数据与人工智能技术赋能英语听说教学变革。本学期{data['classes']}个班级、{data['total_students']}名学生全面激活并投入使用，系统应用已深度融入日常教学，数据覆盖周期为{mr}。\n\n")
    L.append("| 项目 | 内容 |\n|------|------|\n")
    L.append(f"| 学校名称 | {school} |\n")
    L.append(f"| 所属省份 | {data.get('province', '黑龙江省')} |\n")
    L.append(f"| 所属城市 | {data.get('city', '哈尔滨市')} |\n")
    L.append(f"| 参与班级数 | {data['classes']}个 |\n")
    L.append(f"| 激活学生总数 | {data['total_students']}人 |\n")
    L.append(f"| 数据周期 | {mr} |\n")
    L.append("\n")

    # ── 二、激活/应用概况 ──────────────────────────────────
    L.append("## 二、激活/应用概况\n")
    L.append(f"　　{school}在{mr}期间，E听说产品应用覆盖{data['classes']}个班级、{data['total_students']}名学生，全面激活率达100%，形成稳定的常态化应用节奏，教师持续通过系统布置听说作业，为教学减负增效奠定坚实基础。\n\n")
    L.append("**核心应用数据如下：**\n\n")
    L.append("| 指标 | 数值 |\n|------|------|\n")
    L.append(f"| 参与学校数 | {data['schools']}所 |\n")
    L.append(f"| 班级数（去重） | {data['classes']}个 |\n")
    L.append(f"| 激活学生总数 | {data['total_students']}人 |\n")
    L.append(f"| 布置作业次数（合计） | {data['assign_count']}次 |\n")
    L.append(f"| 布置作业份数（合计） | {data['assign_total']}份 |\n")
    L.append(f"| 作业完成率（均值） | {data['completion_rate']}% |\n")
    L.append(f"| 班级平均作业得分率 | {data['score_rate_avg']}% |\n")
    L.append("\n")

    grade_stats = data.get('grade_class_stats', {})
    if grade_stats:
        L.append(f"**各年级作业布置与完成情况对比：**\n\n")
        L.append("| 年级 | 布置作业次数 | 布置作业份数 | 平均完成率 | 平均得分率 |\n")
        L.append("|------|------------|------------|----------|----------|\n")
        for grade in sorted(grade_stats.keys()):
            g = grade_stats[grade]
            L.append(f"| {grade} | {g['hw_times']}次 | {g['hw_count']}份 | {g['completion_rate']}% | {g['score_rate']}% |\n")
        L.append("\n")
        # 数据小结：找完成率最高和得分率最高的年级
        best_completion_grade = max(grade_stats, key=lambda g: grade_stats[g]['completion_rate'])
        best_score_grade = max(grade_stats, key=lambda g: grade_stats[g]['score_rate'])
        L.append(f"　　从各年级横向对比来看，{best_completion_grade}平均完成率最高（{grade_stats[best_completion_grade]['completion_rate']}%），{best_score_grade}平均得分率领先（{grade_stats[best_score_grade]['score_rate']}%），反映出不同年级在应用侧重上存在一定差异。\n\n")

    L.append("> 数据来源：班级数据总览、作业明细\n\n")

    # ── 三、应用情况分析 ────────────────────────────────────
    L.append("## 三、应用情况分析\n")

    L.append("### 3.1 训练内容/栏目介绍\n")
    L.append(f"　　系统应用覆盖四类训练场景，以「同步」日常开口训练为主体，辅以「专项」「模拟」等训练形式，形成较为完整的听说教学支持体系。其中，「同步」训练侧重夯实发音基础与语感积累，「专项」训练侧重提升重点题型能力，「模拟」训练侧重服务阶段性检测与考前演练。\n\n")
    L.append("| 大类 | 次数 | 占比 | 定位说明 |\n|------|------|------|----------|\n")
    cat_meta = {
        '同步':      '课文朗读/跟读等日常基础训练，帮助学生建立标准发音与语感',
        '专项':      '听说题型专项突破练习，针对性强化薄弱题型',
        '模拟':      '听说模拟整套题，含区域精选/单元测试等，模拟真实考试场景',
        '课外拓展':  '趣味配音等拓展训练，提升学习兴趣与语用能力',
        '其他':      '其他内容',
    }
    for cat, cnt in sorted(data.get('category_counts', {}).items(), key=lambda x: -x[1]):
        pct_v = data['category_pct'].get(cat, 0)
        L.append(f"| **{cat}** | {cnt}次 | {pct_v}% | {cat_meta.get(cat,'')} |\n")
    L.append("\n")

    L.append("### 3.2 整体应用数据\n")
    L.append(f"　　在作业应用方面，{data['classes']}个班级教师本周期内合计布置作业**{data['assign_count']}次**（{data['assign_total']}份），作业完成率均值为**{data['completion_rate']}%**，班级平均作业得分率为**{data['score_rate_avg']}%**。具体数据如下：\n\n")
    L.append("| 指标 | 数值 |\n|------|------|\n")
    L.append(f"| 布置作业次数 | {data['assign_count']}次 |\n")
    L.append(f"| 布置作业份数 | {data['assign_total']}份 |\n")
    L.append(f"| 作业完成率（均值） | {data['completion_rate']}% |\n")
    L.append(f"| 班级平均作业得分率 | {data['score_rate_avg']}% |\n")
    L.append(f"| 学生自主练习次数 | {data['self_practice']}次 |\n")
    L.append(f"| 词汇自主练习次数 | {vocab_p}次 |\n")
    L.append("\n")
    L.append(f"　　与此同时，学生自主练习保持一定活跃度，词汇自主练习次数累计达**{vocab_p}次**，生均约{per_student}次，表明学生在课堂训练之外已初步形成较为稳定的自主巩固习惯。\n\n")

    # 动态找峰值月及增长趋势
    peak_m = max(months, key=lambda m: data['monthly_hw'].get(m, 0)) if months else months[0] if months else ''
    peak_cnt = data['monthly_hw'].get(peak_m, 0)
    L.append("### 3.3 应用频次分析\n")
    L.append(f"　　在应用频次方面，{total_hw}次作业分布于{len(months)}个月份，整体呈现常态化推进态势。{months[0] if months else ''}至{months[-1] if months else ''}期间，{peak_m}月作业量最高（{peak_cnt}次），与阶段复习安排基本一致，说明系统应用与学校教学节奏保持较好匹配。\n\n")
    grade_hw = data.get('grade_monthly_hw', {})
    if grade_hw:
        peak_month = max(months, key=lambda m: data['monthly_hw'][m])
        peak_grade = max(actual_grades, key=lambda g: grade_hw.get(g, {}).get(peak_month, 0))
        L.append(f"　　从各年级横向对比来看，{peak_month}月作业量最高，{peak_grade}在当月作业量最大，表明该年级在当月应用节奏中最为活跃。\n\n")

    L.append("### 3.4 应用方式分析\n")
    L.append(f"　　从作业内容结构来看，「同步」训练（课文朗读、跟读等）是学生日常接触最多的形式，合计占比高达**{syn_pct}%**，构成学生每日开口说英语的基础；「专项」训练占比**{sub_pct}%**，主要用于针对性强化重点题型；「模拟」训练占比**{mon_pct}%**，直接服务阶段检测与听说考试备考。整体呈现“日常打基础 + 专项补短板 + 模拟促实战”的组合模式，符合循序渐进的教学与备考规律。\n\n")
    L.append("| 大类 | 占比 | 定位说明 |\n|------|------|----------|\n")
    for cat in ['同步', '专项', '模拟', '课外拓展']:
        pct_v = data['category_pct'].get(cat, 0)
        L.append(f"| {cat} | {pct_v}% | {cat_meta.get(cat,'')} |\n")
    L.append("\n")

    # 四、应用效果分析
    L.append("## 四、应用效果分析\n")

    L.append("### 4.1 成绩数据对比\n")
    grade_scores = data.get('mock_hw_grade_monthly', {})
    # 找最优回升案例（从最低点到最后一个月的最大涨幅）
    best_recovery = None
    best_gain = 0
    for grade, monthly in grade_scores.items():
        sm = sorted(monthly.items())
        if len(sm) >= 2:
            vals = [s for _, s in sm]
            low_idx = vals.index(min(vals))
            if low_idx < len(vals) - 1 and vals[-1] > vals[low_idx]:
                gain = vals[-1] - vals[low_idx]
                if gain > best_gain:
                    best_gain = gain
                    best_recovery = (grade, sm[low_idx][0], sm[low_idx][1], sm[-1][0], sm[-1][1])

    best_trend_text = None
    if best_recovery:
        g, lm, ls, ltm, lts = best_recovery
        gain = round(lts - ls, 2)
        best_trend_text = f"　　**{g}**听说模拟得分率从最低点{lm}的**{ls}%**逐步回升至{ltm}的**{lts}%**，整体提升**{gain}个百分点**"
    elif grade_scores:
        best_g = max(grade_scores.keys(), key=lambda g: len(grade_scores[g]))
        sm = sorted(grade_scores[best_g].items())
        best_trend_text = f"　　**{best_g}**听说模拟月均得分率走势为：{' → '.join([f'{m}{s}%' for m,s in sm])}"
    if best_trend_text:
        L.append(f"{best_trend_text}，具体数据如下：\n\n")

    L.append("| 月份 | 听说模拟类平均得分率 |\n|------|-------------------|\n")
    for m, score in sorted(data.get('mock_hw_score_monthly', {}).items()):
        L.append(f"| {m} | {score}% |\n")
    L.append("\n")
    L.append("**各年级听说模拟得分率趋势：**\n\n")
    # 交叉表：行=月份，列=年级
    all_grades_sorted = sorted(grade_scores.keys())
    L.append("| 月份 | " + " | ".join(all_grades_sorted) + " |\n")
    L.append("|" + "|".join(["------"] * (len(all_grades_sorted)+1)) + "\n")
    for m in months:
        vals = [str(grade_scores.get(g, {}).get(m, '—')) + '%' for g in all_grades_sorted]
        L.append(f"| {m} | " + " | ".join(vals) + " |\n")
    L.append("\n")

    # ── 4.2 题型得分分析（可选，依赖「听说模拟班级总体情况」Excel）──
    if data.get('has_question_type'):
        L.append("### 4.2 题型得分分析\n")
        qt_school = data.get('qt_school', {})
        if qt_school:
            # 题型难度排序
            sorted_qt = sorted(qt_school.items(), key=lambda x: x[1]['mean'])
            hardest_qt  = sorted_qt[0][0]
            mid_qt      = sorted_qt[1][0] if len(sorted_qt) > 1 else ''
            easiest_qt  = sorted_qt[-1][0]
            L.append(f"基于本周期听说模拟数据，共分析**{list(qt_school.values())[0]['count']}条**记录，涵盖情景反应、对话或短文朗读、篇章复述三大题型。全校各题型平均得分率如下：\n\n")
            L.append("| 题型 | 平均得分率 | 难度定位 |\n|------|-----------|----------|\n")
            qt_labels = {
                '对话或短文朗读': '基础题型，整体表现较好',
                '情景反应':       '中等难度题型，表现存在一定差异',
                '篇章复述':       '综合性较强题型，对能力要求较高',
            }
            for qt_name, info in sorted_qt:
                label = qt_labels.get(qt_name, '')
                L.append(f"| {qt_name} | {info['mean']}% | {label} |\n")
            L.append("\n")
            L.append(f"　　从全校横向对比情况看，**{easiest_qt}**得分率最高（{qt_school[easiest_qt]['mean']}%），反映出学生在该题型上的整体掌握情况相对较好；**{hardest_qt}**得分率最低（{qt_school[hardest_qt]['mean']}%），低分率为**{data.get('qt_hr_lr', {}).get(hardest_qt, {}).get('低分率', 'N/A')}%**，应作为下一阶段专项教学与复习提升的重点关注题型。\n\n")

        # 各班薄弱题型（仅展示差异最突出的重点项，总量不超过5个）
        qt_weak_top = data.get('qt_weak_top', [])
        if qt_weak_top:
            L.append("**各班薄弱题型诊断：**\n\n")
            weak_list = []
            for item in qt_weak_top:
                weak_list.append(f"- **{item['班级']}**：**{item['题型']}**相对班级平均水平偏弱，离均差为{item['离均差']}个百分点")

            if weak_list:
                L.append("\n".join(weak_list) + "\n")
                L.append("　　以上诊断聚焦差异最突出的重点班级与题型，便于学校优先开展针对性训练与教学改进。\n\n")

        # 教师效能
        qt_teacher = data.get('qt_teacher_rank', {})
        if qt_teacher:
            L.append("**教师效能排名（综合各题型平均得分率）：**\n\n")
            L.append("| 排名 | 教师 | 综合得分率 | 薄弱题型提示 |\n")
            L.append("|------|------|-----------|------------|\n")
            # 找教师对应的薄弱题型
            qt_t_full = data.get('qt_teacher', {})
            sorted_teachers = sorted(qt_teacher.items(), key=lambda x: -x[1])
            for rank, (t, score) in enumerate(sorted_teachers[:8], 1):
                t_scores = qt_t_full.get(t, {})
                weak_t = min(t_scores.items(), key=lambda x: x[1]) if t_scores else ('', 0)
                weak_str = f"{weak_t[0]}({weak_t[1]}%)" if weak_t[0] else '—'
                L.append(f"| {rank} | {t} | {score}% | {weak_str} |\n")
            L.append("\n")
            best_t = sorted_teachers[0][0] if sorted_teachers else ''
            worst_t = sorted_teachers[-1][0] if sorted_teachers else ''
            if best_t and worst_t and best_t != worst_t:
                best_score = sorted_teachers[0][1]
                worst_score = sorted_teachers[-1][1]
                diff = best_score - worst_score
                L.append(f"　　教师所带班级综合得分率存在**{diff:.1f}个百分点**的差异（{best_t}所带班级最高{best_score}%，{worst_t}所带班级相对较低，为{worst_score}%），建议通过校本教研、课堂观摩与经验交流等方式，促进有效做法在校内共享。\n\n")

    if strong:
        L.append("### 4.3 相关性分析\n")
        L.append("以班级为单位，分析各类学习行为与作业得分率之间的相关性（Pearson相关系数）：\n\n")
        L.append("| 分析维度 | 相关系数 | 样本量 | 强度判定 | 结论 |\n")
        L.append("|---------|---------|--------|---------|------|\n")
        for lbl, r, n in [
            ('词汇自主练习次数 vs 平均得分率', r_v, n_v),
            ('作业完成率 vs 平均得分率',        r_c, n_c),
            ('自主练习次数 vs 平均得分率',      r_s, n_s),
        ]:
            d    = corr_label(r)
            flag = " ✅" if abs(r) >= 0.4 else ""
            L.append(f"| {lbl} | {r:.4f} | {n}个班级 | {d}{flag} | {'正向关联' if r > 0.3 else '需进一步观察'} |\n")
        L.append("\n")
        L.append("**Pearson相关系数理论说明：**\n\n")
        L.append("| 系数范围 | 相关强度 | 统计含义 |\n")
        L.append("|---------|---------|---------|\n")
        L.append("| 0.7 ≤ |r| ≤ 1.0 | 强相关 | 两变量存在明显线性关系 |\n")
        L.append("| 0.4 ≤ |r| < 0.7 | 中等相关 | 两变量存在一定线性关系 |\n")
        L.append("| 0.2 ≤ |r| < 0.4 | 弱相关 | 两变量存在微弱线性关系 |\n")
        L.append("| 0 ≤ |r| < 0.2 | 几乎无相关 | 两变量线性关系较弱 |\n")
        L.append("\n")
        L.append("**强相关发现：**\n")
        for lbl, r, n in strong:
            L.append(f"- **{lbl}**与得分率呈中等及以上正相关（r={r:.4f}，n={n}），表明相关学习行为与学业表现之间具有较为明确的同向关系，可为后续教学管理与过程评价提供参考。\n")
        L.append("\n")

    # 五、典型案例
    L.append("## 五、典型班级分析\n")
    if top:
        top0 = top[0]
        L.append(f"　　以**{tc_name}**作为标杆班级（数据周期内作业总量位居全校前列）：\n\n")
        L.append(f"- 该班共完成**{top0['all_hw_count']}次**作业（所有类目），其中听说模拟**{top0['mock_count']}次**\n")
        L.append(f"- 听说模拟平均得分率高达**{top0['avg_score']}%**，居全校前列\n\n")

        top_all_m = data.get('top_class_all_monthly', {})
        top_mock_m = data.get('top_class_mock_monthly', {})

        if top_all_m and top_mock_m:
            # 合并为一个表：三列
            L.append("| 月份 | 所有类目布置作业次数 | 听说模拟类目布置作业次数 | 得分率 |\n")
            L.append("|------|------------------|----------------------|--------|\n")
            all_months = sorted(set(top_all_m.keys()) | set(top_mock_m.keys()))
            for m in all_months:
                all_cnt = top_all_m.get(m, '—')
                mdata = top_mock_m.get(m, None)
                mock_cnt = mdata['count'] if mdata else '—'
                score    = f"{mdata['score']}%" if mdata else '—'
                L.append(f"| {m} | {all_cnt} | {mock_cnt} | {score} |\n")
            L.append("\n")

        L.append("**全校TOP5班级（按总作业量排名）：**\n\n")
        L.append("| 排名 | 班级 | 总作业次数 | 听说模拟次数 | 平均得分率 |\n")
        L.append("|------|------|----------|------------|--------|\n")
        for i, c in enumerate(top, 1):
            L.append(f"| {i} | {c['display_name']} | {c['all_hw_count']}次 | {c['mock_count']}次 | {c['avg_score']}% |\n")
        L.append("\n")

    # 六、总结与建议
    L.append("## 六、总结与建议\n")
    top_score = top[0]["avg_score"] if top else 0

    # ── 问题诊断（基于实际数据发现）────────────────────────
    issues = []   # 存储具体问题描述，供亮点和问题两节共用

    # 问题1：类别失衡
    if syn_pct >= 75:
        issues.append(f"同步训练占比偏高（{syn_pct}%），专项训练与模拟训练占比相对不足，训练结构仍有优化空间")
    if mon_pct < 10 and sub_pct < 10:
        issues.append(f"专项训练（{sub_pct}%）与模拟训练（{mon_pct}%）总体占比较低，阶段性巩固与实战演练力度有待加强")
    elif mon_pct < 10:
        issues.append(f"模拟训练占比仅{mon_pct}%，阶段性检验与考前演练支撑相对不足，训练链条仍需进一步完善")

    # 问题2：班级间不均衡
    if grade_stats:
        rates = [g['completion_rate'] for g in grade_stats.values()]
        if rates and max(rates) - min(rates) > 20:
            worst_g = min(grade_stats, key=lambda g: grade_stats[g]['completion_rate'])
            best_g  = max(grade_stats, key=lambda g: grade_stats[g]['completion_rate'])
            issues.append(f"年级间作业完成率存在较明显差异：{best_g}最高为{grade_stats[best_g]['completion_rate']}%，{worst_g}为{grade_stats[worst_g]['completion_rate']}%，相差{max(rates)-min(rates):.0f}个百分点，分层推进的均衡性仍需加强")

    # 问题3：自主练习薄弱
    if vocab_p < data['total_students'] * 2:
        issues.append(f"词汇自主练习总量为{vocab_p}次、生均{per_student}次，自主巩固训练频次仍有提升空间")
    if abs(r_v) < 0.2 and vocab_p > 0:
        issues.append(f"词汇自主练习与得分相关性较弱（r={r_v:.4f}），自主练习与学业表现之间的转化效应仍需进一步观察")

    # 问题4：整体完成率偏低
    if data['completion_rate'] < 70:
        issues.append(f"作业平均完成率为{data['completion_rate']}%，与较高质量落实要求相比仍有差距，学生持续性练习习惯仍需进一步培育")

    # ── 亮点提炼（从数据中找最大优势）───────────────────────
    highlights = []

    # 亮点1：覆盖率
    highlights.append(f"应用覆盖面较广：{data['classes']}个班级、{data['total_students']}名学生已纳入系统应用范围，作业完成率均值为{data['completion_rate']}%，常态化使用机制已基本建立")

    # 亮点2：最强相关指标
    corr_list = [('词汇自主练习', r_v, n_v), ('作业完成率', r_c, n_c), ('自主练习次数', r_s, n_s)]
    best_corr = max(corr_list, key=lambda x: abs(x[1]))
    bc_name, bc_r, bc_n = best_corr
    if abs(bc_r) >= 0.4:
        highlights.append(f"关键过程指标与结果指标关联较为明确：{bc_name}与得分呈中等正相关（r={bc_r:.4f}，n={bc_n}个班级），说明相关学习行为对学习成效具有一定支撑作用")
    elif abs(bc_r) >= 0.2:
        highlights.append(f"自主练习积累呈现一定积极作用：{bc_name}与得分呈弱正相关（r={bc_r:.4f}），持续练习对能力提升的长期价值值得进一步跟踪")

    # 亮点3：训练体系
    if syn_pct + mon_pct + sub_pct >= 85:
        highlights.append(f"训练结构较为完整：同步训练占比{syn_pct}%，专项训练占比{sub_pct}%，模拟训练占比{mon_pct}%，能够较好支撑日常教学、专项提升与阶段检测的衔接")
    elif syn_pct >= 50:
        highlights.append(f"日常训练基础较为扎实：同步训练占比{syn_pct}%，朗读跟读等常规训练已较好融入课堂教学过程，有助于巩固发音基础与语感培养")

    # 亮点4：标杆班级
    if top and top[0]['avg_score'] >= 75:
        highlights.append(f"典型班级示范作用较为突出：{tc_name}平均得分率为{top[0]['avg_score']}%，位居全校前列，可为同年级教学改进提供参考样本")

    # ── 题型数据增强亮点（可选）──────────────────────────────
    if data.get('has_question_type'):
        qt_school = data.get('qt_school', {})
        qt_teacher = data.get('qt_teacher_rank', {})
        if qt_school:
            hardest = min(qt_school.items(), key=lambda x: x[1]['mean'])
            easiest = max(qt_school.items(), key=lambda x: x[1]['mean'])
            highlights.append(f"题型表现差异较为清晰：{easiest[0]}得分率最高（{easiest[1]['mean']}%），整体表现相对稳定；{hardest[0]}得分率相对较低（{hardest[1]['mean']}%，低分率{data.get('qt_hr_lr', {}).get(hardest[0], {}).get('低分率', 'N/A')}%），可作为后续专项改进的重点方向")
        if qt_teacher:
            best_t = max(qt_teacher.items(), key=lambda x: x[1])
            worst_t = min(qt_teacher.items(), key=lambda x: x[1])
            if best_t[0] != worst_t[0]:
                highlights.append(f"教师间教学结果存在一定差异：{best_t[0]}所带班级综合得分率为{best_t[1]}%，较{worst_t[0]}（{worst_t[1]}%）高{best_t[1]-worst_t[1]:.1f}个百分点，相关经验可纳入校本教研交流范畴")

    # ── 建议推导（与问题一一对应）────────────────────────────
    suggestions = []
    if any('同步' in iss for iss in issues) or any('实战' in iss for iss in issues):
        suggestions.append(f"优化训练结构：适度提升专项训练与模拟训练频次（建议各占总量15%~20%），进一步完善日常训练与阶段检测相结合的实施体系")
    if any('不均衡' in iss or '差异' in iss for iss in issues):
        suggestions.append(f"加强分层推进：重点关注完成率偏低班级，结合具体原因开展分类指导，推动年级内应用质量更加均衡")
    if any('自主' in iss for iss in issues):
        suggestions.append(f"强化自主学习引导：通过过程性评价与正向激励相结合的方式，提升词汇自主练习覆盖率，逐步培育学生自主巩固习惯")
    if any('完成率' in iss for iss in issues):
        suggestions.append(f"强化作业过程管理：建立作业完成率监测与预警机制，对持续偏低班级开展原因分析与定向改进")
    # 保底建议（必定有）
    if len(suggestions) == 0:
        suggestions.append(f"深化数据应用：持续跟踪核心指标变化，定期复盘教学策略，将数据分析结果更好融入校本教研与教学改进过程")

    # ── 输出 6.1 主要亮点 ──────────────────────────────────
    L.append("### 6.1 主要亮点\n\n")
    for i, h in enumerate(highlights, 1):
        L.append(f"**亮点{i}：{h}**\n\n")

    # ── 输出 6.2 问题与不足 ───────────────────────────────
    L.append("### 6.2 问题与不足\n\n")
    if issues:
        for issue in issues:
            L.append(f"- **{issue}**\n\n")
    else:
        L.append("本期应用中未发现突出结构性问题，各项指标总体平稳。\n\n")

    # ── 输出 6.3 下阶段计划 ────────────────────────────────
    L.append("### 6.3 下阶段计划\n\n")
    for s in suggestions:
        L.append(f"- {s}\n")
    L.append(f"- **听说专项提升**：适度增加听说模拟套卷训练频次，引导学生结合答题录音开展自我诊断与纠音练习，提升训练的针对性与实效性\n\n")

    return ''.join(L)


def make_charts(data):
    charts = {}
    FONT_FAMILY = 'Noto Sans CJK SC, Microsoft YaHei, SimHei, Arial Unicode MS, sans-serif'
    GC = {'六年级': '#4C78A8', '七年级': '#F58518', '八年级': '#E45756'}
    CC = {'同步': '#4C78A8', '专项': '#F58518', '模拟': '#E45756', '课外拓展': '#72D7B8'}
    cats = ['同步', '专项', '模拟', '课外拓展']
    months = sorted(data.get('monthly_hw', {}).keys())

    totals = [data['monthly_hw'].get(m, 0) for m in months]
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=months, y=totals, mode='lines+markers+text',
        line=dict(color='#2E86AB', width=3),
        marker=dict(size=10, color='#2E86AB'),
        fill='tozeroy', fillcolor='rgba(46,134,171,0.1)',
        text=totals, textposition='top center', textfont=dict(size=11, family=FONT_FAMILY, color='#1f2937'),
        name='作业总量', hovertemplate='%{x}<br>作业量：%{y}次<extra></extra>'
    ))
    fig.update_layout(
        title=None,
        xaxis_title='月份', yaxis_title='作业次数',
        height=460, template='plotly_white', hovermode='x unified',
        font=dict(size=13, color='#1f2937', family=FONT_FAMILY),
        margin=dict(l=70, r=40, t=90, b=70),
        xaxis=dict(tickangle=0, automargin=True),
        yaxis=dict(automargin=True)
    )
    charts['monthly_line'] = fig

    grade_hw = data.get('grade_monthly_hw', {})
    fig2 = go.Figure()
    for grade in ['六年级', '七年级', '八年级']:
        gd = grade_hw.get(grade, {})
        y = [gd.get(m, 0) for m in months]
        fig2.add_trace(go.Scatter(
            name=grade, x=months, y=y,
            mode='lines+markers+text', line=dict(width=2.5),
            marker=dict(size=7, color=GC.get(grade, '#999')),
            text=y, textposition='top center', textfont=dict(size=10, family=FONT_FAMILY, color='#1f2937'),
            cliponaxis=False
        ))
    fig2.update_layout(
        title=None,
        xaxis_title='月份', yaxis_title='作业次数',
        height=440, template='plotly_white',
        font=dict(size=13, color='#1f2937', family=FONT_FAMILY),
        margin=dict(l=70, r=40, t=90, b=70),
        xaxis=dict(automargin=True),
        yaxis=dict(automargin=True)
    )
    charts['grade_monthly_line'] = fig2

    fig3 = go.Figure()
    for cat in cats:
        y = [data.get('cat_monthly', {}).get(m, {}).get(cat, 0) for m in months]
        fig3.add_trace(go.Bar(
            name=cat, x=months, y=y, marker_color=CC.get(cat, '#999'),
            text=y, textposition='inside', textfont=dict(size=10, color='#1f2937', family=FONT_FAMILY), cliponaxis=False
        ))
    fig3.update_layout(
        barmode='stack',
        title=None,
        xaxis_title='月份', yaxis_title='作业次数',
        height=460, template='plotly_white',
        font=dict(size=13, color='#1f2937', family=FONT_FAMILY),
        legend=dict(
            orientation='v',
            x=1.02, xanchor='left',
            y=1, yanchor='top',
            font=dict(size=12, color='#1f2937', family=FONT_FAMILY)
        ),
        margin=dict(l=70, r=120, t=40, b=70),
        xaxis=dict(automargin=True),
        yaxis=dict(automargin=True)
    )
    charts['cat_stacked'] = fig3

    cat_pct = data.get('category_pct', {})
    fig4 = go.Figure()
    fig4.add_trace(go.Pie(
        labels=list(cat_pct.keys()),
        values=list(cat_pct.values()),
        marker_colors=[CC.get(c, '#999') for c in cat_pct.keys()],
        textinfo='none',
        hovertemplate='%{label}<br>%{percent}<extra></extra>',
        hole=0.35,
        sort=False,
        showlegend=True
    ))
    fig4.update_layout(
        title=None,
        height=460, template='plotly_white',
        font=dict(size=13, color='#1f2937', family=FONT_FAMILY),
        legend=dict(
            orientation='v',
            x=1.02, xanchor='left',
            y=0.95, yanchor='top',
            font=dict(size=12, color='#1f2937', family=FONT_FAMILY)
        ),
        margin=dict(l=40, r=180, t=30, b=40)
    )
    charts['cat_pie'] = fig4

    mock_scores = data.get('mock_hw_score_monthly', {})
    mock_counts = data.get('mock_hw_count_monthly', {})
    if mock_scores:
        ms = sorted(mock_scores.keys())
        sc = [mock_scores[m] for m in ms]
        mc = [mock_counts.get(m, 0) for m in ms]
        fig5 = make_subplots(specs=[[{"secondary_y": True}]])
        # 柱状图：月均作业次数（典型班级）
        fig5.add_trace(go.Bar(
            x=ms, y=mc, name='作业数',
            marker_color='#F58518', opacity=0.6,
            text=mc, textposition='outside', textfont=dict(size=10, color='#1f2937', family=FONT_FAMILY),
            yaxis='y2'
        ))
        # 折线图：月均得分率趋势
        fig5.add_trace(go.Scatter(
            x=ms, y=sc, name='得分率',
            mode='lines+markers+text',
            line=dict(color='#E45756', width=2.5), marker=dict(size=8, color='#E45756'),
            text=[f"{s}%" for s in sc], textposition='top center', textfont=dict(size=10, color='#1f2937', family=FONT_FAMILY),
            yaxis='y'
        ))
        fig5.update_layout(
            title=None,
            template='plotly_white', height=430,
            font=dict(size=13, color='#1f2937', family=FONT_FAMILY),
            legend=dict(
                orientation='v',
                x=1.02, xanchor='left',
                y=1, yanchor='top',
                font=dict(size=12, color='#1f2937', family=FONT_FAMILY)
            ),
            hovermode='x unified',
            margin=dict(l=70, r=120, t=40, b=70),
            xaxis=dict(automargin=True),
            yaxis=dict(automargin=True)
        )
        fig5.update_layout(yaxis2=dict(title_text='作业次数', overlaying='y', side='right'))
        fig5.update_yaxes(title_text='得分率（%）', range=[0, 100])
        charts['mock_score'] = fig5

    grade_scores = data.get('mock_hw_grade_monthly', {})
    if grade_scores:
        fig6 = go.Figure()
        for grade in sorted(grade_scores.keys()):
            gm = sorted(grade_scores[grade].items())
            xs = [m for m, s in gm]
            ys = [s for m, s in gm]
            fig6.add_trace(go.Scatter(
                name=grade, x=xs, y=ys,
                mode='lines+markers+text', line=dict(width=2.5),
                marker=dict(size=7, color=GC.get(grade, '#999')),
                text=[f"{v}%" for v in ys], textposition='top center', textfont=dict(size=10, family=FONT_FAMILY, color='#1f2937'),
                cliponaxis=False
            ))
        fig6.update_layout(
            title=None,
            xaxis_title='月份', yaxis_title='得分率（%）',
            height=430, template='plotly_white',
            font=dict(size=13, color='#1f2937', family=FONT_FAMILY),
            yaxis=dict(range=[0, 100], automargin=True),
            margin=dict(l=70, r=40, t=90, b=70),
            xaxis=dict(automargin=True)
        )
        charts['grade_score'] = fig6

    top_all_m = data.get('top_class_all_monthly', {})
    top_mock_m = data.get('top_class_mock_monthly', {})
    if top_all_m and top_mock_m:
        all_months = sorted(set(top_all_m.keys()) | set(top_mock_m.keys()))
        sc_t  = [top_mock_m.get(m, {}).get('score', None) for m in all_months]
        ct_t  = [top_all_m.get(m, 0) for m in all_months]
        sc_t_fmt = [f"{s}%" if s is not None else '—' for s in sc_t]

        fig7 = make_subplots(specs=[[{"secondary_y": True}]])
        fig7.add_trace(go.Scatter(
            x=all_months, y=sc_t, name='听说模拟得分率',
            mode='lines+markers+text',
            line=dict(color='#4C78A8', width=2.5), marker=dict(size=8),
            text=sc_t_fmt, textposition='top center', textfont=dict(size=10, family=FONT_FAMILY, color='#1f2937'),
            yaxis='y'
        ))
        fig7.add_trace(go.Bar(
            x=all_months, y=ct_t, name='所有类目布置次数',
            opacity=0.35, marker_color='#F58518', yaxis='y2',
            text=ct_t, textposition='outside', textfont=dict(size=10, family=FONT_FAMILY, color='#1f2937')
        ))
        fig7.update_layout(
            title=None,
            template='plotly_white', height=430,
            font=dict(size=13, color='#1f2937', family=FONT_FAMILY),
            legend=dict(
                orientation='v',
                x=1.02, xanchor='left',
                y=1, yanchor='top',
                font=dict(size=12, color='#1f2937', family=FONT_FAMILY)
            ),
            hovermode='x unified',
            margin=dict(l=70, r=120, t=40, b=70),
            xaxis=dict(automargin=True),
            yaxis=dict(automargin=True)
        )
        fig7.update_layout(yaxis2=dict(title_text='布置次数', overlaying='y', side='right'))
        fig7.update_yaxes(title_text='得分率（%）', range=[0, 100])
        charts['top_class_trend'] = fig7

    return charts
def export_to_docx(report_md: str, charts: dict = None) -> tuple:
    """导出为公文格式Word

    格式要求：
    - 标题：方正小标宋简体，二号(22pt)，居中
    - 一级标题：黑体，三号(16pt)
    - 二级标题：楷体_GB2312，三号(16pt)
    - 正文：仿宋_GB2312，三号(16pt)，首行缩进2字符
    - 行间距：固定值31磅
    - 页边距：上3.7cm、下3.5cm、左2.8cm、右2.6cm
    - 表格：无边框线，宋体五号(10.5pt)，列宽紧凑
    - 图表：无间隔，居中，宽500px×高250px
    """
    try:
        from docx import Document
        from docx.shared import Pt, Cm, Emu
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
    except ImportError:
        return None, "python-docx未安装"

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(3.5)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)

    # ── 辅助函数 ──────────────────────────────────────────────

    def set_font(run, fname, fsize, bold=False):
        run.font.name = fname
        run.font.size = Pt(fsize)
        run.font.bold = bold
        try:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), fname)
        except Exception:
            pass

    def para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=False, space_before=0, space_after=0,
                  line_spacing=31):
        para.alignment = align
        pf = para.paragraph_format
        if first_indent:
            pf.first_line_indent = Pt(32)
        pf.space_before = Pt(space_before)
        pf.space_after = Pt(space_after)
        pf.line_spacing = Pt(line_spacing)

    def add_para(text, fname='仿宋_GB2312', fsize=16, bold=False,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  first_indent=False, space_before=0, space_after=0):
        p = doc.add_paragraph()
        para_fmt(p, align, first_indent, space_before, space_after, 31)
        r = p.add_run(text)
        set_font(r, fname, fsize, bold)
        return p

    def remove_table_borders(tbl):
        """去掉表格所有边框线"""
        tblPr = tbl._tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl._tbl.insert(0, tblPr)
        tblBorders = OxmlElement('w:tblBorders')
        for btype in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
            b = OxmlElement(f'w:{btype}')
            b.set(qn('w:val'), 'none')
            tblBorders.append(b)
        tblPr.append(tblBorders)

    def add_border_table(headers, rows_data):
        """有边框线表格：黑体五号(10.5pt)，紧凑列宽，防御性列数对齐"""
        # 统一列数：取所有行的最大列数，不足者在末尾补空字符串
        all_rows_raw = [headers] + list(rows_data)
        ncol = max(len(row) for row in all_rows_raw) if all_rows_raw else 1
        def pad_row(row):
            return list(row) + [''] * (ncol - len(row))
        headers_padded = pad_row(headers)
        rows_padded    = [pad_row(r) for r in rows_data]

        tbl = doc.add_table(rows=1+len(rows_padded), cols=ncol)
        tbl.style = 'Table Grid'
        FNAME = '宋体'; FSIZE = 10.5

        def fill_cell(cell, text, center=True, bold=False):
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_font(r, FNAME, FSIZE, bold)

        for ci, h in enumerate(headers_padded):
            fill_cell(tbl.rows[0].cells[ci], h, center=True, bold=True)
        for ri, row in enumerate(rows_padded):
            for ci, val in enumerate(row):
                fill_cell(tbl.rows[ri+1].cells[ci], str(val), center=True, bold=False)

        # 紧凑列宽（基于填入后的表格实际内容计算）
        all_content = [headers_padded] + rows_padded
        col_widths = []
        for ci in range(ncol):
            max_len = max(len(str(row[ci])) for row in all_content)
            width_cm = max(1.5, min(max_len * 0.5 + 0.4, 10))
            col_widths.append(Cm(width_cm))
        for ri, row in enumerate(tbl.rows):
            for ci, cell in enumerate(row.cells):
                cell.width = col_widths[ci]

        # 居中：整表水平居中（通过段落对齐实现）
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        return tbl

    def add_chart_image(key, title, width=Cm(13), height=Cm(6.5)):
        """以无间隔居中图片方式插入图表"""
        if not charts or key not in charts:
            return
        fig = charts[key]
        try:
            img_bytes = fig.to_image(format='png', width=1000, height=500, scale=2)
        except Exception:
            note = doc.add_paragraph()
            para_fmt(note, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_before=0, space_after=6, line_spacing=24)
            note_run = note.add_run(f"【{title}未导出：当前环境缺少图表导出依赖 kaleido】")
            set_font(note_run, '宋体', 10.5, False)
            return
        img_io = BytesIO(img_bytes)
        p = doc.add_paragraph()
        para_fmt(p, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_before=0, space_after=0, line_spacing=0)
        run = p.add_run()
        run.add_picture(img_io, width=width, height=height)
        # 图注
        cap = doc.add_paragraph()
        para_fmt(cap, align=WD_ALIGN_PARAGRAPH.CENTER, first_indent=False, space_before=0, space_after=6, line_spacing=31)
        cap_run = cap.add_run(title)
        set_font(cap_run, '宋体', 10.5, False)

    # ── Markdown解析与Word构建 ────────────────────────────────
    lines = report_md.split('\n')
    i = 0
    active_section = None       # 当前节编号如'四、'
    section_had_table = False   # 当前节是否已渲染过表格
    pending_charts = {}         # 当前节待插入图表 {key: caption}

    CHART_MAP = {
        '三、': {},
        '四、': {
            'mock_score':    '图5  听说模拟类月均得分率趋势',
        },
        '五、': {
            'top_class_trend': '图7  标杆班级月度作业量与得分率组合图',
        },
    }

    def flush_section_charts():
        """将当前节所有待插图表插入文档，并标记已处理"""
        global section_had_table
        if pending_charts:
            # 图表与上文（通常是表格）之间空一行
            gap = doc.add_paragraph()
            para_fmt(gap, space_before=0, space_after=0, line_spacing=0)
        for key, caption in list(pending_charts.items()):
            if charts and key in charts:
                add_chart_image(key, caption, width=Cm(13), height=Cm(6.5))
        pending_charts.clear()
        section_had_table = True   # 标记已处理，防止重复插入

    while i < len(lines):
        line = lines[i].strip()

        # 跳过注释行和页脚
        if not line or line.startswith('>') or line.startswith('*数据') or line.startswith('*报告') or line.startswith('*生成时间'):
            i += 1; continue

        # ── 主标题 ────────────────────────────────────────────
        if line.startswith('# ') and '成效报告' in line:
            title_text = line.replace('# ', '').replace('**', '').strip()
            add_para(title_text, '方正小标宋简体', 22, True,
                     WD_ALIGN_PARAGRAPH.CENTER, space_before=12, space_after=12)
            i += 1; continue

        # ── 一级标题（## xxx）───────────────────────────────
        if line.startswith('## '):
            section_text = line.replace('## ', '').strip()
            # 先加载新课明的pending，再把上一节的pending flush出去
            # 这样图表出现在前一节末尾 + 新节内容之前
            new_pending = dict(CHART_MAP.get(section_text[:2], {}))
            if active_section and active_section in CHART_MAP and pending_charts:
                flush_section_charts()
            active_section = section_text[:2]
            section_had_table = False
            pending_charts = new_pending
            add_para(section_text, '黑体', 16, True,
                     WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6)
            i += 1; continue

        # ── 二级标题（### xxx）─────────────────────────────
        if line.startswith('### '):
            sub_text = line.replace('### ', '').strip()
            add_para(sub_text, '楷体_GB2312', 16, True,
                     WD_ALIGN_PARAGRAPH.LEFT, first_indent=False,
                     space_before=6, space_after=3)

            # 按小节精确控制图表位置
            if sub_text == '3.3 应用频次分析':
                if charts and 'monthly_line' in charts:
                    add_chart_image('monthly_line', '图1  月度作业总量趋势', width=Cm(13), height=Cm(6.5))
                if charts and 'grade_monthly_line' in charts:
                    add_chart_image('grade_monthly_line', '图2  各年级月度作业量趋势', width=Cm(13), height=Cm(6.5))
            elif sub_text == '3.4 应用方式分析':
                # 用户要求图3与图4顺序对调：正文内先图4后图3
                if charts and 'cat_pie' in charts:
                    add_chart_image('cat_pie', '图4  作业类型占比分布', width=Cm(13), height=Cm(6.5))
                if charts and 'cat_stacked' in charts:
                    add_chart_image('cat_stacked', '图3  各月各类作业量分布', width=Cm(13), height=Cm(6.5))
            elif sub_text == '4.1 成绩数据对比':
                if charts and 'grade_score' in charts:
                    add_chart_image('grade_score', '图6  各年级听说模拟得分率趋势', width=Cm(13), height=Cm(6.5))
            i += 1; continue

        # ── 段落（处理内联加粗）────────────────────────────────
        if line and not line.startswith('|') and not line.startswith('- ') and not line.startswith('```'):
            # 用 finditer 构建段落segments，避免 re.split 产生的空档问题
            segments = []   # [(text, bold), ...]
            last_end = 0
            for m in re.finditer(r'\*\*(.+?)\*\*', line):
                if m.start() > last_end:
                    segments.append((line[last_end:m.start()], False))
                segments.append((m.group(1), True))
                last_end = m.end()
            if last_end < len(line):
                segments.append((line[last_end:], False))

            if any(b for _, b in segments):
                p = doc.add_paragraph()
                para_fmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, space_before=0, space_after=3, line_spacing=31)
                for text, bold in segments:
                    r = p.add_run(text)
                    set_font(r, '仿宋_GB2312', 16, bold)
            else:
                clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line).strip()
                if clean:
                    add_para(clean, '仿宋_GB2312', 16, False,
                             WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=True, space_before=0, space_after=3)
            i += 1; continue

        # ── 列表项（处理加粗）────────────────────────────────
        if line.startswith('- '):
            clean = re.sub(r'\*\*(.+?)\*\*', r'\1', line).lstrip('- ')
            p = doc.add_paragraph(style='List Bullet')
            para_fmt(p, WD_ALIGN_PARAGRAPH.JUSTIFY, first_indent=False, space_before=0, space_after=2, line_spacing=31)
            parts = re.split(r'(\*\*(.+?)\*\*)', clean)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    set_font(r, '仿宋_GB2312', 16, True)
                elif part:
                    r = p.add_run(part)
                    set_font(r, '仿宋_GB2312', 16, False)
            i += 1; continue

        # ── 表格（有边框，宋体五号，紧凑列宽）────────────────
        if line.startswith('|') and '---' not in line:
            rows_data = []
            j = i
            while j < len(lines) and lines[j].strip().startswith('|'):
                if '---' not in lines[j]:
                    cells = [re.sub(r'\*\*(.+?)\*\*', r'\1', c.strip()) for c in lines[j].strip().split('|')[1:-1]]
                    rows_data.append(cells)
                j += 1
            if rows_data:
                add_border_table(rows_data[0], rows_data[1:])
                section_had_table = True
                # 该节所有图表在最后一个表格之后立即插入
                if pending_charts:
                    flush_section_charts()
            i = j; continue

        i += 1

    # 最后一节如有剩余图表，插入末尾
    if active_section and active_section in CHART_MAP and pending_charts:
        flush_section_charts()

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf, None


